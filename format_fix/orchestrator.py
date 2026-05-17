"""Format-Fix Pipeline orchestrator.

Walks every PDF page, asks each registered SectionHandler in priority order
whether it claims the page, and dispatches to the first match. Each handler
mutates a shared python-docx Document. After all pages are processed, a
final font-discipline pass enforces TNR / 12pt-body / 14pt-bold-heading /
black-text-only.

CLI:
    python -m format_fix.orchestrator <pdf> <university_id> <output_docx> \
        [--skip-pages 1,3,5]

Library:
    from format_fix.orchestrator import run
    run(pdf_path, university_id, output_path, skip_pages=[3])
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

import fitz
from docx import Document

from . import extraction
from .context import Context
from .styles import apply_university_styles, enforce_font_discipline

from .handlers.base            import SectionHandler
from .handlers.cover           import CoverHandler
from .handlers.acknowledgement import AcknowledgementHandler
from .handlers.declaration     import DeclarationHandler
from .handlers.abstract        import AbstractHandler
from .handlers.toc             import TableOfContentsHandler
from .handlers.survey_table    import SurveyTableHandler
from .handlers.form_page       import FormPageHandler
from .handlers.references      import ReferencesHandler
from .handlers.annexures       import AnnexuresHandler
from .handlers.chapter         import ChapterHandler


def _build_handlers() -> list[SectionHandler]:
    """Registered handlers, sorted by priority (lower = checked first).

    Ordering matters because applies_to() decides who claims a page; the
    catch-all ChapterHandler must be last.
    """
    handlers: list[SectionHandler] = [
        SurveyTableHandler(),       # 10
        FormPageHandler(),          # 20
        CoverHandler(),             # 30
        AcknowledgementHandler(),   # 40
        DeclarationHandler(),       # 50
        AbstractHandler(),          # 60
        TableOfContentsHandler(),   # 70
        ReferencesHandler(),        # 80
        AnnexuresHandler(),         # 85
        ChapterHandler(),           # 100  catch-all
    ]
    handlers.sort(key=lambda h: h.priority)
    return handlers


def _render_chapter_inline(doc, elements, ctx) -> None:
    """DOCX-specific chapter renderer with interleaved tables/images.

    Replicates the body-paragraph rendering of
    handlers.chapter.ChapterHandler.render() but processes an
    interleaved element stream produced by
    extraction_docx.parse_docx_to_pages():

        elements = [
            ("p",   block_5tuple),
            ("p",   block_5tuple),
            ("tbl", TableRows),         # << source-position preserved
            ("p",   block_5tuple),
            ("img", (blob, content_type)),
            ...
        ]

    This is the Step 4.6e fix for "tables and charts consolidated
    together." Calling chapter.render() with just the paragraph blocks
    and appending tables after dumped all of a page's tables at the
    end of its text. By walking the source-order stream we emit each
    table / image at the exact paragraph boundary where it originally
    appeared.

    The chapter handler module itself is NOT modified (Rule 1). We
    reuse its private helpers `_classify` and `_flush_kv` as
    collaborators -- the underscore is a Python convention, not
    enforcement, and these helpers are intentionally factored out so
    they can be shared.

    DOCX-specific simplifications vs chapter.py:
      * Skip extraction.merge_body_blocks / merge_bullet_continuations
        -- the PDF "Canva visual line break" problem doesn't exist in
        DOCX; each python-docx paragraph is already one logical block.
      * Skip post-loop image embedding via ctx.pdf -- DOCX images come
        in inline as ("img", blob) elements, not as fitz xrefs.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from . import extraction
    from . import extraction_docx
    from .handlers import chapter as _chapter

    # First-pass scan: does this page contain any "real h1" heading?
    # Same in_toc_section gating as chapter.render -- a page with no
    # h1 inside a ToC region is suppressed so the auto-TOC doesn't get
    # echoed in the body.
    has_real_h1 = False
    for kind, payload in elements:
        if kind != "p":
            continue
        text, max_sz, _dom, _bold, _align = payload
        if _chapter._classify((text or "").strip(), max_sz, ctx.body_pt, 0) == "h1":
            has_real_h1 = True
            break

    if has_real_h1:
        ctx.in_toc_section = False
    elif ctx.in_toc_section:
        # Body text is suppressed, but customer data (tables / images)
        # must still reach the output -- emit those defensively.
        for kind, payload in elements:
            if kind == "tbl":
                extraction_docx.emit_table(doc, payload)
            elif kind == "img":
                blob, ctype = payload
                extraction_docx.emit_image(doc, blob, ctype)
        return

    kv_buffer: list[tuple[str, str]] = []

    def _flush_kv_if_any() -> None:
        nonlocal kv_buffer
        if kv_buffer:
            _chapter._flush_kv(doc, kv_buffer)
            kv_buffer = []

    for kind, payload in elements:
        if kind == "tbl":
            _flush_kv_if_any()
            extraction_docx.emit_table(doc, payload)
            continue
        if kind == "img":
            _flush_kv_if_any()
            blob, ctype = payload
            extraction_docx.emit_image(doc, blob, ctype)
            continue

        # kind == "p"
        text, max_sz, _dom, _bold, _align = payload
        ln = (text or "").strip()
        if not ln:
            continue
        if extraction.is_stray_line(ln):
            continue
        if extraction.TOC_LEADER_RE.search(ln):
            continue

        # Key-value detection (same logic as chapter.render -- tightened
        # to avoid prose lines like "Very Satisfied (45.8%): 55 ..." being
        # mistaken for K:V form fields).
        kv_m = extraction.KV_RE.match(ln)
        looks_like_prose_kv = (
            kv_m is not None
            and (
                "%" in kv_m.group(1)
                or "(" in kv_m.group(1)
                or len(kv_m.group(2)) > 80
                or len(ln) > 120
            )
        )
        if (kv_m
                and not extraction.ALL_CAPS_RE.match(ln)
                and not looks_like_prose_kv):
            kv_buffer.append((kv_m.group(1), kv_m.group(2)))
            continue
        elif kv_buffer:
            _flush_kv_if_any()

        cls = _chapter._classify(ln, max_sz, ctx.body_pt, 0)
        if cls == "skip":
            continue

        # DOCX-style heading override (Step 4.6d).
        # _chapter._classify gates h2 on ALL_CAPS_RE which fails on
        # numbered sub-headings like "3.1 Sampling Design" or
        # "4.1 Data Analysis and Interpretation" -- they fall through
        # to plain body and lose bold + the Heading 2 paragraph style.
        # extraction_docx bumps Heading 1 -> 22pt, Heading 2 -> 16pt,
        # Heading 3 -> 14pt, so the source paragraph's max_size is a
        # reliable signal of its intended heading level. Use those
        # thresholds directly when the underlying paragraph was bold
        # (genuine headings nearly always are) so we don't accidentally
        # promote a body paragraph that happens to have an explicit
        # 14pt run.
        if _bold:
            if max_sz >= 20:
                cls = "h1"
            elif max_sz >= 16:
                cls = "h2"
            elif max_sz >= 14 and cls == "body":
                cls = "h3"

        if cls == "h1":
            doc.add_page_break()
            h = doc.add_heading(ln.strip(" ?:"), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if cls == "h2":
            h = doc.add_heading(ln.strip(" ?:"), level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if cls == "h3":
            h = doc.add_heading(ln.strip(" ?:"), level=3)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if extraction.BULLET_RE.match(ln):
            cleaned = extraction.BULLET_RE.sub("", ln).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(cleaned)
            continue

        if extraction.NUM_LIST_RE.match(ln):
            cleaned = extraction.NUM_LIST_RE.sub("", ln).strip()
            p = doc.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(cleaned)
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(ln)

    _flush_kv_if_any()


def run(pdf_path: str | Path,
        university_id: str,
        output_path: str | Path,
        skip_pages: list[int] | None = None,
        verbose: bool = False) -> dict:
    """Build a formatted DOCX from a PDF using section-handler dispatch.

    Returns a small summary dict with counts (pages_processed, claims by
    handler name) so callers can log/display what each agent did.
    """
    pdf_path    = Path(pdf_path)
    output_path = Path(output_path)
    pdf         = fitz.open(str(pdf_path))

    try:
        ctx = Context.build(pdf, university_id, skip_pages=skip_pages)
        doc = Document()
        apply_university_styles(doc, ctx.config)

        handlers = _build_handlers()
        claims: dict[str, int] = {h.name: 0 for h in handlers}
        claims["__skipped__"] = 0
        claims["__nohandler__"] = 0

        for page_no in range(len(pdf)):
            if page_no in ctx.skip_set:
                claims["__skipped__"] += 1
                if verbose:
                    print(f"  p{page_no+1}: SKIPPED (user requested)")
                continue

            blocks = extraction.blocks_with_font(pdf, page_no)

            picked: SectionHandler | None = None
            for h in handlers:
                if h.applies_to(blocks, page_no, ctx):
                    picked = h
                    break

            if picked is None:
                claims["__nohandler__"] += 1
                if verbose:
                    print(f"  p{page_no+1}: no handler claimed (skipped)")
                continue

            if verbose:
                print(f"  p{page_no+1}: -> {picked.name}")
            picked.render(doc, blocks, page_no, ctx)
            claims[picked.name] += 1

        enforce_font_discipline(doc, ctx)
        doc.save(str(output_path))
        return {
            "output":  str(output_path),
            "pages":   len(pdf),
            "claims":  claims,
        }
    finally:
        pdf.close()


def run_from_docx(docx_bytes: bytes,
                  university_id: str,
                  output_path: str | Path,
                  skip_pages: list[int] | None = None,
                  verbose: bool = False) -> dict:
    """DOCX entry point — parallel to run() for PDF input.

    Parses the DOCX into virtual pages of 5-tuple blocks (matching the
    PDF blocks_with_font contract), then dispatches the existing handler
    pipeline over those pages. A blank fitz.Document with the same page
    count is created so ctx.pdf remains valid for handlers that touch
    it — image extraction in chapter/form_page handlers degrades to a
    no-op on blank pages, which is acceptable since DOCX images aren't
    re-embedded in this pipeline.

    Returns the same summary dict shape as run() plus source="docx".
    """
    from . import extraction_docx

    output_path = Path(output_path)

    pages_blocks, pages_tables, pages_images, pages_elements, n_pages = \
        extraction_docx.parse_docx_to_pages(docx_bytes)
    if n_pages == 0:
        n_pages = 1
        pages_blocks   = [[]]
        pages_tables   = [[]]
        pages_images   = [[]]
        pages_elements = [[]]

    # Synthesize a blank fitz.Document with the same page count so the
    # existing Context.build + handlers continue to function. They will
    # find no embedded images, which is fine for DOCX input.
    pdf = fitz.open()
    try:
        for _ in range(n_pages):
            pdf.new_page()  # A4 default

        ctx = Context.build(pdf, university_id, skip_pages=skip_pages)
        doc = Document()
        apply_university_styles(doc, ctx.config)

        # Skip FormPageHandler for DOCX input: it was designed for scanned
        # PDF survey/evaluation pages and renders the page as a PNG via
        # ctx.pdf. For DOCX, ctx.pdf is a blank shim — rendering it would
        # produce blank PNGs and silently drop the text content (e.g. the
        # "TRAINEE EVALUATION" pages of aswathy_Project.docx, which is
        # how 62% of body content was being lost). The other handlers
        # consume the 5-tuple blocks and work identically for DOCX
        # virtual pages.
        handlers = [h for h in _build_handlers() if h.name != "form_page"]
        claims: dict[str, int] = {h.name: 0 for h in handlers}
        claims["__skipped__"]   = 0
        claims["__nohandler__"] = 0

        for page_no in range(n_pages):
            if page_no in ctx.skip_set:
                claims["__skipped__"] += 1
                if verbose:
                    print(f"  p{page_no+1}: SKIPPED (user requested)")
                continue

            blocks = pages_blocks[page_no]

            picked: SectionHandler | None = None
            for h in handlers:
                if h.applies_to(blocks, page_no, ctx):
                    picked = h
                    break

            if picked is None:
                claims["__nohandler__"] += 1
                if verbose:
                    print(f"  p{page_no+1}: no handler claimed (skipped)")
                # Tables / images on a no-handler page would still be
                # lost; emit them defensively so customer data isn't
                # silently dropped.
                for tbl_rows in pages_tables[page_no]:
                    extraction_docx.emit_table(doc, tbl_rows)
                for blob, ctype in pages_images[page_no]:
                    extraction_docx.emit_image(doc, blob, ctype)
                continue

            if verbose:
                print(f"  p{page_no+1}: -> {picked.name}")

            if picked.name == "chapter":
                # Step 4.6e: walk the interleaved element stream so
                # tables and images sit at their original source
                # positions in the rendered text, not all clumped at
                # the end of the section. The chapter handler module
                # is NOT modified; _render_chapter_inline reuses its
                # _classify and _flush_kv helpers.
                _render_chapter_inline(doc, pages_elements[page_no], ctx)
                claims["chapter"] += 1
            else:
                # Non-chapter handlers (cover / ack / decl / refs /
                # annex / toc) keep their existing PDF-style render.
                # Tables and images in those sections are rare and we
                # still emit them defensively after the handler so
                # nothing customer-supplied is lost.
                picked.render(doc, blocks, page_no, ctx)
                claims[picked.name] += 1

                for tbl_rows in pages_tables[page_no]:
                    extraction_docx.emit_table(doc, tbl_rows)
                for blob, ctype in pages_images[page_no]:
                    extraction_docx.emit_image(doc, blob, ctype)

        enforce_font_discipline(doc, ctx)
        doc.save(str(output_path))
        return {
            "output":  str(output_path),
            "pages":   n_pages,
            "claims":  claims,
            "source":  "docx",
        }
    finally:
        pdf.close()


def _main() -> int:
    p = argparse.ArgumentParser(
        prog="python -m format_fix.orchestrator",
        description="Section-handler PDF -> DOCX pipeline.",
    )
    p.add_argument("pdf",     type=str, help="path to source PDF")
    p.add_argument("uni",     type=str, help="university config id (e.g. ignou)")
    p.add_argument("out",     type=str, help="path to output DOCX")
    p.add_argument("--skip-pages", type=str, default="",
                    help="comma-separated 1-based page numbers to drop")
    p.add_argument("--verbose", "-v", action="store_true")
    args = p.parse_args()

    skip_pages: list[int] | None = None
    if args.skip_pages.strip():
        skip_pages = [int(x) for x in args.skip_pages.split(",")
                       if x.strip().isdigit()]

    summary = run(args.pdf, args.uni, args.out,
                   skip_pages=skip_pages, verbose=args.verbose)

    print()
    print(f"OK  output : {summary['output']}")
    print(f"OK  pages  : {summary['pages']}")
    print( "    claims :")
    for name, n in summary["claims"].items():
        if n:
            print(f"      {name:14s} {n}")
    return 0


if __name__ == "__main__":
    sys.exit(_main())
