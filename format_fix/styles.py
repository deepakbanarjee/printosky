"""Document-level style infrastructure.

Two public functions:
  apply_university_styles(doc, config)
    Sets up the Normal and Heading 1/2/3 paragraph styles using the
    university config (body font, body size, heading sizes, line spacing).

  enforce_font_discipline(doc, ctx)
    Final pass that walks every run in every paragraph + every table cell
    and forces:
      - Times New Roman (or whatever ctx.body_font is)
      - body size for body paragraphs, +2pt bold for Heading-styled paras
      - Black text colour (RGB 0,0,0)
      - No italic, no underline
"""
from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

if TYPE_CHECKING:
    from docx import Document
    from .context import Context


# ---------------------------------------------------------------------------
# Style setup
# ---------------------------------------------------------------------------

def apply_university_styles(doc: "Document", config: dict) -> None:
    """Configure Normal + Heading 1/2/3 styles from the university config."""
    body_font = config.get("body_font", "Times New Roman")
    body_size = int(config.get("body_size_pt", 12))
    spacing   = float(config.get("line_spacing", 1.5))
    sp_before = int(config.get("para_space_before_pt", 0))
    sp_after  = int(config.get("para_space_after_pt", 6))

    normal = doc.styles["Normal"]
    normal.font.name = body_font
    normal.font.size = Pt(body_size)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = spacing
    pf.space_before      = Pt(sp_before)
    pf.space_after       = Pt(sp_after)

    for level in (1, 2, 3):
        hcfg = config.get(f"heading{level}", {})
        size = int(hcfg.get("size_pt", body_size + (4 - level)))
        bold = bool(hcfg.get("bold", True))
        try:
            hstyle = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        hstyle.font.name = body_font
        hstyle.font.size = Pt(size)
        hstyle.font.bold = bold
        hstyle.font.color.rgb = RGBColor(0, 0, 0)

    margins = config.get("margins") or {}
    for section in doc.sections:
        # Enforce A4 (210 x 297 mm) unconditionally — academic standard
        # for Indian universities. Without this, python-docx defaults to
        # US Letter (216 x 279 mm), which is wrong for our customers.
        # PDF inputs that were Letter-sized get re-paginated to A4 here,
        # which is the intended behaviour.
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)

        if "left_cm" in margins:
            section.left_margin = Cm(float(margins["left_cm"]))
        if "right_cm" in margins:
            section.right_margin = Cm(float(margins["right_cm"]))
        if "top_cm" in margins:
            section.top_margin = Cm(float(margins["top_cm"]))
        if "bottom_cm" in margins:
            section.bottom_margin = Cm(float(margins["bottom_cm"]))


# ---------------------------------------------------------------------------
# Final font discipline pass
# ---------------------------------------------------------------------------

def _force_run(run, *, font_name: str, size_pt: int,
                bold: bool | None = None) -> None:
    """Stamp a run with TNR + size + black + no italic/underline."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    font_name)
    rFonts.set(qn("w:hAnsi"),    font_name)
    rFonts.set(qn("w:cs"),       font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    run.font.italic    = False
    run.font.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    for color_el in rPr.findall(qn("w:color")):
        rPr.remove(color_el)
    new_color = OxmlElement("w:color")
    new_color.set(qn("w:val"), "000000")
    rPr.append(new_color)


def enforce_font_discipline(doc: "Document", ctx: "Context") -> None:
    """Final pass: every visible run becomes TNR body or heading-sized."""
    body_font = ctx.body_font
    body_size = ctx.body_size_pt
    head_size = ctx.heading_size_pt

    def _walk(para) -> None:
        sn = (para.style.name or "") if para.style else ""
        is_heading = sn.startswith("Heading")
        size = head_size if is_heading else body_size
        bold = True if is_heading else None
        for run in para.runs:
            _force_run(run, font_name=body_font, size_pt=size, bold=bold)

    for para in doc.paragraphs:
        _walk(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _walk(para)
