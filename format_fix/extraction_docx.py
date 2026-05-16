"""DOCX -> per-page blocks parallel of extraction.blocks_with_font().

Why this exists
---------------
The format_fix orchestrator was designed around PyMuPDF (`fitz`) and walks
PDF pages, asking each handler "is this your page?" by inspecting the
first few blocks. The handlers consume 5-tuples
`(text, max_size, dom_size, bold, align)` and don't care HOW they were
produced.

When the customer uploads a DOCX, there is no PDF, but python-docx
exposes paragraphs in document order with style + run-level font info.
We can synthesize the same 5-tuple stream from a DOCX and group it into
"virtual pages" so the existing handler-dispatch loop works unchanged.

How virtual pages are inferred
------------------------------
Real DOCX files may or may not have explicit page breaks. We split the
paragraph stream at:

1. Explicit page-break runs  (`<w:br w:type="page"/>`)
2. Section-title lines that match any handler's recognised heading
   (ACKNOWLEDGEMENT / DECLARATION / ABSTRACT / TABLE OF CONTENTS /
   REFERENCES / BIBLIOGRAPHY / WORKS CITED / ANNEXURE(S) /
   APPENDIX(ES) / CHAPTER N)

This produces a page layout the handlers' `applies_to(blocks, page_no, ctx)`
can dispatch over without any modification.

Font-size mapping
-----------------
python-docx exposes `run.font.size` only when explicitly set on the run.
For paragraphs that inherit size from the style (most student docs), we
fall back to bumped sizes based on `paragraph.style.name`:

  - "Title"          -> 24pt bold
  - "Heading 1"      -> 22pt bold   (must exceed body_pt + 8 so the
                                     chapter handler's _classify()
                                     returns "h1" -- anything <= 20pt
                                     falls through to "h2", which in
                                     turn lets the in_toc_section flag
                                     suppress entire body pages)
  - "Heading 2"      -> 16pt bold
  - "Heading 3"      -> 14pt bold
  - everything else  -> 12pt (the body-size default the orchestrator uses)

These values are signals to the dispatcher only -- they never reach
the rendered DOCX. apply_university_styles() rewrites the actual
"Heading 1/2/3" style fonts before save().
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


# 5-tuple shape:  (text, max_size, dom_size, bold, align)
Block = tuple[str, float, float, bool, str]


# Union of every handler's _TITLE_VARIANTS plus chapter detection.
# Hitting any of these in the paragraph stream starts a new virtual page.
_SECTION_TITLES: frozenset[str] = frozenset({
    "ACKNOWLEDGEMENT", "ACKNOWLEDGEMENTS", "ACKNOWLEDGMENT", "ACKNOWLEDGMENTS",
    "DECLARATION",
    "ABSTRACT",
    "CONTENTS", "TABLE OF CONTENTS", "TABLE OF CONTENT",
    "BIBLIOGRAPHY", "REFERENCES", "WORKS CITED",
    "ANNEXURE", "ANNEXURES", "APPENDIX", "APPENDICES",
})

_CHAPTER_RE = re.compile(r"^\s*CHAPTER\s+\d+\b", re.IGNORECASE)


_ALIGN_MAP: dict = {
    WD_ALIGN_PARAGRAPH.LEFT:    "left",
    WD_ALIGN_PARAGRAPH.CENTER:  "center",
    WD_ALIGN_PARAGRAPH.RIGHT:   "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "left",
}


def _has_page_break(paragraph) -> bool:
    """Detect <w:br w:type='page'/> in any run of the paragraph."""
    for run in paragraph.runs:
        brs = run.element.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
        )
        for br in brs:
            t = br.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
            )
            if t == "page":
                return True
    return False


def _paragraph_to_block(paragraph) -> Block | None:
    """Convert a single python-docx Paragraph to a 5-tuple, or None to skip."""
    text = (paragraph.text or "").strip()
    if not text:
        return None

    sizes: list[float] = []
    bolds: list[bool]  = []
    for run in paragraph.runs:
        sz = run.font.size
        if sz is not None:
            sizes.append(sz.pt)
        bolds.append(bool(run.bold))

    max_size = max(sizes) if sizes else 12.0
    dom_size = max_size
    bold     = any(bolds)

    style_name = ((paragraph.style.name or "") if paragraph.style else "").lower()
    if "title" in style_name:
        max_size = max(max_size, 24.0)
        dom_size = max(dom_size, 24.0)
        bold = True
    elif "heading 1" in style_name:
        # 22pt so the chapter handler's _classify() crosses its
        # body_pt + 8 threshold and returns "h1". 16pt was misclassified
        # as "h2" and let in_toc_section suppress body pages.
        max_size = max(max_size, 22.0)
        dom_size = max(dom_size, 22.0)
        bold = True
    elif "heading 2" in style_name:
        max_size = max(max_size, 16.0)
        dom_size = max(dom_size, 16.0)
        bold = True
    elif "heading 3" in style_name:
        max_size = max(max_size, 14.0)
        dom_size = max(dom_size, 14.0)
        bold = True

    align = _ALIGN_MAP.get(paragraph.alignment, "left")
    return (text, max_size, dom_size, bold, align)


def _is_section_boundary(text: str) -> bool:
    """Does this line start a new virtual page (section header)?"""
    s = text.strip(" ?:.").upper()
    if s in _SECTION_TITLES:
        return True
    if _CHAPTER_RE.match(s):
        return True
    return False


# A "table_rows" is a list of rows; each row is a list of cell text strings.
# Stored separately from blocks so handlers don't need to learn a new shape.
TableRows = list[list[str]]

# (image_blob_bytes, content_type)  -- e.g. (b"\xff\xd8\xff...", "image/jpeg")
ImageBlob = tuple[bytes, str]


# Namespace constants for DOCX OOXML drawing elements
_W_DRAWING = qn("w:drawing")
_W_PICT    = qn("w:pict")
_A_BLIP    = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
_R_EMBED   = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"


def _extract_table_rows(tbl) -> TableRows:
    """Convert a python-docx Table to a list-of-rows-of-cell-strings."""
    rows: TableRows = []
    for row in tbl.rows:
        cells: list[str] = []
        for cell in row.cells:
            cells.append((cell.text or "").strip())
        if cells:
            rows.append(cells)
    return rows


def _extract_paragraph_images(paragraph, doc_part) -> list[ImageBlob]:
    """Return (blob, content_type) tuples for every inline image in a paragraph.

    Walks each run for <w:drawing> elements, finds the <a:blip> inside,
    follows the r:embed relationship id back to the actual image part in
    the DOCX zip, returns raw blob bytes plus MIME type.

    Quiet on missing relationships: a misformed/orphan drawing is skipped
    rather than raised, so one bad image never wipes the whole document.
    """
    out: list[ImageBlob] = []
    for run in paragraph.runs:
        for blip in run.element.findall(".//" + _A_BLIP):
            rId = blip.get(_R_EMBED)
            if not rId:
                continue
            try:
                part = doc_part.related_parts.get(rId)
            except Exception:
                part = None
            if part is None:
                continue
            try:
                blob = part.blob
                ctype = part.content_type
            except Exception:
                continue
            if blob and len(blob) > 200:  # skip 0-byte / corrupt blips
                out.append((blob, ctype))
    return out


def parse_docx_to_pages(
    docx_bytes: bytes,
) -> tuple[
    list[list[Block]],
    list[list[TableRows]],
    list[list[ImageBlob]],
    list[list[tuple]],
    int,
]:
    """Parse a DOCX byte-string into virtual pages of blocks + tables + images.

    Walks doc.element.body in document order so paragraphs, tables and
    images interleave correctly. Returns FOUR per-page lists:

      page_blocks[i]   -- 5-tuple paragraph blocks (for handler dispatch
                          that expects the legacy block contract)
      page_tables[i]   -- TableRows (legacy "all tables at end" emission)
      page_images[i]   -- ImageBlob  (legacy "all images at end" emission)
      page_elements[i] -- INTERLEAVED stream of
                          ("p", Block) | ("tbl", TableRows) | ("img", ImageBlob)
                          in source document order. Use this when
                          rendering needs to preserve "table-in-the-
                          middle-of-the-text" placement -- chapter pages
                          in Step 4.6e.
      n_pages          -- len(page_blocks) (always >= 1)

    The three legacy lists (blocks/tables/images) plus the new
    page_elements list are returned together so callers that only
    consume one of the projections don't pay the cost of computing
    the others themselves.
    """
    doc = Document(io.BytesIO(docx_bytes))
    doc_part = doc.part

    pages:         list[list[Block]]      = [[]]
    page_tables:   list[list[TableRows]]  = [[]]
    page_images:   list[list[ImageBlob]]  = [[]]
    page_elements: list[list[tuple]]      = [[]]

    para_iter = iter(doc.paragraphs)
    tbl_iter  = iter(doc.tables)

    P_TAG = qn("w:p")
    T_TAG = qn("w:tbl")

    def _new_page() -> None:
        pages.append([])
        page_tables.append([])
        page_images.append([])
        page_elements.append([])

    for elem in doc.element.body.iterchildren():
        tag = elem.tag

        if tag == P_TAG:
            try:
                para = next(para_iter)
            except StopIteration:
                continue
            block = _paragraph_to_block(para)
            had_explicit_break = _has_page_break(para)

            imgs = _extract_paragraph_images(para, doc_part)

            if block is None:
                # Empty paragraph that still carries an image (unusual
                # but observed in some templates).
                for img in imgs:
                    page_images[-1].append(img)
                    page_elements[-1].append(("img", img))
                if had_explicit_break and pages[-1]:
                    _new_page()
                continue

            text = block[0]
            # Section-title boundary -> new virtual page BEFORE adding
            # so handlers' first-5-blocks scan finds the title.
            if _is_section_boundary(text) and pages[-1]:
                _new_page()

            pages[-1].append(block)
            page_elements[-1].append(("p", block))
            for img in imgs:
                page_images[-1].append(img)
                page_elements[-1].append(("img", img))

            if had_explicit_break:
                _new_page()

        elif tag == T_TAG:
            try:
                tbl = next(tbl_iter)
            except StopIteration:
                continue
            rows = _extract_table_rows(tbl)
            if rows:
                page_tables[-1].append(rows)
                page_elements[-1].append(("tbl", rows))

        # Other tags (sectPr, sdt, etc.) are intentionally ignored.

    while (len(pages) > 1
           and not pages[-1]
           and not page_tables[-1]
           and not page_images[-1]
           and not page_elements[-1]):
        pages.pop()
        page_tables.pop()
        page_images.pop()
        page_elements.pop()

    if not pages:
        pages = [[]]
        page_tables = [[]]
        page_images = [[]]
        page_elements = [[]]

    return pages, page_tables, page_images, page_elements, len(pages)


def emit_image(doc, blob: bytes, content_type: str = "") -> None:
    """Render an image blob into `doc`.

    Width defaults to 5.5 inches (~14cm) -- safe for A4 with default
    margins. Height auto-scales to preserve aspect ratio.

    Quiet on python-docx errors: if the blob is corrupt or in a format
    python-docx can't handle (rare WMF/EMF), the image is skipped rather
    than killing the whole render.

    Decorative-image filtering (logos repeated on every page) is NOT
    done here -- that's the PDF chapter handler's job for fitz xrefs.
    For DOCX we err on the side of preserving the customer's image even
    if it duplicates a header logo a few times.
    """
    if not blob or len(blob) < 200:
        return
    try:
        doc.add_picture(io.BytesIO(blob), width=Inches(5.5))
        # Center the last paragraph (the one add_picture inserted)
        if doc.paragraphs:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        # Format unsupported / corrupt blob -- skip silently.
        pass


def emit_table(doc, rows: TableRows) -> None:
    """Render a TableRows matrix as a native Word table in `doc`.

    Used by run_from_docx after each per-page handler render. First row
    is treated as a header (bold). Style is "Table Grid" so borders are
    visible. Defensively handles ragged rows by padding to the widest
    row's column count.
    """
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    if n_cols == 0:
        return

    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        tbl.style = "Table Grid"
    except KeyError:
        # Some python-docx setups need the style registered; fall back
        # quietly so we never lose data over styling.
        pass

    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ""
            cell = tbl.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
