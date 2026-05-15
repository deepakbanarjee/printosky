"""FormPageHandler — Canva-designed evaluation forms.

Pages whose layout has no detectable border rules but contain Canva-style
'TRAINEE EVALUATION' / 'PARTICIPANT EVALUATION' forms get rendered as a
full-page PNG instead of flat text (Canva forms have no border lines for
PyMuPDF table detection, and column data interleaves badly when extracted
as flowing text).

Once detected, ctx.in_form_section flips True so the next continuation
page (no keyword) is also rendered as image. The flag clears when a real
chapter divider (large font heading) appears.
"""
from __future__ import annotations

import io
from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .. import extraction
from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


_FORM_KEYWORDS = (
    "TRAINEE EVALUATION", "PARTICIPANT EVALUATION",
    "EVALUATION FORM", "ASSESSMENT FORM",
)


def _is_keyword_form_page(blocks) -> tuple[bool, str]:
    for text, _max, _dom, _bold, _align in blocks:
        u = text.strip(" ?:.").upper()
        for kw in _FORM_KEYWORDS:
            if kw in u and len(text) < 80:
                return True, text.strip(" ?:.")
    return False, ""


def _has_real_h1(blocks, body_pt: float) -> bool:
    for text, max_sz, _dom, _bold, _align in blocks:
        if (max_sz > body_pt + 8
                and text[:1].isupper()
                and len(text) >= 4):
            return True
    return False


class FormPageHandler(SectionHandler):
    """Detects keyword forms; rolls over to next page until real h1."""

    priority = 20
    name = "form_page"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        if page_no < ctx.front_matter_page_limit:
            return False
        is_keyword, _ = _is_keyword_form_page(blocks)
        if is_keyword:
            return True
        if ctx.in_form_section and not _has_real_h1(blocks, ctx.body_pt):
            return True
        return False

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        is_keyword, heading = _is_keyword_form_page(blocks)
        if is_keyword:
            ctx.in_form_section = True
            if heading:
                h = doc.add_heading(heading, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            png = extraction.render_page_png(ctx.pdf, page_no)
            doc.add_picture(io.BytesIO(png), width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except Exception:
            pass
