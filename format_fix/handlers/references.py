"""ReferencesHandler — Bibliography / References / Works Cited section.

Detects pages whose first heading-sized block reads BIBLIOGRAPHY, REFERENCES,
or WORKS CITED. Renders the title as Heading 1, then formats body as a
numbered list (preserves source numbering if present, otherwise emits each
entry as its own justified paragraph). Sub-headings like 'JOURNAL ARTICLES'
and 'WEBSITES' become Heading 2.
"""
from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH

from .. import extraction
from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


_TITLE_VARIANTS = ("BIBLIOGRAPHY", "REFERENCES", "WORKS CITED")


class ReferencesHandler(SectionHandler):
    """Pages whose first heading-sized block reads BIBLIOGRAPHY/REFERENCES."""

    priority = 80
    name = "references"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        for text, max_sz, _dom, _bold, _align in blocks[:8]:
            s = (text or "").strip(" ?:.").upper()
            if s in _TITLE_VARIANTS and max_sz >= ctx.body_pt:
                return True
        return False

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        title_emitted = False
        doc.add_page_break()
        ctx.in_toc_section = False

        blocks = extraction.merge_body_blocks(blocks, ctx.body_pt)
        for text, max_sz, _dom, _bold, _align in blocks:
            s = (text or "").strip()
            if not s or extraction.PAGE_NUM_RE.match(s):
                continue
            if extraction.is_stray_line(s):
                continue
            if extraction.BODY_PAGENUM_RE.match(s) and max_sz <= ctx.body_pt + 1:
                continue

            up = s.strip(" ?:.").upper()

            if not title_emitted and up in _TITLE_VARIANTS:
                h = doc.add_heading(up, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_emitted = True
                continue

            if (extraction.ALL_CAPS_RE.match(s)
                    and 4 <= len(s) <= 40
                    and up not in _TITLE_VARIANTS):
                h2 = doc.add_heading(s.strip(" ?:."), level=2)
                h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(s)
