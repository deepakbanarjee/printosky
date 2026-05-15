"""AbstractHandler — short summary of the project (academic abstract).

Detects pages whose first heading-sized block reads 'ABSTRACT'. Renders the
title as Heading 1, body as justified Normal paragraphs, and a 'Keywords:'
line at the end (if present) as a separate paragraph.
"""
from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH

from .. import extraction
from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


_TITLE_VARIANTS = ("ABSTRACT",)


class AbstractHandler(SectionHandler):
    """Pages whose first heading-sized block reads ABSTRACT."""

    priority = 60
    name = "abstract"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        if page_no >= ctx.front_matter_page_limit:
            return False
        for text, _max, _dom, _bold, _align in blocks[:5]:
            s = (text or "").strip(" ?:.").upper()
            if s in _TITLE_VARIANTS:
                return True
        return False

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        doc.add_page_break()
        title = doc.add_heading("ABSTRACT", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        blocks = extraction.merge_body_blocks(blocks, ctx.body_pt)
        for text, _max, _dom, _bold, _align in blocks:
            s = (text or "").strip()
            if not s or extraction.PAGE_NUM_RE.match(s):
                continue
            if extraction.is_stray_line(s):
                continue
            up = s.strip(" ?:.").upper()
            if up in _TITLE_VARIANTS:
                continue
            p = doc.add_paragraph()
            if s.lower().startswith("keywords"):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(s)
                run.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(s)
