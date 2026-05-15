"""ChapterHandler — catch-all body handler.

Lowest priority. Claims any page that no other handler claimed. Renders
the page's text blocks with classification (h1 / h2 / body / bullet /
numbered list / KV pair) and embeds non-decorative images.
"""
from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches

import fitz

from .. import extraction
from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


def _classify(text: str, max_size: float, body_size: float,
                page_no: int) -> str:
    """Return 'h1' | 'h2' | 'body' | 'skip'."""
    s = text.strip()
    if not s or extraction.PAGE_NUM_RE.match(s):
        return "skip"
    if extraction.BODY_PAGENUM_RE.match(s) and max_size <= body_size + 1:
        return "skip"
    is_alpha_heading = (
        s[:1].isupper()
        and sum(1 for c in s if c.isalpha() and c.isupper())
            >= sum(1 for c in s if c.isalpha() and c.islower())
    )
    looks_like_form_label = bool(re.match(
        r"^(Topic|Venue|Duration|Date|Place|Subject|Title)\b\s*[:.]",
        s, re.IGNORECASE,
    ))
    if re.match(r"^CHAPTER\s+\d+", s, re.IGNORECASE):
        return "h1"
    if (max_size > body_size + 8 and is_alpha_heading
            and 4 <= len(s) <= 80
            and not looks_like_form_label):
        return "h1"
    if (extraction.ALL_CAPS_RE.match(s) and 4 <= len(s) <= 60
            and max_size >= body_size + 2):
        return "h2"
    if (extraction.NUM_HEADING_RE.match(s)
            and max_size >= body_size + 1
            and len(s) <= 80):
        return "h2"
    return "body"


def _flush_kv(doc, kv_buffer):
    """Render a KV run as a 2-column borderless table or fallback paragraphs."""
    if len(kv_buffer) >= 3:
        tbl = doc.add_table(rows=len(kv_buffer), cols=2)
        tbl.autofit = False
        for r_idx, (k, v) in enumerate(kv_buffer):
            cells = tbl.rows[r_idx].cells
            cells[0].width = Cm(6.5)
            cells[1].width = Cm(9.5)
            cells[0].text = k.rstrip(":.").strip()
            cells[1].text = v.strip()
            for run in cells[0].paragraphs[0].runs:
                run.bold = True
        doc.add_paragraph()
    else:
        for k, v in kv_buffer:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(f"{k}: {v}")


class ChapterHandler(SectionHandler):
    """Catch-all body handler. Lowest priority — claims everything else."""

    priority = 100
    name = "chapter"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        return True   # always — last in priority order

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        # Inter-block merging: paragraphs first, then bullet continuations.
        blocks = extraction.merge_body_blocks(blocks, ctx.body_pt)
        blocks = extraction.merge_bullet_continuations(blocks, ctx.body_pt)

        # Detect a real chapter heading anywhere on this page. If found,
        # clear the ToC-continuation flag; otherwise, if we're still inside
        # a ToC region, suppress the whole page (Word's auto ToC re-creates
        # the listing from Heading styles).
        has_real_h1 = any(
            _classify((t or "").strip(), m, ctx.body_pt, page_no) == "h1"
            for t, m, _d, _b, _a in blocks
        )
        if has_real_h1:
            ctx.in_toc_section = False
        elif ctx.in_toc_section:
            return

        kv_buffer: list[tuple[str, str]] = []
        emitted_chapter_h1 = False

        for text, max_sz, _dom, _bold, _align in blocks:
            ln = (text or "").strip()
            if not ln:
                continue
            if extraction.is_stray_line(ln):
                continue
            # Suppress stray TOC-leader lines that escaped the ToC handler
            # (e.g. "Findings of the Study .................... 32").
            if extraction.TOC_LEADER_RE.search(ln):
                continue

            # Key-value detection (groups consecutive K:V lines). Tightened
            # so prose lines like "Very Satisfied (45.8%): 55 respondents..."
            # don't masquerade as form fields. Real K:V pairs are short
            # labels with short values; survey interpretation prose has long
            # values and parenthesised percentages embedded.
            kv_m = extraction.KV_RE.match(ln)
            looks_like_prose_kv = (
                kv_m is not None
                and (
                    "%" in kv_m.group(1)
                    or "(" in kv_m.group(1)
                    or len(kv_m.group(2)) > 80
                    or len(ln) > 120
                )
            )
            if (kv_m
                    and not extraction.ALL_CAPS_RE.match(ln)
                    and not looks_like_prose_kv):
                kv_buffer.append((kv_m.group(1), kv_m.group(2)))
                continue
            elif kv_buffer:
                _flush_kv(doc, kv_buffer)
                kv_buffer = []

            cls = _classify(ln, max_sz, ctx.body_pt, page_no)
            if cls == "skip":
                continue
            if cls == "h1":
                # New chapter → force a page break before every subsequent
                # h1. The first h1 in this section follows the previous
                # handler's page break (or the natural flow start).
                if emitted_chapter_h1:
                    doc.add_page_break()
                else:
                    doc.add_page_break()
                emitted_chapter_h1 = True
                h = doc.add_heading(ln.strip(" ?:"), level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
            if cls == "h2":
                h = doc.add_heading(ln.strip(" ?:"), level=2)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            if extraction.BULLET_RE.match(ln):
                cleaned = extraction.BULLET_RE.sub("", ln).strip()
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(cleaned)
                continue

            if extraction.NUM_LIST_RE.match(ln):
                cleaned = extraction.NUM_LIST_RE.sub("", ln).strip()
                p = doc.add_paragraph(style="List Number")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(cleaned)
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(ln)

        if kv_buffer:
            _flush_kv(doc, kv_buffer)

        # Embed non-decorative images on this page
        page = ctx.pdf[page_no]
        for info in page.get_images(full=True):
            xref = info[0]
            if xref in ctx.decorative_xrefs:
                continue
            try:
                pix = fitz.Pixmap(ctx.pdf, xref)
                if pix.alpha or pix.n > 4:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                png = pix.tobytes("png")
                pix = None
                doc.add_picture(io.BytesIO(png), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                continue
