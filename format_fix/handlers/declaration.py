"""DeclarationHandler — student's declaration of original work."""
from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH

from .. import extraction
from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


_TITLE_VARIANTS = ("DECLARATION",)
_SIG_PREFIXES   = ("Place:", "Date:", "Reg No", "Reg. No",
                    "EN No", "Enrollment")


def _is_signature_line(s: str) -> bool:
    return any(s.startswith(p) for p in _SIG_PREFIXES)


def _looks_like_author_name(s: str) -> bool:
    words = s.split()
    if not (1 < len(words) <= 4):
        return False
    for w in words:
        if not w or not w[0].isalpha():
            continue
        if not (w[:1].isupper() or w.isupper()):
            return False
    return True


class DeclarationHandler(SectionHandler):
    """Pages whose first heading-sized block reads DECLARATION."""

    priority = 50
    name = "declaration"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        if page_no >= ctx.front_matter_page_limit:
            return False
        for text, _max, _dom, _bold, _align in blocks[:5]:
            s = (text or "").strip(" ?:.").upper()
            if s in _TITLE_VARIANTS:
                return True
        return False

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        doc.add_page_break()
        title = doc.add_heading("DECLARATION", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        blocks = extraction.merge_body_blocks(blocks, ctx.body_pt)
        body_lines:  list[str] = []
        sig_lines:   list[str] = []
        author_line: str | None = None

        for text, max_sz, _dom, _bold, _align in blocks:
            s = (text or "").strip()
            if not s or extraction.PAGE_NUM_RE.match(s):
                continue
            if extraction.is_stray_line(s):
                continue
            up = s.strip(" ?:.").upper()
            if up in _TITLE_VARIANTS:
                continue
            if max_sz > ctx.body_pt + 2 and _looks_like_author_name(s):
                author_line = s
                continue
            if _is_signature_line(s):
                sig_lines.append(s)
                continue
            body_lines.append(s)

        for line in body_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(line)

        if author_line:
            doc.add_paragraph()
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ar = ap.add_run(author_line)
            ar.bold = True

        for line in sig_lines:
            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sp.add_run(line)
