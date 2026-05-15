"""CoverHandler — the very first page of the project.

Claims page 0 only. Renders blocks in source Y-order (no reordering, no
content invention). Title-sized blocks become Heading 1 so the font
discipline pass at the end keeps them at 14pt bold; everything else is
Normal at 12pt. Block alignment from the source bbox is preserved.
"""
from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH

from .. import extraction
from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


def _set_para_align(p, blk_align: str) -> None:
    if blk_align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif blk_align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


class CoverHandler(SectionHandler):
    """Page 0 only. Source-order, layout-driven, no invented lines."""

    priority = 30
    name = "cover"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        return page_no == 0

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        for text, max_sz, _dom, _bold, blk_align in blocks:
            s = (text or "").strip()
            if not s or extraction.PAGE_NUM_RE.match(s):
                continue
            if extraction.is_stray_line(s):
                continue

            # Promote to Heading 1 ONLY for short title lines. Body-length
            # boilerplate ("In partial fulfillment of the requirements ...")
            # must stay Normal even if rendered slightly larger in the source.
            is_title_size = (
                max_sz > ctx.body_pt + 2
                and len(s) <= 80
                and not s.lower().startswith((
                    "in partial", "in the partial",
                    "submitted", "guided",
                    "under the guidance",
                ))
            )
            if is_title_size:
                p = doc.add_heading(s.strip(" ?:."), level=1)
            else:
                p = doc.add_paragraph()
                p.add_run(s)
            _set_para_align(p, blk_align)
