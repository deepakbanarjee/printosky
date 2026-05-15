"""TableOfContentsHandler — auto-generated ToC.

Detects pages whose title block reads CONTENTS or TABLE OF CONTENTS and
inserts a Word auto-update TOC field. The student's source ToC entries
are not transcribed — Word regenerates the ToC from the Heading 1/2 styles
applied by downstream chapter handlers.

A placeholder text appears as the field result; the student must
right-click and choose 'Update Field' once after opening the doc.
"""
from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .. import extraction  # noqa: F401  (kept for API symmetry across handlers)
from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


_TITLE_VARIANTS = ("CONTENTS", "TABLE OF CONTENTS")


def _add_toc_field(doc) -> None:
    """Insert a Word TOC field that auto-populates from Heading 1/2/3."""
    p = doc.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = OxmlElement("w:t")
    placeholder.text = ("[Right-click and choose 'Update Field' to populate "
                         "the Table of Contents.]")
    run._r.append(placeholder)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


class TableOfContentsHandler(SectionHandler):
    """Pages whose first heading-sized block reads (TABLE OF) CONTENTS."""

    priority = 70
    name = "toc"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        if page_no >= ctx.front_matter_page_limit:
            return False
        for text, _max, _dom, _bold, _align in blocks[:6]:
            s = (text or "").strip(" ?:.").upper()
            if s in _TITLE_VARIANTS:
                return True
        return False

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        doc.add_page_break()
        title = doc.add_heading("TABLE OF CONTENTS", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        _add_toc_field(doc)
        doc.add_paragraph()
        # Mark ToC region active. Continuation pages (more leader-dot ToC
        # entries with page numbers) are then suppressed by ChapterHandler
        # until a real CHAPTER heading clears the flag.
        ctx.in_toc_section = True
