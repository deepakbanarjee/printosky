"""AnnexuresHandler — Annexures / Appendices / Appendix back-matter.

Detects pages whose first heading-sized block reads ANNEXURES / ANNEXURE /
APPENDIX / APPENDICES. Emits title as Heading 1; body content rendered
verbatim as justified paragraphs. Per the user's directive (this session),
the annexures section is treated as ONE BLOCK — no attempt to parse internal
sub-structure (questionnaires, raw survey data, etc.).
"""
from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH

from .. import extraction
from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


_TITLE_VARIANTS = ("ANNEXURES", "ANNEXURE", "APPENDIX", "APPENDICES")


class AnnexuresHandler(SectionHandler):
    """Pages whose first heading-sized block reads ANNEXURES/APPENDIX."""

    priority = 85
    name = "annexures"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        for text, max_sz, _dom, _bold, _align in blocks[:8]:
            s = (text or "").strip(" ?:.").upper()
            if s in _TITLE_VARIANTS and max_sz >= ctx.body_pt:
                return True
        return False

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        title_emitted = False
        doc.add_page_break()
        ctx.in_toc_section = False

        blocks = extraction.merge_body_blocks(blocks, ctx.body_pt)
        for text, max_sz, _dom, _bold, _align in blocks:
            s = (text or "").strip()
            if not s or extraction.PAGE_NUM_RE.match(s):
                continue
            if extraction.is_stray_line(s):
                continue
            if extraction.BODY_PAGENUM_RE.match(s) and max_sz <= ctx.body_pt + 1:
                continue

            up = s.strip(" ?:.").upper()

            if not title_emitted and up in _TITLE_VARIANTS:
                h = doc.add_heading(up, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_emitted = True
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(s)
