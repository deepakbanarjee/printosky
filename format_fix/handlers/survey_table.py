"""SurveyTableHandler — Yes/No survey question pages (Charts and Figures).

Detects pages where ctx.survey_pages has parsed survey-table data (already
done in Context.build()). Renders each table as:
  - bold caption "Table N: <statement>"
  - 4-column Word table (SL.NO / Variables / No. of Participants / Percentage %)
  - matplotlib bar chart (legend outside plot, value labels above bars)
  - italic figure caption "Figure N: response distribution"
"""
from __future__ import annotations

import io
from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .base import SectionHandler

if TYPE_CHECKING:
    from docx import Document
    from ..context import Context


_LIKERT_COLORS = ["#16a34a", "#65a30d", "#737373", "#f59e0b", "#dc2626"]


def _setup_axes(ax, values):
    ax.set_ylim(0, max(values) * 1.45 if max(values) else 1)
    ax.set_ylabel("No. of Participants", fontsize=12)
    ax.set_xlabel("Response", fontsize=12)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", alpha=0.25, linestyle=":")


def _render_yesno_chart(table: dict) -> bytes | None:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import Patch

        plt.rcParams["font.family"] = "Times New Roman"
        fig, ax = plt.subplots(figsize=(5.5, 3.4), dpi=160)
        labels = ["TRUE", "FALSE"]
        values = [table["true_n"], table["false_n"]]
        colors = ["#2563eb", "#dc2626"]
        bars   = ax.bar(labels, values, color=colors, width=0.55)
        for bar, val, pct in zip(bars, values,
                                  [table["true_pct"], table["false_pct"]]):
            ax.text(bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + 0.5,
                    f"{val} ({pct}%)",
                    ha="center", va="bottom",
                    fontsize=11, fontweight="bold")
        _setup_axes(ax, values)
        ax.legend(
            handles=[Patch(facecolor=colors[0], label="TRUE"),
                     Patch(facecolor=colors[1], label="FALSE")],
            loc="center left", bbox_to_anchor=(1.02, 0.5),
            frameon=True, fontsize=11,
            title="Response", title_fontsize=11,
        )
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


def _render_likert_chart(table: dict) -> bytes | None:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import Patch

        plt.rcParams["font.family"] = "Times New Roman"
        rows = table["rows"]
        labels = [r["label"] for r in rows]
        values = [r["count"] for r in rows]
        pcts   = [r["pct"]   for r in rows]
        colors = _LIKERT_COLORS[: len(rows)]
        fig, ax = plt.subplots(figsize=(6.4, 3.6), dpi=160)
        short = {"Strongly agree": "S.Agree", "Agree": "Agree",
                  "Neutral": "Neutral", "Disagree": "Disagree",
                  "Strongly disagree": "S.Disagree"}
        x_labels = [short.get(l, l) for l in labels]
        bars = ax.bar(x_labels, values, color=colors, width=0.6)
        for bar, val, pct in zip(bars, values, pcts):
            ax.text(bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + 0.5,
                    f"{val} ({pct}%)",
                    ha="center", va="bottom",
                    fontsize=10, fontweight="bold")
        _setup_axes(ax, values)
        ax.legend(
            handles=[Patch(facecolor=c, label=l)
                     for c, l in zip(colors, labels)],
            loc="center left", bbox_to_anchor=(1.02, 0.5),
            frameon=True, fontsize=10,
            title="Response", title_fontsize=11,
        )
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


def _add_yesno_block(doc, t: dict) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cap.add_run(f"Table {t['n']}: {t['statement']}")
    r.bold = True

    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Table Grid"
    headers = ["SL. NO", "Variables", "No. of Participants", "Percentage %"]
    for i, head in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = head
        for r_ in cell.paragraphs[0].runs:
            r_.bold = True
    rows = [
        ("1", "TRUE",  str(t["true_n"]),  str(t["true_pct"])),
        ("2", "FALSE", str(t["false_n"]), str(t["false_pct"])),
        ("",  "TOTAL", str(t["total_n"]), str(t["total_pct"])),
    ]
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            tbl.rows[r_idx].cells[c_idx].text = val
    doc.add_paragraph()

    png_bytes = _render_yesno_chart(t)
    if png_bytes:
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap2 = doc.add_paragraph()
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cap2.add_run(f"Figure {t['n']}: response distribution")
        rr.italic = True
    doc.add_paragraph()


def _add_likert_block(doc, t: dict) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cap.add_run(f"Table {t['n']}: {t['statement']}")
    r.bold = True

    rows = t["rows"]
    n_rows = len(rows) + 2  # header + data rows + total
    tbl = doc.add_table(rows=n_rows, cols=4)
    tbl.style = "Table Grid"
    headers = ["SL. NO", "Option", "No. of Respondents", "Percentage %"]
    for i, head in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = head
        for r_ in cell.paragraphs[0].runs:
            r_.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = tbl.rows[r_idx].cells
        cells[0].text = str(r_idx)
        cells[1].text = row["label"]
        cells[2].text = str(row["count"])
        pct = row["pct"]
        cells[3].text = (f"{int(pct)}%" if pct == int(pct) else f"{pct}%")
    total_cells = tbl.rows[-1].cells
    total_cells[0].text = ""
    total_cells[1].text = "TOTAL"
    total_cells[2].text = str(t["total_n"])
    tp = t["total_pct"]
    total_cells[3].text = (f"{int(tp)}%" if tp == int(tp) else f"{tp}%")
    for c in total_cells:
        for run in c.paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    png_bytes = _render_likert_chart(t)
    if png_bytes:
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap2 = doc.add_paragraph()
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cap2.add_run(f"Chart {t['n']}: response distribution")
        rr.italic = True
    doc.add_paragraph()


def _add_one_block(doc, t: dict) -> None:
    kind = t.get("kind", "yesno")
    if kind == "likert":
        _add_likert_block(doc, t)
    else:
        _add_yesno_block(doc, t)


class SurveyTableHandler(SectionHandler):
    """Pages where Context detected Yes/No survey table data."""

    priority = 10
    name = "survey_table"

    def applies_to(self, blocks, page_no, ctx) -> bool:
        return page_no in ctx.survey_pages

    def render(self, doc: "Document", blocks, page_no, ctx: "Context") -> None:
        # On the *first* survey page in this run, insert a section break +
        # a Heading 1 ("CHARTS AND TABLES") so the analysis chapter stands
        # apart from the preceding narrative chapter. Subsequent survey
        # pages flow under the same heading.
        if not getattr(ctx, "_emitted_survey_header", False):
            doc.add_page_break()
            h = doc.add_heading("CHARTS AND TABLES", level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                ctx._emitted_survey_header = True  # type: ignore[attr-defined]
            except Exception:
                pass
        for t in ctx.survey_pages.get(page_no, []):
            _add_one_block(doc, t)
        ctx.in_form_section = False
        ctx.in_toc_section  = False
