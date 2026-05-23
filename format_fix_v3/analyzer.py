"""Document analyzer + engine router.

Pure-local, zero-cost pre-pass that inspects an uploaded file and decides
which formatting engine should handle it. No API calls, no DB, no network.

Routing logic (2-way):
  - DOCX that already has Word heading styles (Heading 1/2/3 actually used)
    -> "v2_structured": read the student's own structure directly. Accurate
       (12/12 chapters), ~zero API cost. The student already did the work
       of marking chapters; don't throw it away and re-guess with Vision.
  - Everything else (PDF, or DOCX with no heading styles)
    -> "v4_vision": Claude Vision + chunking + async. Must infer structure
       visually because there's none to read.

The analyzer also surfaces diagnostic facts (page count, text-layer
presence, heading counts) so callers/operators can see WHY a route was
chosen, and so future work can add finer routing (e.g. clean text-layer
PDFs -> v2's PDF font-heuristic path to save API cost).
"""
from __future__ import annotations

import io
from typing import Any

# Minimum non-empty heading-styled paragraphs for a DOCX to count as
# "structured". 3 is deliberately low — even a short report with
# intro/body/conclusion marked as headings is worth reading directly.
_MIN_HEADINGS_FOR_STRUCTURED = 3

# A PDF page averaging more than this many characters has a real text
# layer (vs a scanned image). Informational only — PDFs still route to
# Vision today, but this flags scanned docs where the text hint is empty.
_MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER = 100


def _detect_type(src_bytes: bytes) -> str:
    """Return 'pdf' | 'docx' | 'unknown' from magic bytes."""
    if len(src_bytes) < 4:
        return "unknown"
    if src_bytes[:4] == b"%PDF":
        return "pdf"
    if src_bytes[:4] == b"PK\x03\x04":
        return "docx"
    return "unknown"


def _docx_heading_stats(src_bytes: bytes) -> dict[str, Any]:
    """Count non-empty heading-styled paragraphs in a DOCX."""
    from docx import Document
    doc = Document(io.BytesIO(src_bytes))

    heading_counts: dict[str, int] = {}
    for p in doc.paragraphs:
        name = (p.style.name or "") if p.style else ""
        if not p.text.strip():
            continue
        if name.startswith("Heading ") or name == "Title":
            heading_counts[name] = heading_counts.get(name, 0) + 1

    total_headings = sum(heading_counts.values())
    h1_count = heading_counts.get("Heading 1", 0)

    return {
        "heading_counts":  heading_counts,
        "total_headings":  total_headings,
        "h1_count":        h1_count,
        "paragraph_count": len(doc.paragraphs),
        "table_count":     len(doc.tables),
    }


def _pdf_text_stats(src_bytes: bytes) -> dict[str, Any]:
    """Compute page count + average text-layer density for a PDF."""
    import fitz
    pdf = fitz.open(stream=src_bytes, filetype="pdf")
    n = pdf.page_count
    total_chars = 0
    for i in range(n):
        total_chars += len(pdf[i].get_text("text").strip())
    pdf.close()
    avg = total_chars / max(1, n)
    return {
        "page_count":          n,
        "avg_chars_per_page":  round(avg, 1),
        "has_text_layer":      avg >= _MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER,
    }


def analyze(src_bytes: bytes) -> dict[str, Any]:
    """Inspect a file and recommend an engine. No API/DB/network.

    Returns:
        {
            "input_type":   "pdf" | "docx" | "unknown",
            "page_count":   int | None,
            "route":        "v2_structured" | "v4_vision" | "reject",
            "reason":       str,
            ...diagnostic fields depending on type...
        }
    """
    input_type = _detect_type(src_bytes)

    if input_type == "unknown":
        return {
            "input_type": "unknown",
            "page_count": None,
            "route":      "reject",
            "reason":     "not a PDF or DOCX (bad magic bytes)",
        }

    if input_type == "docx":
        try:
            stats = _docx_heading_stats(src_bytes)
        except Exception as exc:
            return {
                "input_type": "docx",
                "page_count": None,
                "route":      "v4_vision",
                "reason":     f"DOCX parse failed ({type(exc).__name__}); "
                              f"falling back to Vision",
            }

        structured = stats["total_headings"] >= _MIN_HEADINGS_FOR_STRUCTURED
        if structured:
            route = "v2_structured"
            reason = (
                f"DOCX has {stats['total_headings']} heading-styled "
                f"paragraphs ({stats['h1_count']} Heading-1) -> read the "
                f"existing structure directly (no Vision needed)"
            )
        else:
            route = "v4_vision"
            reason = (
                f"DOCX has only {stats['total_headings']} heading-styled "
                f"paragraphs (student used manual formatting) -> Vision"
            )
        return {
            "input_type":     "docx",
            "page_count":     None,
            "route":          route,
            "reason":         reason,
            "heading_counts": stats["heading_counts"],
            "total_headings": stats["total_headings"],
            "h1_count":       stats["h1_count"],
            "table_count":    stats["table_count"],
        }

    # input_type == "pdf"
    try:
        stats = _pdf_text_stats(src_bytes)
    except Exception as exc:
        return {
            "input_type": "pdf",
            "page_count": None,
            "route":      "v4_vision",
            "reason":     f"PDF parse failed ({type(exc).__name__}); "
                          f"defaulting to Vision",
        }

    route = "v4_vision"
    if stats["has_text_layer"]:
        reason = (
            f"PDF, {stats['page_count']} pages, has a text layer "
            f"({stats['avg_chars_per_page']} chars/page avg) -> Vision "
            f"(no heading styles in PDF to read)"
        )
    else:
        reason = (
            f"PDF, {stats['page_count']} pages, scanned/image-only "
            f"({stats['avg_chars_per_page']} chars/page) -> Vision "
            f"(text hint will be sparse)"
        )
    return {
        "input_type":         "pdf",
        "page_count":         stats["page_count"],
        "route":              route,
        "reason":             reason,
        "avg_chars_per_page": stats["avg_chars_per_page"],
        "has_text_layer":     stats["has_text_layer"],
    }
