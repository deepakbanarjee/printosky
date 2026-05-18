"""Render Claude Vision JSON to .docx.

Consumes the per-page JSON dicts returned by vision.analyze_page and
produces an A4 Word document with:
  - Times New Roman 12 body / Heading styles, all black
  - 1.5 line spacing on body paragraphs
  - Page break BEFORE every chapter (level 1 heading)
  - Per-run bold/italic spans (the runs[] schema)
  - Content-fit table columns within the printable area
  - Numbered/bulleted lists with auto-numbering (no double prefix)
  - Title-page logo capped at 2" so the page fits one A4 sheet
  - Tightened spacing for name lists (title page, acknowledgement)
  - Merged "Name, Designation" line for guide blocks
  - Optional page numbers (Roman for front matter, Arabic from Ch1)
  - Optional header / footer text

Page-image pairing: each page JSON may have N "image" elements; the
caller supplies the actual image bytes extracted from that PDF page
(in source order). We embed them at the position the image element
appears in the JSON stream.
"""
from __future__ import annotations

import io
import re
from datetime import date
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

PAGE_BREAK_BEFORE_KINDS = {
    "title", "certificate", "acknowledgement", "declaration",
    "abstract", "toc", "list_of_figures", "list_of_tables",
    "references", "appendix",
}

# Single-unit pages where NO page-break should fire between H1s (because
# a multi-line title gets emitted as several consecutive H1 elements).
NO_INTRA_H1_BREAK_KINDS = {
    "title", "certificate", "acknowledgement", "declaration",
    "abstract",
}

SKIP_TYPES = {"page_number", "footer", "skip"}

# H1 must look like a real chapter to keep its level. Otherwise demote.
_REAL_CHAPTER_RE = re.compile(
    r"^\s*("
    r"CHAPTER\s+[\divxIVX]+|"
    r"\d+\.\s+[A-Z]|"
    r"PART\s+[\divxIVX]+|"
    r"APPENDIX\s+[A-Z\d]|"
    r"REFERENCES\s*$|"
    r"BIBLIOGRAPHY\s*$|"
    r"INTRODUCTION\s*$|"
    r"CONCLUSION\s*$|"
    r"ACKNOWLEDGEMENT|"
    r"ABSTRACT\s*$|"
    r"CONTENTS\s*$"
    r")",
    re.IGNORECASE,
)

# Printable width inside a 1-inch-margin A4 page: 21cm - 2.54cm = ~6.5"
PRINTABLE_WIDTH_INCHES = 6.5
TITLE_PAGE_LOGO_MAX_INCHES = 2.0
DEFAULT_IMAGE_MAX_INCHES = 5.5
TABLE_CELL_PADDING_TWIPS = 40   # 2px in twentieths-of-a-point (approx)


# ---------------------------------------------------------------------------
# Filename helper

_FILENAME_SAFE_RE = re.compile(r"[^A-Za-z0-9._-]+")
_JUNK_NAME_RE = re.compile(r"^[a-f0-9-]{8,}$", re.IGNORECASE)


def smart_download_filename(
    original_filename: str | None,
    project_title: str | None = None,
) -> str:
    """Return a sensible download name for the formatted .docx.

    Priority:
      1. Original filename (with extension swapped to .docx, suffix
         "_formatted") IF it looks like a real name (>=3 alphanumeric
         chars, not a UUID).
      2. Project title from page-1 detection (sanitized).
      3. Fallback: "Printosky_Project_<yyyy-mm-dd>.docx"
    """
    if original_filename:
        stem = original_filename.rsplit(".", 1)[0].strip()
        if (len(stem) >= 3
                and any(c.isalpha() for c in stem)
                and not _JUNK_NAME_RE.match(stem)):
            safe = _FILENAME_SAFE_RE.sub("_", stem).strip("_")
            if safe:
                return f"{safe}_formatted.docx"

    if project_title:
        safe = _FILENAME_SAFE_RE.sub("_", project_title.strip()).strip("_")
        if safe and len(safe) >= 3:
            return f"{safe}_formatted.docx"

    return f"Printosky_Project_{date.today().isoformat()}.docx"


# ---------------------------------------------------------------------------
# Document setup

def _set_line_spacing(paragraph_format, spacing: float = 1.5) -> None:
    paragraph_format.line_spacing = spacing


def _force_black(run) -> None:
    """Explicitly black so any inherited theme colour gets overridden."""
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _setup_doc(line_spacing: float = 1.5) -> "Document":
    """A4, 1-inch margins, Times New Roman 12 black body, 1.5 line spacing."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    _set_line_spacing(normal.paragraph_format, line_spacing)

    for level, size in [(1, 18), (2, 14), (3, 13)]:
        try:
            h = doc.styles[f"Heading {level}"]
            h.font.name = "Times New Roman"
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.keep_with_next = True
            h.paragraph_format.widow_control = True
            _set_line_spacing(h.paragraph_format, line_spacing)
        except KeyError:
            pass

    return doc


# ---------------------------------------------------------------------------
# Headers / footers / page numbers

def _add_page_field(paragraph) -> None:
    """Insert a Word PAGE field into a paragraph (current page number)."""
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char1)
    run._element.append(instr_text)
    run._element.append(fld_char2)


def _set_page_number_format(section, fmt: str) -> None:
    """Set page-number format on a section: 'decimal' or 'lowerRoman'."""
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    if fmt == "decimal":
        pg_num_type.set(qn("w:start"), "1")


def _apply_page_numbers(section, position: str = "center") -> None:
    """Put a PAGE field in the section footer."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    align = ALIGN_MAP.get(position, WD_ALIGN_PARAGRAPH.CENTER)
    p.alignment = align
    for child in list(p._element):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            p._element.remove(child)
    _add_page_field(p)


def _apply_header(section, text: str) -> None:
    if not text:
        return
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for child in list(p._element):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            p._element.remove(child)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    _force_black(run)


def _apply_footer_text(section, text: str) -> None:
    """Add an EXTRA text line ABOVE the page-number line."""
    if not text:
        return
    footer = section.footer
    p = footer.add_paragraph()
    footer._element.insert(0, p._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    _force_black(run)


# ---------------------------------------------------------------------------
# Element rendering

def _add_runs(paragraph, runs: list[dict[str, Any]]) -> None:
    """Add multiple runs (text + bold + italic). Forces TNR + black."""
    for r in runs or []:
        text = r.get("text", "") or ""
        if not text:
            continue
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        if r.get("bold"):
            run.bold = True
        if r.get("italic"):
            run.italic = True
        _force_black(run)


def _flat_text(runs: list[dict[str, Any]]) -> str:
    return "".join((r.get("text") or "") for r in (runs or []))


def _emit_heading(doc, el: dict[str, Any], force_page_break: bool) -> None:
    level = el.get("level") or 1
    if level not in (1, 2, 3):
        level = 2
    if force_page_break and level == 1:
        doc.add_page_break()
    h = doc.add_heading(level=level)
    h.alignment = ALIGN_MAP.get(el.get("alignment", "left"),
                                 WD_ALIGN_PARAGRAPH.LEFT)
    _add_runs(h, el.get("runs", []))


def _emit_body(doc, el: dict[str, Any], *, tight: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = ALIGN_MAP.get(el.get("alignment", "justify"),
                                 WD_ALIGN_PARAGRAPH.JUSTIFY)
    if tight:
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    _add_runs(p, el.get("runs", []))


_NUM_LIST_PREFIX_RE = re.compile(r"^\s*(\(?\d+[.)]?|[ivxIVX]+\.)\s+")
_BULLET_LIST_PREFIX_RE = re.compile(r"^\s*[•●○\-*]\s+")


def _strip_list_prefix(text: str, style: str) -> str:
    if style == "List Number":
        return _NUM_LIST_PREFIX_RE.sub("", text, count=1)
    return _BULLET_LIST_PREFIX_RE.sub("", text, count=1)


def _emit_list_item(doc, el: dict[str, Any]) -> None:
    style = ("List Number" if el.get("list_style") == "number"
             else "List Bullet")
    p = doc.add_paragraph(style=style)
    p.alignment = ALIGN_MAP.get(el.get("alignment", "left"),
                                 WD_ALIGN_PARAGRAPH.LEFT)
    runs = list(el.get("runs", []))
    if runs:
        first = dict(runs[0])
        first["text"] = _strip_list_prefix(first.get("text", "") or "", style)
        runs[0] = first
    _add_runs(p, runs)


def _emit_caption(doc, el: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runs = list(el.get("runs", [])) or [{"text": ""}]
    runs = [{**r, "bold": True} for r in runs]
    _add_runs(p, runs)


# ---------------------------------------------------------------------------
# Tables (content-fit columns, total <= printable width)

def _set_cell_margins(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(twips))
        el.set(qn("w:type"), "dxa")


def _disable_table_autofit(table) -> None:
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _compute_column_widths(rows: list[list[Any]],
                            total_width_inches: float) -> list[float]:
    """Content-aware column widths."""
    if not rows:
        return []
    n_cols = max(len(r) for r in rows)
    col_chars = [1] * n_cols
    for row in rows:
        for c_idx in range(n_cols):
            if c_idx < len(row):
                cell = row[c_idx]
                if isinstance(cell, dict):
                    text = _flat_text(cell.get("runs", []))
                else:
                    text = str(cell or "")
                longest = max((len(line) for line in text.splitlines()),
                              default=len(text))
                col_chars[c_idx] = max(col_chars[c_idx], longest)
    MIN_COL_FRAC = 0.05
    MAX_COL_FRAC = 0.55
    raw_total = sum(col_chars)
    weights = [c / raw_total for c in col_chars]
    weights = [min(MAX_COL_FRAC, max(MIN_COL_FRAC, w)) for w in weights]
    s = sum(weights)
    weights = [w / s for w in weights]
    return [total_width_inches * w for w in weights]


def _emit_table(doc, el: dict[str, Any]) -> None:
    rows = el.get("table_rows") or []
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _disable_table_autofit(tbl)

    widths_in = _compute_column_widths(rows, PRINTABLE_WIDTH_INCHES)

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.width = Inches(widths_in[c_idx])
            _set_cell_margins(cell, TABLE_CELL_PADDING_TWIPS)

            raw = row[c_idx] if c_idx < len(row) else None
            if isinstance(raw, dict):
                cell_runs = raw.get("runs", [{"text": ""}])
                cell_align = raw.get("alignment", "left")
            else:
                cell_runs = [{"text": str(raw) if raw is not None else ""}]
                cell_align = "left"

            p = cell.paragraphs[0]
            p.alignment = ALIGN_MAP.get(cell_align, WD_ALIGN_PARAGRAPH.LEFT)
            if r_idx == 0:
                cell_runs = [{**r, "bold": True} for r in cell_runs]
            _add_runs(p, cell_runs)

    doc.add_paragraph()


def _emit_image(doc, image_bytes: bytes,
                position: str | None,
                page_kind: str | None,
                title_logo_seen: dict[str, bool]) -> None:
    if not image_bytes:
        return
    try:
        if page_kind == "title" and not title_logo_seen.get("seen"):
            width = Inches(TITLE_PAGE_LOGO_MAX_INCHES)
            title_logo_seen["seen"] = True
        elif position == "full_width":
            width = Inches(PRINTABLE_WIDTH_INCHES)
        else:
            width = Inches(min(DEFAULT_IMAGE_MAX_INCHES,
                                PRINTABLE_WIDTH_INCHES))
        doc.add_picture(io.BytesIO(image_bytes), width=width)
        if position in ("centered", "full_width", None):
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Heuristics: name-list spacing, lecturer-merge

_NAME_LINE_RE = re.compile(
    r"^[A-Z][A-Z .]{1,40}"
    r"(\s+(REG\.?\s*NO|ROLL\.?\s*NO|ID)\s*[:.]\s*[\w/]+)?$"
)


def _is_name_list_element(el: dict[str, Any]) -> bool:
    if el.get("type") != "body":
        return False
    if el.get("alignment") not in ("center", "left"):
        return False
    text = _flat_text(el.get("runs", [])).strip()
    if not text or len(text) > 80:
        return False
    return bool(_NAME_LINE_RE.match(text)) or (
        text.isupper() and 4 <= len(text) <= 50 and " " in text
    )


_DESIGNATION_HINTS = (
    "lecturer", "professor", "assistant", "associate", "head",
    "hod", "principal", "guide", "supervisor", "dean", "coordinator",
)


def _looks_like_designation(el: dict[str, Any]) -> bool:
    if el.get("type") != "body":
        return False
    text = _flat_text(el.get("runs", [])).strip().lower()
    if not text or len(text) > 80:
        return False
    return any(h in text for h in _DESIGNATION_HINTS)


def _try_merge_lecturer(prev_heading_el: dict[str, Any] | None,
                         designation_el: dict[str, Any]) -> dict[str, Any] | None:
    """If the previous element was a guide-name heading and this is a
    designation, return a merged single-line body element."""
    if not prev_heading_el or prev_heading_el.get("type") != "heading":
        return None
    if (prev_heading_el.get("level") or 0) < 3:
        return None
    name_runs = list(prev_heading_el.get("runs", []))
    desig_runs = list(designation_el.get("runs", []))
    if not name_runs or not desig_runs:
        return None
    merged_runs = []
    merged_runs.extend({**r, "bold": True} for r in name_runs)
    merged_runs.append({"text": ", "})
    merged_runs.extend({**r, "italic": True} for r in desig_runs)
    return {
        "type": "body",
        "alignment": designation_el.get("alignment", "center"),
        "runs": merged_runs,
    }


# ---------------------------------------------------------------------------
# Main render

def render_pages(
    pages: list[dict[str, Any]],
    page_images: list[list[bytes]],
    *,
    line_spacing: float = 1.5,
    page_numbers: bool = True,
    page_number_style: str = "roman_then_arabic",
    page_number_position: str = "center",
    header_text: str = "",
    footer_text: str = "",
) -> bytes:
    """Render the full document.

    Args:
        pages: List of per-page JSON dicts from vision.analyze_page.
        page_images: Parallel list - page_images[i] is the list of image
            bytes extracted from the source for PDF/DOCX page i+1.
        line_spacing: Multiplier (default 1.5).
        page_numbers: If True, add a PAGE field to the footer.
        page_number_style: "roman_then_arabic" inserts a section break
            before the first chapter so front matter uses i/ii/iii and
            chapters use 1/2/3. "decimal" uses arabic throughout. "off"
            disables (overrides page_numbers=True).
        page_number_position: "left" / "center" / "right".
        header_text: Optional text in the page header.
        footer_text: Optional text in the footer ABOVE the page number.

    Returns:
        DOCX file bytes.
    """
    doc = _setup_doc(line_spacing=line_spacing)

    front_section = doc.sections[0]
    use_roman = (
        page_numbers
        and page_number_style == "roman_then_arabic"
    )
    if page_numbers and page_number_style != "off":
        if use_roman:
            _set_page_number_format(front_section, "lowerRoman")
        else:
            _set_page_number_format(front_section, "decimal")
        _apply_page_numbers(front_section, position=page_number_position)

    _apply_header(front_section, header_text)
    _apply_footer_text(front_section, footer_text)

    prev_kind: str | None = None
    title_logo_seen = {"seen": False}
    first_chapter_break_inserted = False

    for p_idx, page in enumerate(pages):
        kind = page.get("page_kind", "other")

        kind_change_break = (
            kind in PAGE_BREAK_BEFORE_KINDS
            and prev_kind is not None
            and prev_kind != kind
        )
        if kind_change_break:
            doc.add_page_break()

        if (use_roman
                and kind == "chapter"
                and not first_chapter_break_inserted):
            new_section = doc.add_section()
            new_section.page_width = front_section.page_width
            new_section.page_height = front_section.page_height
            new_section.left_margin = front_section.left_margin
            new_section.right_margin = front_section.right_margin
            new_section.top_margin = front_section.top_margin
            new_section.bottom_margin = front_section.bottom_margin
            _set_page_number_format(new_section, "decimal")
            if page_numbers:
                _apply_page_numbers(new_section,
                                     position=page_number_position)
            _apply_header(new_section, header_text)
            _apply_footer_text(new_section, footer_text)
            first_chapter_break_inserted = True

        imgs = list(page_images[p_idx]) if p_idx < len(page_images) else []
        img_cursor = 0

        elements = list(page.get("elements", []))

        # Pre-pass: lecturer-merge.
        merged_elements: list[dict[str, Any]] = []
        i = 0
        while i < len(elements):
            el = elements[i]
            nxt = elements[i + 1] if i + 1 < len(elements) else None
            if (nxt
                    and el.get("type") == "heading"
                    and (el.get("level") or 0) >= 3
                    and _looks_like_designation(nxt)):
                merged = _try_merge_lecturer(el, nxt)
                if merged is not None:
                    merged_elements.append(merged)
                    i += 2
                    continue
            merged_elements.append(el)
            i += 1
        elements = merged_elements

        # Detect name-list runs for tight spacing.
        name_run_flags = [False] * len(elements)
        run_start: int | None = None
        for idx, el in enumerate(elements):
            if _is_name_list_element(el):
                if run_start is None:
                    run_start = idx
            else:
                if run_start is not None and idx - run_start >= 2:
                    for j in range(run_start, idx):
                        name_run_flags[j] = True
                run_start = None
        if run_start is not None and len(elements) - run_start >= 2:
            for j in range(run_start, len(elements)):
                name_run_flags[j] = True

        # Pre-pass: H1 dedup + downgrade.
        # 1. Drop H1 if it duplicates the previous-emitted H1 text (case-
        #    insensitive, ignoring whitespace) within the last 3 elements.
        #    Catches the Claude "REVIEW OF LITERATURE on page-end + page-
        #    start" pattern.
        # 2. Downgrade H1 to H2 if the text doesn't look like a real
        #    chapter heading (CHAPTER N, N. TITLE, REFERENCES, etc.).
        #    Stops random sub-sections (PILOT STUDY, TOPIC, PREFACE) from
        #    triggering chapter-level page breaks.
        deduped: list[dict[str, Any]] = []
        for el in elements:
            if el.get("type") == "heading":
                cur_text = _flat_text(el.get("runs", [])).strip().lower()
                # Look back at the last 3 elements for an identical heading
                # text (regardless of level - catches the case where the
                # first occurrence already got downgraded to H2 but the
                # source repeats it again).
                tail = deduped[-3:] if deduped else []
                is_dup = any(
                    (p.get("type") == "heading"
                     and _flat_text(p.get("runs", [])).strip().lower()
                         == cur_text)
                    for p in tail
                )
                if is_dup and cur_text:
                    continue
                # H1 downgrade: only keep H1 if it looks like a real chapter
                if ((el.get("level") or 0) == 1
                        and not _REAL_CHAPTER_RE.match(cur_text)):
                    el = {**el, "level": 2}
            deduped.append(el)
        elements = deduped

        # Title-page H1 grouping: on pages like "title" / "certificate" /
        # "acknowledgement" we treat the WHOLE page as one unit. Even if
        # Claude returned multiple H1s (because it split a multi-line
        # title), we don't page-break between them.
        no_intra_h1_break = kind in NO_INTRA_H1_BREAK_KINDS

        first_h1_on_page = True

        for idx, el in enumerate(elements):
            etype = el.get("type", "body")
            if etype in SKIP_TYPES:
                continue
            if etype == "heading":
                level = el.get("level") or 1
                # Force a page break for H1 EXCEPT:
                #  - First H1 on a fresh-section page (kind-change already broke)
                #  - Any H1 on a single-unit page (title/cert/ack/etc)
                force_break = (
                    level == 1
                    and not (first_h1_on_page and kind_change_break)
                    and not no_intra_h1_break
                )
                _emit_heading(doc, el, force_page_break=force_break)
                if level == 1:
                    first_h1_on_page = False
            elif etype == "body":
                _emit_body(doc, el, tight=name_run_flags[idx])
            elif etype == "list_item":
                _emit_list_item(doc, el)
            elif etype == "caption":
                _emit_caption(doc, el)
            elif etype == "table":
                _emit_table(doc, el)
            elif etype == "image":
                blob = imgs[img_cursor] if img_cursor < len(imgs) else b""
                img_cursor += 1
                _emit_image(doc, blob, el.get("image_position"),
                             kind, title_logo_seen)
            else:
                _emit_body(doc, el)

        # Leftover-image emission: for DOCX inputs, fitz strips images
        # during DOCX->PDF conversion, so Claude never sees them and
        # never emits "image" elements. The orchestrator extracts the
        # actual image bytes directly from the DOCX and hands them in
        # via page_images. If we've consumed fewer image elements than
        # extracted images for this page, emit the leftovers now (at
        # end of page) so they aren't lost.
        while img_cursor < len(imgs):
            _emit_image(doc, imgs[img_cursor], "centered",
                         kind, title_logo_seen)
            img_cursor += 1

        prev_kind = kind

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Project title extraction (for smart filename)

def extract_project_title(pages: list[dict[str, Any]]) -> str | None:
    """Find the project title from the title page, if present."""
    if not pages:
        return None
    title_page = next((p for p in pages
                        if p.get("page_kind") == "title"), None)
    if not title_page:
        return None
    SKIP_PHRASES = (
        "project report", "submitted", "fulfillment", "fulfilment",
        "degree of", "diploma in", "guided by", "department of",
    )
    candidates = []
    for el in title_page.get("elements", []):
        if el.get("type") != "heading":
            continue
        text = _flat_text(el.get("runs", [])).strip()
        if not text or len(text) < 4 or len(text) > 80:
            continue
        if any(s in text.lower() for s in SKIP_PHRASES):
            continue
        score = (3 - (el.get("level") or 3)) * 100 + len(text)
        candidates.append((score, text))
    if not candidates:
        return None
    return max(candidates, key=lambda x: x[0])[1]
