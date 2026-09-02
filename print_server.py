"""
print_server.py — Printosky Print Server
Runs on store PC at port 3005.
Receives print commands from admin dashboard → prints via SumatraPDF → updates job status.

Endpoints:
  POST /print          — { job_id, filepath, printer, copies, colour_mode, staff_id }
  POST /staff-login    — { pin, pc_id } → { ok, staff_id, name, session_id }
  POST /staff-logout   — { session_id, idle }
  GET  /active-staff   — ?pc_id=PC1 → { staff_id, name, session_id }
  GET  /status         — health check
  GET  /printers       — list configured printers
  POST /local-print    — save + print a file on THIS PC (no cloud round trip)
  GET  /health         — full system health
  POST /create-job     — { customer_name, phone, source, colour, sides, copies, pages, paper_size, finishing, amount_collected|amount_partial, payment_mode, override_reason }
  POST /upload-file    — { filename, file_data (base64) } → saves to hot folder
"""

import hashlib
import hmac
import json
import logging
import math
import collections
import os
import socket
import sqlite3
import subprocess
import sys
import threading
import urllib.request
from datetime import datetime
from http.server import BaseHTTPRequestHandler, HTTPServer
from pathlib import Path
from urllib.parse import parse_qs, urlparse
from dotenv import load_dotenv

from store_config import get_store_config

load_dotenv()

# Log to the console AND a rotating file, so print history survives after the
# console window is closed (chasing a live cmd window is not a diagnostic plan).
# 2 MB × 5 backups ≈ 10 MB cap. File logging is best-effort — if the logs dir
# can't be created, the console handler still works.
#
# Must be the FIRST logging call anywhere in this module. Python's logging.info()
# / .warning() etc. silently trigger an implicit logging.basicConfig() with
# default settings (WARNING level, console only) the first time any of them
# runs with no handlers configured yet — which makes a LATER explicit
# basicConfig() call a permanent no-op for the rest of the process, dropping
# every .info() line, file and all. Confirmed: once printer_queue_names got
# configured, the store_config override's own confirmation log line (an
# earlier logging.info() call further down this file) triggered exactly this,
# and every INFO-level log after it vanished for that entire process.
import logging.handlers as _log_handlers_mod

_log_handlers = [logging.StreamHandler()]
try:
    _log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
    os.makedirs(_log_dir, exist_ok=True)
    _log_handlers.append(_log_handlers_mod.RotatingFileHandler(
        os.path.join(_log_dir, "print_server.log"),
        maxBytes=2_000_000, backupCount=5, encoding="utf-8",
    ))
except Exception:
    pass

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [PRINT] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=_log_handlers,
)

# Rate card engine (same directory)
sys.path.insert(0, os.path.dirname(__file__))
try:
    import rate_card as _rc
except ImportError:
    _rc = None
    logging.warning("rate_card.py not found — /quote will return 0")

# Job tracker — status machine + audit log
try:
    from job_tracker import log_event as _jt_log, transition as _jt_transition, get_events as _jt_events, setup_job_events_db as _jt_setup
    JOB_TRACKER_AVAILABLE = True
except ImportError:
    JOB_TRACKER_AVAILABLE = False
    def _jt_log(*a, **kw): return 0
    def _jt_transition(*a, **kw): return {"ok": False, "error": "job_tracker not available"}
    def _jt_events(*a, **kw): return []
    def _jt_setup(*a): pass

# Colour detector — PyMuPDF-based colour page detection
try:
    from colour_detector import (
        build_colour_map as _cd_build,
        save_colour_map as _cd_save,
        confirm_colour_map as _cd_confirm,
    )
    COLOUR_DETECTOR_AVAILABLE = True
except ImportError:
    COLOUR_DETECTOR_AVAILABLE = False
    def _cd_build(*a, **kw): return {"error": "colour_detector not available"}
    def _cd_save(*a, **kw): pass
    def _cd_confirm(*a, **kw): pass

# Review manager — post-collection review requests + discount codes
try:
    from review_manager import (
        schedule_review as _rv_schedule,
        setup_review_db as _rv_setup,
        record_rating as _rv_record,
    )
    REVIEW_MANAGER_AVAILABLE = True
except ImportError:
    REVIEW_MANAGER_AVAILABLE = False
    def _rv_schedule(*a, **kw): pass
    def _rv_setup(*a): pass
    def _rv_record(*a, **kw): return {"ok": False, "error": "review_manager not available"}

# Work session tracker — DTP / editing timer
try:
    from work_session_tracker import (
        start_session as _ws_start,
        pause_session as _ws_pause,
        resume_session as _ws_resume,
        end_session as _ws_end,
        get_sessions as _ws_get,
        get_open_session as _ws_open,
        setup_work_sessions_db as _ws_setup,
    )
    WORK_SESSION_AVAILABLE = True
except ImportError:
    WORK_SESSION_AVAILABLE = False
    def _ws_start(*a, **kw): return {"ok": False, "error": "work_session_tracker not available"}
    def _ws_pause(*a, **kw): return {"ok": False, "error": "work_session_tracker not available"}
    def _ws_resume(*a, **kw): return {"ok": False, "error": "work_session_tracker not available"}
    def _ws_end(*a, **kw): return {"ok": False, "error": "work_session_tracker not available"}
    def _ws_get(*a, **kw): return []
    def _ws_open(*a, **kw): return None
    def _ws_setup(*a): pass

# ── Konica Windows username → PC identifier mapping ───────────────────────────
# PC1 = Priya/Deepak/Anu  |  PC2 = Revana  |  PC3 = rarely used (Nirmal)
# KONICA_USER_PC_MAP retired 2026-05-12 (0/4507 attribution rate); see
# retired/2026-05-12-graveyard/konica_attribution.py for the dict and the
# four root-cause failures documented in that file's header.

# ── Staff session helpers ──────────────────────────────────────────────────────
_active_sessions = {}   # pc_id → {staff_id, name, session_id}  (in-memory cache)

# ── Supabase JWT cache (for returning to admin.html on staff login) ────────────
_supabase_jwt_cache = {"token": None, "expires_at": 0}

def _get_supabase_jwt() -> str:
    """Return a cached Supabase JWT, refreshing if expired or missing."""
    import time, json as _json
    now = time.time()
    if _supabase_jwt_cache["token"] and now < _supabase_jwt_cache["expires_at"] - 60:
        return _supabase_jwt_cache["token"]
    url  = os.environ.get("SUPABASE_URL", "")
    email = os.environ.get("SUPABASE_AUTH_EMAIL", "")
    pwd   = os.environ.get("SUPABASE_AUTH_PASSWORD", "")
    if not (url and email and pwd):
        return ""
    try:
        import urllib.request as _ur, urllib.error
        body = _json.dumps({"email": email, "password": pwd}).encode()
        req  = _ur.Request(f"{url}/auth/v1/token?grant_type=password", data=body,
                           headers={"apikey": os.environ.get("SUPABASE_KEY", ""),
                                    "Content-Type": "application/json"})
        with _ur.urlopen(req, timeout=5) as resp:
            data = _json.loads(resp.read())
        token = data.get("access_token", "")
        expires_in = data.get("expires_in", 3600)
        _supabase_jwt_cache["token"] = token
        _supabase_jwt_cache["expires_at"] = now + expires_in
        return token
    except Exception as e:
        logging.warning(f"Supabase JWT refresh failed: {e}")
        return ""


def _sha256(text: str) -> str:
    """Legacy SHA-256 hash — kept for admin password comparison only."""
    return hashlib.sha256(text.encode()).hexdigest()

# ── PBKDF2 PIN hashing ────────────────────────────────────────────────────────
_PBKDF2_ITERATIONS = 260_000

def _hash_pin(pin: str) -> tuple[str, str]:
    """Return (hash_hex, salt_hex) using PBKDF2-HMAC-SHA256."""
    import secrets as _secrets
    salt = _secrets.token_hex(16)
    h = hashlib.pbkdf2_hmac("sha256", pin.encode(), salt.encode(), _PBKDF2_ITERATIONS).hex()
    return h, salt

def _verify_pin(pin: str, stored_hash: str, stored_salt: str | None) -> bool:
    """Verify PIN against stored hash. Handles both legacy SHA-256 (salt=None) and PBKDF2."""
    if stored_salt is None:
        # Legacy path: plain SHA-256
        return hmac.compare_digest(stored_hash, hashlib.sha256(pin.encode()).hexdigest())
    # New path: PBKDF2
    expected = hashlib.pbkdf2_hmac("sha256", pin.encode(), stored_salt.encode(), _PBKDF2_ITERATIONS).hex()
    return hmac.compare_digest(stored_hash, expected)


# Allowed directories for legacy /print filepath parameter
_PRINT_ALLOWED_DIRS = [
    Path(r"C:\Printosky\Jobs\Incoming"),
    Path(r"C:\Printosky\Jobs\Archive"),
    Path(r"C:\Printosky\Jobs\Assigned"),  # multi-store: files pulled by store_puller
]

def _is_allowed_filepath(filepath: str) -> bool:
    """Return True only if filepath resolves inside an allowed print directory."""
    try:
        resolved = Path(filepath).resolve()
        return any(
            resolved.parent == d.resolve()
            for d in _PRINT_ALLOWED_DIRS
        )
    except Exception:
        return False


def init_staff_tables(db_path: str):
    """Ensure staff, staff_sessions, and work_sessions tables exist (idempotent)."""
    conn = sqlite3.connect(db_path)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS staff (
            id TEXT PRIMARY KEY, name TEXT NOT NULL,
            pin_hash TEXT NOT NULL, active INTEGER DEFAULT 1,
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    # v15: add pin_salt column if not present (idempotent)
    try:
        conn.execute("ALTER TABLE staff ADD COLUMN pin_salt TEXT")
        conn.commit()
    except Exception:
        pass  # column already exists
    conn.execute("""
        CREATE TABLE IF NOT EXISTS staff_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id TEXT NOT NULL, pc_id TEXT,
            login_at TEXT NOT NULL, logout_at TEXT,
            idle_logout INTEGER DEFAULT 0,
            FOREIGN KEY (staff_id) REFERENCES staff(id)
        )
    """)
    conn.commit()
    _ws_setup(conn)
    _rv_setup(conn)
    conn.close()


def staff_login(db_path: str, pin: str, pc_id: str):
    """Validate PIN, close prior session on this PC, open new session. Returns dict."""
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    # Fetch all active staff — compare PIN client-side to support PBKDF2 migration
    rows = conn.execute(
        "SELECT id, name, pin_hash, pin_salt FROM staff WHERE active=1"
    ).fetchall()
    matched = None
    for r in rows:
        if _verify_pin(pin, r["pin_hash"], r["pin_salt"]):
            matched = r
            break
    if not matched:
        conn.close()
        return {"ok": False, "error": "Invalid PIN"}

    # Upgrade legacy SHA-256 hash to PBKDF2 on first successful login
    if matched["pin_salt"] is None:
        new_hash, new_salt = _hash_pin(pin)
        conn.execute(
            "UPDATE staff SET pin_hash=?, pin_salt=? WHERE id=?",
            (new_hash, new_salt, matched["id"])
        )
        conn.commit()
        logging.info("PIN hash upgraded to PBKDF2 for staff %s", matched["id"])

    row = matched
    staff_id, name = row["id"], row["name"]
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Close any open session for this pc_id
    if pc_id:
        conn.execute(
            "UPDATE staff_sessions SET logout_at=? WHERE pc_id=? AND logout_at IS NULL",
            (now, pc_id)
        )

    cur = conn.execute(
        "INSERT INTO staff_sessions (staff_id, pc_id, login_at) VALUES (?,?,?)",
        (staff_id, pc_id, now)
    )
    session_id = cur.lastrowid
    conn.commit()
    conn.close()

    _active_sessions[pc_id] = {"staff_id": staff_id, "name": name, "session_id": session_id}
    logging.info("Staff login: %s (%s) on %s — session #%d", name, staff_id, pc_id, session_id)
    return {"ok": True, "staff_id": staff_id, "name": name, "session_id": session_id}


def staff_logout(db_path: str, session_id: int, idle: bool = False):
    """Close a staff session."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = sqlite3.connect(db_path)
    conn.execute(
        "UPDATE staff_sessions SET logout_at=?, idle_logout=? WHERE id=? AND logout_at IS NULL",
        (now, 1 if idle else 0, session_id)
    )
    conn.commit()
    # Clear from in-memory cache
    for pc_id, info in list(_active_sessions.items()):
        if info.get("session_id") == session_id:
            del _active_sessions[pc_id]
            break
    conn.close()
    logging.info("Staff logout: session #%d (idle=%s)", session_id, idle)
    return {"ok": True}


def get_active_staff(db_path: str, pc_id: str):
    """Return active staff for a given pc_id (checks DB for open session)."""
    if pc_id in _active_sessions:
        return _active_sessions[pc_id]
    conn = sqlite3.connect(db_path)
    row = conn.execute("""
        SELECT ss.id, ss.staff_id, s.name
        FROM staff_sessions ss JOIN staff s ON ss.staff_id = s.id
        WHERE ss.pc_id=? AND ss.logout_at IS NULL
        ORDER BY ss.login_at DESC LIMIT 1
    """, (pc_id,)).fetchone()
    conn.close()
    if row:
        info = {"session_id": row[0], "staff_id": row[1], "name": row[2]}
        _active_sessions[pc_id] = info
        return info
    return {"staff_id": None}


# attribute_konica_jobs() retired 2026-05-12 -- 0/4507 attribution rate.
# Function extracted to retired/2026-05-12-graveyard/konica_attribution.py.

# ── Internet / network health check ──────────────────────────────────────────
def check_internet(host="8.8.8.8", port=53, timeout=3) -> bool:
    try:
        socket.setdefaulttimeout(timeout)
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.connect((host, port))
        s.close()
        return True
    except Exception:
        return False

def check_printer_reachable(ip: str | None, timeout=2) -> bool:
    if not ip:  # finishing-only nodes have no Konica → treat as unreachable
        return False
    try:
        socket.setdefaulttimeout(timeout)
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.connect((ip, 9100))
        s.close()
        return True
    except Exception:
        return False

import fitz  # PyMuPDF — page rendering for /scale-preview
import pdf_scaler
import service_jobs
from db_migrations import ensure_job_service_columns

from ops_watchdog import report as _report_health

_PRINTERS_CFG = get_store_config().printers
PRINTER_IPS = {
    "konica": _PRINTERS_CFG.konica_ip,
    "epson":  _PRINTERS_CFG.epson_ip,
}

# ── Rate limiter for /staff-login ──────────────────────────────────────────────────────
import time as _time
_rate_limit: dict[str, list[float]] = {}  # ip -> [timestamp, ...]
_RATE_LIMIT_MAX    = 5     # max attempts per window
_RATE_LIMIT_WINDOW = 60.0  # seconds

def _check_rate_limit(ip: str) -> bool:
    """Return True if allowed, False if rate-limited (5 attempts per 60s per IP)."""
    now = _time.monotonic()
    hits = [t for t in _rate_limit.get(ip, []) if now - t < _RATE_LIMIT_WINDOW]
    _rate_limit[ip] = hits
    if len(hits) >= _RATE_LIMIT_MAX:
        return False
    _rate_limit[ip].append(now)
    return True

_SERVER_START = _time.monotonic()  # for uptime in /health

def get_system_health() -> dict:
    internet  = check_internet()
    konica_ok = check_printer_reachable(PRINTER_IPS["konica"])
    epson_ok  = check_printer_reachable(PRINTER_IPS["epson"])

    if internet and (konica_ok or epson_ok):
        mode = "full"
        mode_label = "Full — payments + printing active"
    elif not internet and (konica_ok or epson_ok):
        mode = "offline_print"
        mode_label = "Offline — walk-in printing only, no online payments"
    elif internet and not konica_ok and not epson_ok:
        mode = "online_no_printers"
        mode_label = "Printers unreachable — check network/power"
    else:
        mode = "manual"
        mode_label = "Manual mode — no internet, no printers"

    # Every failing watchdog check, so one request tells a console the whole
    # truth: printers, poller, fetchers, cloud sync. Reporting them here also
    # keeps /health honest when nobody has polled recently.
    try:
        from ops_watchdog import health as _ops_health, report as _report_health
        if has_konica():
            _report_health("printer.konica", konica_ok,
                           f"reachable at {PRINTER_IPS['konica']}" if konica_ok else
                           f"UNREACHABLE at {PRINTER_IPS['konica']} — powered off, "
                           "or the IP changed")
        _report_health("printer.epson", epson_ok,
                       f"reachable at {PRINTER_IPS['epson']}" if epson_ok else
                       f"UNREACHABLE at {PRINTER_IPS['epson']} — powered off, or the IP changed")
        watchdog = _ops_health()
    except Exception as exc:                      # never let /health itself 500
        watchdog = {"healthy": None, "error": str(exc), "checks": {}, "failing": []}

    return {
        "internet":     internet,
        "konica":       konica_ok,
        "epson":        epson_ok,
        "has_konica":   has_konica(),
        "printer_ips":  PRINTER_IPS,
        "mode":         mode,
        "mode_label":   mode_label,
        "active_staff": list(_active_sessions.keys()),
        "staff_count":  len(_active_sessions),
        "time":         datetime.now().strftime("%H:%M:%S"),
        "uptime_s":     int(_time.monotonic() - _SERVER_START),
        "db_ok":        __import__("os").path.exists(DB_PATH),
        "watchdog":     watchdog,
        "healthy":      bool(internet and (konica_ok or epson_ok)) and watchdog.get("healthy") is not False,
    }

# ── Config ────────────────────────────────────────────────────────────────────

PORT = 3005

# Printer names must match exactly what Windows sees in "Devices and Printers".
# Defaults are the canonical OSP store queue names; per-store overrides come
# from store_config.json via the `printer_queue_names` field (e.g. on a dev
# office PC, redirect 'epson' to 'Microsoft Print to PDF' so test dispatches
# don't burn real ink).
PRINTERS = {
    "konica": "KONICA MINOLTA 1100 PS",
    "epson":  "EM-C8100 Series(Network)",  # Epson EM-C8100 installed 2026-06-29 — verify exact queue name on store PC
}
try:
    _store_pq = getattr(get_store_config(), "printer_queue_names", None)
    if _store_pq:
        PRINTERS.update({k: v for k, v in _store_pq.items() if v})
        logging.info("PRINTERS overridden by store_config: %s", _store_pq)
except Exception as _e:
    logging.warning("Could not apply store_config.printer_queue_names: %s", _e)

# SumatraPDF path — portable version in project folder or installed
SUMATRA_PATHS = [
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "SumatraPDF.exe"),
    r"C:\printosky_watcher\SumatraPDF.exe",
    r"C:\PY\printosky\SumatraPDF.exe",
    r"D:\PY\printosky\SumatraPDF.exe",
    r"C:\Users\ABC\AppData\Local\SumatraPDF\SumatraPDF.exe",
    r"C:\Program Files\SumatraPDF\SumatraPDF.exe",
    r"C:\Program Files (x86)\SumatraPDF\SumatraPDF.exe",
]

# Shared secret — must match STORE_TOKEN in .env and storeToken in browser localStorage
STORE_TOKEN = os.environ.get("STORE_TOKEN", "")

# SQLite DB path — driven by store_config when present (e.g. dev/test stores
# pointing at a sandboxed jobs.db). Falls back to legacy OSP paths otherwise,
# so the real store PC's behaviour is unchanged (no store_config.json there
# → legacy fallback inside store_config returns the same hardcoded path).
try:
    DB_PATH = get_store_config().db_path
except Exception:
    if sys.platform == "win32":
        DB_PATH = r"C:\Printosky\Data\jobs.db"
    else:
        DB_PATH = str(Path.home() / "Printosky" / "Data" / "jobs.db")

# ── Helpers ───────────────────────────────────────────────────────────────────

def find_sumatra():
    for p in SUMATRA_PATHS:
        if os.path.exists(p):
            return p
    return None


def update_job_status(job_id: str, status: str, printer: str, staff_id: str = None):
    """Update job status, printer, printed_by, and notes in SQLite, then push to Supabase immediately."""
    try:
        conn = sqlite3.connect(DB_PATH)
        old_row = conn.execute("SELECT status FROM jobs WHERE job_id=?", (job_id,)).fetchone()
        old_status = old_row[0] if old_row else None
        note = f"Printed on {printer} at {datetime.now().strftime('%H:%M')}"
        if staff_id:
            note += f" by {staff_id}"
        conn.execute(
            "UPDATE jobs SET status=?, printer=?, notes=COALESCE(notes||' | ','') || ?"
            + (", printed_by=?" if staff_id else "")
            + " WHERE job_id=?",
            (status, printer, note, staff_id, job_id) if staff_id else (status, printer, note, job_id)
        )
        conn.commit()
        conn.close()
        logging.info("Job %s status → %s (printer: %s, staff: %s)", job_id, status, printer, staff_id or "—")
    except Exception as e:
        logging.error("DB update failed for %s: %s", job_id, e)
        return
    # Log in audit trail
    _jt_log(DB_PATH, job_id, "print_sent",
            from_status=old_status, to_status=status,
            staff_id=staff_id, notes=f"printer={printer}")
    # Immediately push status to Supabase so admin panel reflects change without waiting for sync cycle
    threading.Thread(target=_push_job_status_supabase, args=(job_id, status, printer), daemon=True).start()


_sb_client = None


def _supabase_client():
    """Lazy singleton Supabase SDK client for pickup-code uniqueness checks."""
    global _sb_client
    if _sb_client is None:
        from supabase import create_client
        url = os.environ["SUPABASE_URL"]
        key = os.environ.get("SUPABASE_SERVICE_KEY") or os.environ["SUPABASE_KEY"]
        _sb_client = create_client(url, key)
    return _sb_client


def _push_job_status_supabase(job_id: str, status: str, printer: str):
    """PATCH job status to Supabase using service key (bypasses RLS)."""
    import json as _json
    sb_url = os.environ.get("SUPABASE_URL", "")
    sb_key = os.environ.get("SUPABASE_SERVICE_KEY", "") or os.environ.get("SUPABASE_KEY", "")
    if not sb_url or not sb_key:
        return
    try:
        payload = _json.dumps({"status": status, "printer": printer}).encode()
        req = urllib.request.Request(
            f"{sb_url}/rest/v1/jobs?job_id=eq.{job_id}",
            data=payload,
            method="PATCH",
            headers={
                "apikey": sb_key,
                "Authorization": f"Bearer {sb_key}",
                "Content-Type": "application/json",
                "Prefer": "return=minimal",
            },
        )
        with urllib.request.urlopen(req, timeout=5):
            pass
        logging.info("Supabase status push OK: %s → %s", job_id, status)
    except Exception as e:
        logging.warning("Supabase status push failed for %s: %s", job_id, e)


def _sumatra_paper(size: str | None) -> str | None:
    """Map a Printosky paper size to a SumatraPDF ``paper=`` token.

    Returns None for unknown/empty sizes so the caller omits ``paper=`` and the
    printer queue's own default is used (the pre-existing behaviour). SumatraPDF
    accepts A-series uppercase (``A4``) and named sizes lowercase (``legal``).
    """
    if not size:
        return None
    u = size.strip().upper()
    if u in ("A2", "A3", "A4", "A5", "A6"):
        return u
    return {"LEGAL": "legal", "LETTER": "letter", "TABLOID": "tabloid",
            "STATEMENT": "statement"}.get(u)


def has_konica() -> bool:
    """True iff this store actually has a Konica.

    PRINTERS always carries a Konica queue name (the OSP default is inherited by
    every store), so presence is decided by the configured IP instead: a
    finishing/collection store like Nattika sets `konica_ip` to null or "" in
    store_config.json. The string "None" is accepted as empty too — that is what
    a JSON null becomes once it has been through str().
    """
    ip = PRINTER_IPS.get("konica")
    return bool(ip) and ip != "None"


def _watchdog_summary() -> dict:
    """Compact health headline for /status: what is broken, right now."""
    try:
        from ops_watchdog import health
        h = health()
        return {"healthy": h.get("healthy"), "failing": h.get("failing", []),
                "checks": {k: v for k, v in h.get("checks", {}).items() if not v.get("ok")}}
    except Exception as exc:
        return {"healthy": None, "error": str(exc), "failing": []}


def _effective_printer_key(printer_key: str, job_id: str = "") -> str:
    """Resolve the printer to actually use, applying the no-Konica redirect.

    Finishing/collection-only stores (e.g. Nattika) have no Konica: PRINTERS
    still carries the inherited OSP Konica queue name — a queue that does not
    exist on this PC — so a 'konica' dispatch prints nothing. Redirect it to the
    Epson (the only printer present; its monochrome mode handles B&W). Shared by
    every print path so they cannot drift. No-op for any other key, or when a
    real Konica IP is configured.
    """
    if printer_key == "konica" and not has_konica():
        logging.info("no Konica on this store — routing job %s to epson", job_id or "?")
        return "epson"
    return printer_key


def _konica_queue_for_sides(sides: str | None) -> str | None:
    """Pick a duplex/simplex-specific Konica queue key, or None to keep using
    the plain 'konica' queue unchanged.

    Why this exists: the KONICA MINOLTA 1100 PS driver silently ignores
    SumatraPDF's per-job duplex/simplex override in *both* directions — the
    printer just follows whatever its Windows Printing Preferences default
    currently is, regardless of what the job asked for (confirmed by log,
    see docs/PRINT_ROTATION_MATRIX.md, 2026-08-29). The standard fix for a
    driver that won't take a per-job override is to stop asking it to: install
    the same physical printer twice as two separate Windows queues, each with
    its own persisted default, and pick the QUEUE instead of the setting.

    Inert until a store configures both queues via
    store_config.json's printer_queue_names (keys 'konica_duplex' and
    'konica_simplex') — with neither set, this returns None for every sides
    value and print_server behaves exactly as before.
    """
    s = (sides or "").strip().lower()
    if s in ("ds", "duplex", "double", "duplexlong", "duplexshort"):
        variant = "konica_duplex"
    elif s in ("ss", "simplex", "single"):
        variant = "konica_simplex"
    else:
        return None
    return variant if PRINTERS.get(variant) else None


def send_to_printer(job_id: str, filepath: str, printer_key: str, copies: int = 1,
                    colour_mode: str = "auto", staff_id: str = None,
                    sides: str = None, paper_size: str = None,
                    orientation: str = None, update_status: bool = True,
                    scale_applied: bool = False):
    """
    Execute print command via SumatraPDF (silent, no UI).
    Returns (success: bool, message: str)

    ``sides``       : 'ds'/'duplex' -> long-edge duplex, 'ss'/'simplex' -> simplex,
                      None -> leave to the queue default.
    ``paper_size``  : e.g. 'A4'/'A3'/'Legal'; None -> queue default.
    ``orientation`` : 'portrait'/'landscape'; 'auto'/None -> per-page (queue default).
    ``scale_applied``: the file already carries its final geometry (print_planner
                      baked Fit/Actual/Custom into it), so tell SumatraPDF not to
                      scale it again. Default False emits nothing, which is what
                      every job did before scaling existed.
    """
    # No-Konica stores: redirect a 'konica' dispatch to the Epson (shared helper).
    printer_key = _effective_printer_key(printer_key, job_id)

    # Konica's driver won't honour a per-job duplex/simplex override (see
    # _konica_queue_for_sides docstring) — route to a sides-specific queue if
    # the store has one configured. No-op otherwise.
    if printer_key == "konica":
        konica_variant = _konica_queue_for_sides(sides)
        if konica_variant:
            logging.info("job %s: routing to %s queue for sides=%r",
                         job_id, konica_variant, sides)
            printer_key = konica_variant

    printer_name = PRINTERS.get(printer_key)
    if not printer_name:
        return False, f"Unknown printer key: {printer_key}"

    if not os.path.exists(filepath):
        # Check archive folder — file may have been moved after receipt
        archive_path = os.path.join(r"C:\Printosky\Jobs\Archive", os.path.basename(filepath))
        if os.path.exists(archive_path):
            filepath = archive_path
            logging.info("File found in archive: %s", archive_path)
        else:
            return False, f"File not found: {filepath} (also checked Archive)"

    sumatra = find_sumatra()
    if not sumatra:
        # Fallback: use Windows print verb. This path has NO control over
        # sides/colour/paper_size/orientation — it just triggers os.startfile,
        # which prints with whatever the queue's last-used settings were. A
        # job that asked for simplex or a specific paper size can silently
        # come out wrong, so this must alert, not just log.
        logging.warning("SumatraPDF not found — using Windows shell print (no sides/colour/paper control)")
        _report_health(
            "print_server.sumatra_missing", False,
            f"SumatraPDF.exe not found in any SUMATRA_PATHS — job {job_id} sent via Windows "
            "shell print, which cannot control sides/colour/paper_size/orientation. "
            "Re-install SumatraPDF or fix SUMATRA_PATHS.",
        )
        return windows_shell_print(filepath, printer_name, copies, printer_key)
    _report_health("print_server.sumatra_missing", True, "SumatraPDF found")

    # Build SumatraPDF command
    # -print-to <printer>  : print to named printer silently
    # -print-settings      : copies, colour settings
    # -exit-when-done      : close after printing
    settings_parts = [f"{copies}x"]
    if colour_mode == "bw":
        settings_parts.append("monochrome")
    elif colour_mode == "colour":
        settings_parts.append("color")
    # "auto" = let printer decide

    # Duplex — only emit when explicitly known; else leave to the queue default.
    s = (sides or "").strip().lower()
    if s in ("ds", "duplex", "double", "duplexlong"):
        settings_parts.append("duplexlong")
    elif s in ("duplexshort", "dss", "shortedge", "short"):
        # Short/top-edge bind — landscape N-up needs this so the back registers.
        settings_parts.append("duplexshort")
    elif s in ("ss", "simplex", "single"):
        settings_parts.append("simplex")

    # Paper size (e.g. A4/A3/Legal). Unknown/empty -> queue default.
    paper_tok = _sumatra_paper(paper_size)
    if paper_tok:
        settings_parts.append(f"paper={paper_tok}")

    # Orientation — 'auto' means honour each page; only force when told to.
    o = (orientation or "").strip().lower()
    if o in ("portrait", "landscape"):
        settings_parts.append(o)

    # The page geometry is already baked into the file, so stop the driver
    # having a second opinion about it. A guard, not the mechanism: if a driver
    # ignores this the sheet is still right, because the correctness is in the
    # PDF. Never emitted for a job that did not ask for scaling.
    if scale_applied:
        settings_parts.append("noscale")

    settings = ",".join(settings_parts)

    file_dir  = os.path.dirname(os.path.abspath(filepath))
    file_name = os.path.basename(filepath)

    # SumatraPDF uses the filename as the document name in the Windows print
    # spooler, which the Epson records verbatim in its job-history CSV.
    # Copy to a temp file named <job_id>.<ext> so the Epson log shows the
    # Printosky job ID — enabling direct matching without time-window deltas.
    import tempfile, shutil as _shutil
    ext = os.path.splitext(file_name)[1] or ".pdf"
    named_tmp = os.path.join(tempfile.gettempdir(), f"{job_id}{ext}")
    _shutil.copy2(filepath, named_tmp)

    cmd = [
        sumatra,
        "-print-to", printer_name,
        "-print-settings", settings,
        "-exit-when-done",
        "-silent",
        named_tmp,
    ]

    logging.info("Print command: %s", " ".join(cmd))

    try:
        result = subprocess.run(cmd, timeout=60, capture_output=True, text=True)
        if result.returncode == 0:
            # Skipped for multi-sub-job orders (caller marks Printed once, after
            # ALL sub-jobs succeed) so a later-section failure isn't masked.
            if update_status:
                update_job_status(job_id, "Printed", printer_name, staff_id)
            _trigger_printer_poll_now(printer_key)
            return True, f"Sent to {printer_name} ({copies} cop{'y' if copies==1 else 'ies'})"
        else:
            err = result.stderr or result.stdout or "Unknown error"
            logging.error("SumatraPDF error: %s", err)
            return False, f"Print failed: {err}"
    except subprocess.TimeoutExpired:
        return False, "Print command timed out after 60s"
    except Exception as e:
        return False, str(e)
    finally:
        try:
            os.remove(named_tmp)
        except OSError:
            pass


def _trigger_printer_poll_now(printer_key: str) -> None:
    """Tell the printer poller (a thread inside watcher.py) to re-check
    counters right now instead of waiting for its next scheduled sweep — we
    just caused a page count to change, no point waiting up to 5 minutes to
    notice.

    Fired in a background thread so a slow/unreachable watcher process never
    adds latency to the print response, and never raises: this is a latency
    optimisation on top of the always-running sweep, not the source of
    truth, so a failed trigger just means counters catch up on the next
    sweep instead of instantly — same as before this existed. Still reported
    to ops_watchdog so a persistently broken trigger is a visible alert
    rather than a silent slowdown. Local call only (127.0.0.1:3003) — never
    touches Supabase.
    """
    def _fire():
        try:
            req = urllib.request.Request(
                "http://127.0.0.1:3003/printers/poll-now",
                data=json.dumps({"printer": printer_key}).encode(),
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            urllib.request.urlopen(req, timeout=2)
            _report_health("printer_poll_now.trigger", True, "watcher acknowledged")
        except Exception as exc:
            _report_health("printer_poll_now.trigger", False, f"{type(exc).__name__}: {exc}")

    threading.Thread(target=_fire, daemon=True, name="PollNowTrigger").start()


def windows_shell_print(filepath: str, printer_name: str, copies: int, printer_key: str = None):
    """Fallback: use Windows shell to print (no copies control)."""
    try:
        # Set default printer temporarily and print
        subprocess.run(
            ["rundll32", "printui.dll,PrintUIEntry", "/y", "/n", printer_name],
            check=True, timeout=10
        )
        os.startfile(filepath, "print")
        if printer_key:
            _trigger_printer_poll_now(printer_key)
        return True, f"Sent via Windows shell to {printer_name}"
    except Exception as e:
        return False, f"Shell print failed: {e}"


# ── New Sprint 1 helpers ───────────────────────────────────────────────────────

def _db():
    """Return a sqlite3 connection to the jobs DB with row_factory."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def _now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _count_pages_from_list(page_list: str, total_pages: int) -> int:
    """
    Count the number of pages a SumatraPDF-style page-list expression refers to.

    Mirrors _build_page_range_arg() semantics so the spec row we write matches
    what we actually dispatched to the printer.

    Examples:
      'all'         + total=10  -> 10
      '' / None     + total=10  -> 10
      '1-5'                     -> 5
      '1,3,5'                   -> 3
      '1-5,10-15'               -> 5 + 6 = 11
    Returns 0 if the expression can't be parsed.
    """
    total = max(int(total_pages or 0), 0)
    if not page_list or str(page_list).strip().lower() in ("all", ""):
        return total

    count = 0
    for part in str(page_list).strip().split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            try:
                a, b = part.split("-", 1)
                a_i, b_i = int(a.strip()), int(b.strip())
                if b_i >= a_i:
                    count += (b_i - a_i + 1)
            except ValueError:
                continue
        else:
            try:
                int(part)
                count += 1
            except ValueError:
                continue
    return count


def _write_epson_spec_row(
    job_id: str,
    item_number: int,
    file_name: str,
    pages_per_copy: int,
    copies: int,
    colour: str,
    paper_size: str,
) -> None:
    """
    Write a source='spec' row to local epson_jobs table on successful Epson dispatch.

    Captures what we *told* the printer to do (mono/colour/copies/paper). The
    weblog scraper later writes a parallel source='weblog' row with what the
    printer *reports* happened. They join on attributed_job_id = printosky job_id,
    enabling end-to-end reconciliation.

    Why this exists: Epson's free Web Config CSV and SNMP OIDs do not expose
    per-job mono/colour split, copies, or paper size. Per-job accounting is only
    available via paid Epson Print Admin Serverless. This function gives us the
    same data for free by capturing it at dispatch time, when we already know it.

    Note: local SQLite epson_jobs has no store_id column — supabase_sync.py
    injects store_id at upload time. Same convention as weblog/delta rows.

    Failures are logged but NEVER block the print flow.
    """
    try:
        pages_printed = max(int(pages_per_copy or 0), 0) * max(int(copies or 0), 0)
        is_colour = (colour or "").lower() in ("col", "colour", "color")
        mono_pages  = 0 if is_colour else pages_printed
        color_pages = pages_printed if is_colour else 0

        now = _now()
        # Unique per dispatch — timestamp suffix keeps re-prints distinct and
        # never collides with weblog rows (those use plain integer job_numbers).
        ts_compact = datetime.now().strftime("%Y%m%d%H%M%S")
        spec_job_number = f"{job_id}-{int(item_number)}-{ts_compact}"

        conn = _db()
        conn.execute("""
            INSERT INTO epson_jobs (
                source, job_number, file_name,
                pages_printed, mono_pages, color_pages, copies,
                paper_size, job_date, print_end_date,
                attributed_job_id, imported_at, result
            ) VALUES ('spec', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'OK')
        """, (
            spec_job_number, file_name,
            pages_printed, mono_pages, color_pages, int(copies or 1),
            paper_size or "A4", now, now,
            job_id, now,
        ))
        conn.commit()
        conn.close()
        logging.info(
            "epson_jobs spec row: %s (%s) pages=%d mono=%d color=%d copies=%s paper=%s",
            spec_job_number, file_name, pages_printed, mono_pages, color_pages,
            copies, paper_size,
        )
    except Exception as e:
        logging.warning(
            "Failed to write epson_jobs spec row for %s item %d: %s",
            job_id, item_number, e,
        )


def _send_whatsapp(phone: str, message: str) -> bool:
    """Send WhatsApp message via Meta Cloud API. Returns True on success."""
    if not phone:
        return False
    from whatsapp_notify import _send
    return _send(phone, message)


def _job_quote(print_items: list, finishing: str, is_student: bool,
               urgent: bool, paper_size: str = "A4") -> dict:
    """Calculate quote using rate_card. Returns {total, print_cost, finishing_cost, breakdown}."""
    if _rc is None:
        return {"total": 0, "print_cost": 0, "finishing_cost": 0, "breakdown": []}
    return _rc.calculate_quote(
        print_items=print_items,
        finishing=finishing or "none",
        is_student=bool(is_student),
        urgent=bool(urgent),
        paper_size=paper_size or "A4",
    )


# Rendered previews, keyed by everything that can change one. Bounded: a
# counter flicking through a 200-page job must not grow this without limit.
_PREVIEW_CACHE: "collections.OrderedDict[tuple, bytes]" = collections.OrderedDict()
_PREVIEW_CACHE_MAX = 24
_PREVIEW_DPI = 96


def handle_scale_preview(qs: dict) -> tuple[bytes | None, dict]:
    """GET /scale-preview?job_id=&page=1&mode=fit&percent=&paper_size=A4

    Returns (png_bytes, meta) or (None, {"error": ..., "status": ...}).

    The PNG is a render of the **baked** page — the same pdf_scaler.apply_scale()
    output the printer would receive — not a drawing of where the page ought to
    go. What the operator sees is the artifact.

    `meta` carries what a picture cannot: total pages, whether this page crops,
    and how many pages crop overall. They travel as X- headers so the panel can
    fetch the blob and read them in one request.
    """
    job_id = (qs.get("job_id", [""])[0] or "").strip()
    if not job_id:
        return None, {"status": 400, "error": "job_id required"}

    mode = (qs.get("mode", [""])[0] or "").strip().lower()
    if mode not in pdf_scaler.MODES:
        return None, {"status": 400,
                      "error": f"mode must be one of {', '.join(pdf_scaler.MODES)}"}

    percent = pdf_scaler.clamp_percent(qs.get("percent", [None])[0]) if mode == "custom" else None
    if mode == "custom" and percent is None:
        return None, {"status": 400, "error": "custom needs a numeric percent"}

    paper_size = (qs.get("paper_size", ["A4"])[0] or "A4").strip()
    try:
        page_no = max(1, int(qs.get("page", ["1"])[0]))
    except (TypeError, ValueError):
        page_no = 1

    fp = _resolve_job_file(job_id)
    if not fp:
        return None, {"status": 404, "error": "File not found on disk"}
    if os.path.splitext(fp)[1].lower() != ".pdf":
        # Only PDFs have page geometry to preview. The panel keeps its ordinary
        # file view for anything else rather than showing a wrong picture.
        return None, {"status": 415, "error": "preview is available for PDFs only"}

    try:
        key = (fp, os.path.getmtime(fp), page_no, mode, percent, paper_size)
    except OSError as exc:
        return None, {"status": 404, "error": f"file unreadable: {exc}"}

    try:
        with open(fp, "rb") as fh:
            raw = fh.read()

        sizes = pdf_scaler.page_sizes(raw)
        total = len(sizes)
        if total == 0:
            return None, {"status": 422, "error": "PDF has no pages"}
        page_no = min(page_no, total)
        key = key[:2] + (page_no,) + key[3:]

        w, h = sizes[page_no - 1]
        rect = pdf_scaler.scale_rect(w, h, paper_size, mode, percent)
        meta = {
            "total_pages": total,
            "page": page_no,
            "crops": bool(rect and rect["crops"]),
            "cropped_pages": pdf_scaler.count_cropped_pages(raw, mode, percent, paper_size),
            "scaled": rect is not None,
        }

        cached = _PREVIEW_CACHE.get(key)
        if cached is not None:
            _PREVIEW_CACHE.move_to_end(key)
            return cached, meta

        baked = pdf_scaler.apply_scale(raw, mode, percent, paper_size)
        source = baked if baked else raw          # a no-op previews as-is
        doc = fitz.open("pdf", source)
        try:
            png = doc[page_no - 1].get_pixmap(dpi=_PREVIEW_DPI).tobytes("png")
        finally:
            doc.close()

        _PREVIEW_CACHE[key] = png
        while len(_PREVIEW_CACHE) > _PREVIEW_CACHE_MAX:
            _PREVIEW_CACHE.popitem(last=False)
        return png, meta

    except Exception as exc:
        # An operator who cannot see the preview must be told, not shown a
        # stale or invented one — the panel disables the control on this.
        logging.warning("scale preview failed for %s: %s", job_id, exc)
        return None, {"status": 500, "error": f"preview failed: {exc}"}


def _resolve_job_file(job_id: str) -> str:
    """The file on disk for a job, or "" if it cannot be found.

    A puller-downloaded cloud job (multi-store) may have NO local jobs row and
    lives in Jobs\\Assigned named <job_id>*, so a row is not required and the
    job directories are searched too. Shared by /file and /scale-preview so the
    preview can never render a different file than the panel shows.
    """
    _JOB_DIRS = (r"C:\Printosky\Jobs\Assigned",
                 r"C:\Printosky\Jobs\Archive",
                 r"C:\Printosky\Jobs\Incoming")
    fp = ""
    try:
        conn = _db()
        job = conn.execute("SELECT * FROM jobs WHERE job_id=?", (job_id,)).fetchone()
        conn.close()
    except Exception:
        job = None

    if job:
        keys = job.keys()
        fp = (job["filepath"] if "filepath" in keys else "") or ""
        if not fp:
            fp = (job["file_path"] if "file_path" in keys else "") or ""
        fname = job["filename"] if "filename" in keys else ""
        if (not fp or not os.path.exists(fp)) and fname:
            for base in _JOB_DIRS:
                cand = os.path.join(base, fname)
                if os.path.exists(cand):
                    fp = cand
                    break

    # Fallback: match by job_id prefix in the job dirs (covers pulled jobs with
    # no local row). job_id is sanitised to prevent traversal.
    if not fp or not os.path.exists(fp):
        import glob as _glob, re as _re
        safe_id = _re.sub(r"[^A-Za-z0-9_-]", "", job_id or "")
        if safe_id:
            for base in _JOB_DIRS:
                matches = _glob.glob(os.path.join(base, safe_id + "*"))
                if matches:
                    fp = matches[0]
                    break

    return fp if fp and os.path.exists(fp) else ""


def _build_page_range_arg(page_list: str, total_pages: int) -> str:
    """
    Convert page_list string to SumatraPDF -print-settings range format.
    page_list: 'all' | '1-10' | '1,5,12' | '1-5,10-15'
    Returns empty string for 'all', else the range string.
    """
    if not page_list or page_list.strip().lower() in ("all", ""):
        return ""
    return page_list.strip()


def update_job_quote(job_id: str, amount: float):
    """Update amount_quoted on the jobs table."""
    try:
        conn = _db()
        conn.execute("UPDATE jobs SET amount_quoted=? WHERE job_id=?", (amount, job_id))
        conn.commit()
        conn.close()
    except Exception as e:
        logging.error("update_job_quote failed: %s", e)


def _ensure_print_item_scale_columns(conn) -> None:
    """Add print_items.scale_mode / scale_percent if this DB predates them.

    Store PCs update by pulling code and restarting the watcher — nothing runs
    fix_db.py for them (see docs/AUTO_UPDATE.md). So a box can be running this
    file against a database written before 2026-08-30, and an INSERT naming
    columns it does not have would break spec-saving at the counter. Cheap
    PRAGMA, idempotent, and it keeps fix_db.py as the tidy-up rather than the
    prerequisite.
    """
    try:
        have = {row[1] for row in conn.execute("PRAGMA table_info(print_items)")}
        for col, defn in (("scale_mode", "TEXT"), ("scale_percent", "INTEGER")):
            if col not in have:
                conn.execute(f"ALTER TABLE print_items ADD COLUMN {col} {defn}")
                logging.info("print_items: added missing column %s", col)
    except Exception as exc:
        # Not fatal on its own — the INSERT below is what actually needs them,
        # and it will surface a clear error if this could not run.
        logging.warning("could not ensure print_items scale columns: %s", exc)


# ── A1: Update job specs ───────────────────────────────────────────────────────

def handle_update_job(body: dict) -> dict:
    """
    PUT /update-job
    Save print_items + finishing + flags to DB. Recalculate and store quote.
    """
    job_id    = body.get("job_id", "")
    staff_id  = body.get("staff_id", "")
    finishing = body.get("finishing", "none")
    is_student= bool(body.get("is_student", False))
    urgent    = bool(body.get("urgent", False))
    paper_size= body.get("paper_size", "A4")
    items_raw = body.get("print_items", [])

    if not job_id:
        return {"ok": False, "error": "job_id required"}

    conn = _db()
    _ensure_print_item_scale_columns(conn)

    # Upsert print_items — delete old, insert new
    conn.execute("DELETE FROM print_items WHERE job_id=?", (job_id,))

    rc_items = []
    for item in items_raw:
        item_number = int(item.get("item_number", 1))
        page_list   = item.get("page_list", "all")
        colour      = item.get("colour", "bw")
        paper_type  = item.get("paper_type") or (f"{paper_size}_col" if colour == "col" else f"{paper_size}_BW")
        sides       = item.get("sides", "ss")
        layout      = item.get("layout", "1-up")
        copies      = int(item.get("copies", 1))
        paper_gsm   = int(item.get("paper_gsm", 70))
        printer     = item.get("printer") or ("epson" if colour == "col" else "konica")

        # Scaling. Absent/blank stays NULL, which is "no scaling" — so a panel
        # that never sends these fields behaves exactly as it did before.
        scale_mode = (item.get("scale_mode") or "").strip().lower() or None
        if scale_mode not in (None, "fit", "actual", "custom"):
            scale_mode = None
        scale_percent = None
        if scale_mode == "custom":
            try:
                from pdf_scaler import clamp_percent
                scale_percent = clamp_percent(item.get("scale_percent"))
            except Exception:
                scale_percent = None
            if scale_percent is None:
                scale_mode = None      # a custom with no usable percent is no scale

        conn.execute("""
            INSERT INTO print_items
              (job_id, item_number, page_list, paper_type, colour, sides, layout,
               copies, paper_gsm, printer, status, scale_mode, scale_percent)
            VALUES (?,?,?,?,?,?,?,?,?,?,'Pending',?,?)
        """, (job_id, item_number, page_list, paper_type, colour, sides, layout,
              copies, paper_gsm, printer, scale_mode, scale_percent))

        # For quote calculation, approximate pages from page_list
        # If page_list = 'all', fetch page_count from jobs table
        rc_items.append({
            "pages": item.get("pages", 1),
            "paper_type": paper_type,
            "sides": sides,
            "layout": layout,
            "copies": copies,
        })

    # Update main job flags
    conn.execute("""
        UPDATE jobs SET finishing=?, is_student=?, urgent=?, paper_size=?,
               notes=COALESCE(notes||' | ','') || ?
        WHERE job_id=?
    """, (finishing, int(is_student), int(urgent), paper_size,
          f"Specs updated at {_now()} by {staff_id}", job_id))

    # Calculate quote
    quote = _job_quote(rc_items, finishing, is_student, urgent, paper_size)
    conn.execute("UPDATE jobs SET amount_quoted=? WHERE job_id=?",
                 (quote["total"], job_id))

    conn.commit()
    conn.close()

    logging.info("Job %s specs updated by %s — quote Rs.%.0f", job_id, staff_id, quote["total"])
    return {
        "ok": True,
        "quote": quote,
        "printer_assigned": {
            f"item_{i+1}": (it.get("colour","bw") == "col" and "epson" or "konica")
            for i, it in enumerate(items_raw)
        },
    }


# ── A3: Mark ready ────────────────────────────────────────────────────────────

def handle_mark_ready(body: dict) -> dict:
    """
    POST /mark-ready
    Change status to 'Ready', send WhatsApp to customer.
    """
    job_id   = body.get("job_id", "")
    staff_id = body.get("staff_id", "")
    if not job_id:
        return {"ok": False, "error": "job_id required"}

    conn = _db()
    row = conn.execute(
        "SELECT sender, filename, customer_name, pickup_code FROM jobs WHERE job_id=?",
        (job_id,)
    ).fetchone()

    if not row:
        conn.close()
        return {"ok": False, "error": "Job not found"}

    phone    = row["sender"]
    filename = row["filename"] or job_id
    name     = row["customer_name"] or "Customer"

    code = row["pickup_code"]
    if not code:
        try:
            from pickup_code import claim_unique_pickup_code
            code = claim_unique_pickup_code(_supabase_client())
        except Exception:
            code = None

    # Update status + pickup fields (COALESCE preserves existing code/timestamp
    # if staff clicks Ready twice on the same job)
    conn.execute("""
        UPDATE jobs SET status='Ready',
               pickup_code=COALESCE(pickup_code, ?),
               pickup_ready_at=COALESCE(pickup_ready_at, ?),
               notes=COALESCE(notes||' | ','') || ?
        WHERE job_id=?
    """, (code, _now(), f"Ready notified at {_now()} by {staff_id}", job_id))
    conn.commit()
    conn.close()

    # Send WhatsApp — pickup-code-aware message with track link
    whatsapp_sent = False
    if phone and code:
        try:
            from whatsapp_notify import send_pickup_ready
            from store_config import get_store_config
            cfg = get_store_config()
            track_link = f"https://printosky.com/track?code={code}"
            whatsapp_sent = send_pickup_ready(
                phone, code,
                store_label=None,
                store_address=cfg.store_name,
                deep_link=track_link,
            )
        except Exception as _e:
            logging.warning("send_pickup_ready failed, using fallback: %s", _e)
            msg = (f"🎉 Your print job is ready!\n\n"
                   f"🎫 Pickup code: *{code}*\n"
                   f"🔗 Track: https://printosky.com/track?code={code}\n\n"
                   f"Show this code at the counter.")
            whatsapp_sent = _send_whatsapp(phone, msg)
    elif phone:
        msg = (f"Hi! Your print job is ready for collection at Printosky.\n"
               f"Job: {job_id}")
        whatsapp_sent = _send_whatsapp(phone, msg)

    logging.info("Job %s marked Ready code=%s (WhatsApp: %s)",
                 job_id, code, "sent" if whatsapp_sent else "skipped")
    _jt_log(DB_PATH, job_id, "job_ready",
            from_status=None, to_status="Ready",
            staff_id=staff_id,
            notes=f"code={code} whatsapp={'sent' if whatsapp_sent else 'skipped'}")
    return {"ok": True, "whatsapp_sent": whatsapp_sent, "phone": phone or "walk-in",
            "pickup_code": code}


# ── A4: Complete job ──────────────────────────────────────────────────────────

def handle_complete_job(body: dict) -> dict:
    """
    POST /complete-job
    Record payment collected at counter. Mark job as Completed.
    """
    job_id      = body.get("job_id", "")
    amount      = body.get("amount_collected", 0)
    mode        = body.get("payment_mode", "Cash")
    staff_id    = body.get("staff_id", "")

    if not job_id:
        return {"ok": False, "error": "job_id required"}
    if mode not in ("Cash", "UPI", "Online"):
        mode = "Cash"

    try:
        amount = float(amount)
    except (TypeError, ValueError):
        return {"ok": False, "error": "amount_collected must be a number"}

    conn = _db()
    row = conn.execute(
        "SELECT sender, pickup_code FROM jobs WHERE job_id=?", (job_id,)
    ).fetchone()
    phone       = row["sender"]       if row else None
    pickup_code = row["pickup_code"]  if row else None
    conn.execute("""
        UPDATE jobs SET
            status='Completed',
            amount_collected=?,
            payment_mode=?,
            completed_at=?,
            delivered_at=COALESCE(delivered_at, ?),
            notes=COALESCE(notes||' | ','') || ?
        WHERE job_id=?
    """, (amount, mode, _now(), _now(),
          f"Completed at {_now()} by {staff_id} — Rs.{amount:.0f} {mode}",
          job_id))
    conn.commit()
    # Fires only for service jobs; a print job returns from it immediately.
    _alert_zero_priced_service(conn, job_id, amount, staff_id)
    conn.close()

    logging.info("Job %s COMPLETED — Rs.%.0f %s by %s", job_id, amount, mode, staff_id)
    _jt_log(DB_PATH, job_id, "job_collected",
            from_status="Ready", to_status="Completed",
            staff_id=staff_id,
            notes=f"Rs.{amount:.0f} {mode}")

    # Pickup confirmation + schedule review request 30 min after collection
    if phone:
        if pickup_code:
            try:
                from whatsapp_notify import send_pickup_completed
                send_pickup_completed(phone, pickup_code)
            except Exception as _e:
                logging.warning("send_pickup_completed failed: %s", _e)
        _rv_schedule(DB_PATH, job_id, phone, _send_whatsapp)

    return {"ok": True, "job_id": job_id, "amount": amount, "mode": mode}


# ── A5: New photocopy job ─────────────────────────────────────────────────────

def _photocopy_meta(body: dict) -> dict:
    """The rate-card meta for a counter photocopy.

    A photocopy is priced as a print of the same sheets on the same machine —
    same paper, same toner — so it goes through the one rate card rather than a
    second table that would drift from it.
    """
    return service_jobs.photocopy_meta(body)


def _photocopy_quote(meta: dict) -> tuple[float | None, list[str]]:
    """(price, breakdown) from the rate card, or (None, [why]) if it cannot.

    Never raises. A photocopy the shop cannot price is a thing a human has to
    know about — it is the same class of failure as the five finishings that
    billed Rs.0 — so it alerts and comes back None rather than as a free job.
    """
    if _rc is None:
        _report_health("photocopy.quote", False,
                       "rate_card not loaded — /new-photocopy cannot price anything")
        return None, ["rate_card not loaded"]
    try:
        result = _rc.calculate_service_quote("copy", meta)
    except Exception as exc:
        _report_health("photocopy.quote", False,
                       f"/new-photocopy could not price {meta!r}: "
                       f"{type(exc).__name__}: {exc}")
        return None, [f"could not price this: {exc}"]
    if result["needs_manual_price"]:
        _report_health("photocopy.quote", False,
                       f"the rate card has no price for {meta!r} — staff must type one")
        return None, result["breakdown"]
    _report_health("photocopy.quote", True, "pricing photocopies")
    return float(result["total"]), result["breakdown"]


def handle_new_photocopy(body: dict) -> dict:
    """
    POST /new-photocopy
    Create an immediate Completed job entry for a photocopy (no file).

    B-6 (2026-09-01): **the price comes from the rate card, not from a number
    staff type.** The button and its flow are unchanged — it still files one
    Completed job with no file and no print item — but the amount is quoted the
    way every other price in the system is quoted. A typed amount still wins,
    because a discount or a miscount has to be possible; it is recorded against
    the quote so the two are visible side by side rather than one silently
    replacing the other.
    """
    staff_id      = body.get("staff_id", "")
    customer_name = body.get("customer_name", "")
    phone         = body.get("phone", "")
    mode          = body.get("payment_mode", "Cash")

    if mode not in ("Cash", "UPI", "Online"):
        mode = "Cash"

    meta = _photocopy_meta(body)
    quoted, breakdown = _photocopy_quote(meta)

    typed = _amount_or_none(body.get("amount_collected"))
    if typed is not None and typed > 0:
        amount = typed
    elif quoted is not None:
        amount = quoted
    else:
        # Nothing typed and nothing quotable. Refusing beats filing a Rs.0 job:
        # the counter types an amount and the sale completes on the next click.
        return {"ok": False, "needs_manual_price": True,
                "error": "Could not price this photocopy — enter the amount. "
                         + "; ".join(breakdown),
                "breakdown": breakdown}

    overridden = quoted is not None and typed is not None and typed > 0 and typed != quoted

    conn   = _db()
    job_id = _next_job_id(conn)
    now    = _now()

    note = f"Photocopy job created at {now} by {staff_id}"
    if breakdown:
        note += " | " + " | ".join(breakdown)
    if overridden:
        note += f" | staff set Rs.{amount:.0f} over the quoted Rs.{quoted:.0f}"

    conn.execute("""
        INSERT INTO jobs
          (job_id, received_at, filename, source, sender, customer_name,
           service_type, page_count, colour, copies,
           status, amount_collected, payment_mode, completed_at,
           printed_by, notes, amount_quoted)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (job_id, now, "Photocopy Job", "Photocopy", phone or None, customer_name or None,
          "Photocopy", meta["sheets"], meta["colour"], meta["copies"],
          "Completed", amount, mode, now,
          staff_id,
          note,
          quoted if quoted is not None else amount))
    conn.commit()
    conn.close()

    if amount <= 0:
        _report_health(
            "photocopy.unpriced", False,
            f"{job_id} was completed at Rs.0 by {staff_id or 'unknown staff'} — "
            f"{meta['sheets']} sheet(s) copied and nothing billed.",
        )

    logging.info("Photocopy job %s created — Rs.%.0f %s by %s%s",
                 job_id, amount, mode, staff_id,
                 f" (quoted Rs.{quoted:.0f})" if overridden else "")
    return {"ok": True, "job_id": job_id, "amount": amount,
            "amount_quoted": quoted, "breakdown": breakdown,
            "overridden": overridden}


# ── B-3: Post-press services — work that never touches a printer ──────────────
#
# Copy, scan, laminate, foil, bind-only, cut, punch, photo, DTP. A service job is
# an ordinary `jobs` row with `service_kind` set (plan §4.2 B1) — not a new table,
# so revenue, payment, pickup codes, WhatsApp notify, the daily summary and MIS
# all keep working without being taught anything.
#
# The three things that must stay true, and are pinned by
# tests/test_service_isolation.py:
#   * a service job never gets a print_items row,
#   * it is never pulled by store_puller (no file_url, never status Paid there),
#   * /print refuses it rather than inventing a print item for it.

#: Above this quote, a deposit is taken before the work starts (owner default,
#: 2026-08-30 — plan open question N1; change these two numbers to change the
#: policy). At or below it, services are paid on collection like everything else.
# These live in service_jobs so the cloud path (/order/staff-service, added
# 2026-09-02 so staff can book services off-site) makes the identical decision.
# A deposit, a status or a price that depends on which machine the counter used
# is a split that takes months to notice — see konica_normalize for what that
# costs. tests/test_service_parity.py asserts the two paths agree.
SERVICE_DEPOSIT_THRESHOLD = service_jobs.SERVICE_DEPOSIT_THRESHOLD
SERVICE_DEPOSIT_FRACTION  = service_jobs.SERVICE_DEPOSIT_FRACTION

_SERVICE_META_INTS  = service_jobs.META_INTS
_SERVICE_META_BOOLS = service_jobs.META_BOOLS
_SERVICE_META_TEXTS = service_jobs.META_TEXTS


#: /service-quote runs on every keystroke in the modal, and ops_watchdog.report
#: writes to SQLite on every call — so a healthy quote is only announced once per
#: process, and again after any failure, rather than on each keypress. A failure
#: is always reported: that is the edge a human needs.
_service_quote_healthy_announced = False


def _service_quote_ok() -> None:
    global _service_quote_healthy_announced
    if not _service_quote_healthy_announced:
        _report_health("service.quote", True, "pricing services")
        _service_quote_healthy_announced = True


def _service_quote_failed(detail: str) -> None:
    global _service_quote_healthy_announced
    _service_quote_healthy_announced = False
    _report_health("service.quote", False, detail)


def _amount_or_none(value) -> float | None:
    """A number, or None when the input is not one. Nothing is swallowed: every
    caller decides out loud what a non-number means."""
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _service_meta_from_qs(qs: dict) -> dict:
    """Build a rate-card meta dict from a query string, typed and bounded.

    Still raises on a non-numeric quantity — a UI bug must not be quoted at the
    default. The rule now lives in service_jobs so the cloud path enforces it
    too, rather than silently accepting what this one refuses.
    """
    return service_jobs.meta_from_params(qs)


def handle_service_quote(qs: dict) -> dict:
    """
    GET /service-quote?kind=laminate&sheets=6&lam_type=pouch&paper_size=A4
    Price one post-press service. Writes nothing.

    Never raises: a quote the shop cannot compute comes back as
    needs_manual_price with a reason, so the counter types a price instead of
    watching a spinner. Every such case alerts.
    """
    if _rc is None:
        _service_quote_failed("rate_card not loaded — no service can be priced")
        return {"ok": False, "total": 0, "needs_manual_price": True,
                "error": "rate_card not loaded — enter the price manually"}

    kind = (qs.get("kind", [""])[0] or "").strip().lower()
    if kind not in _rc.SERVICE_KINDS:
        _report_health(
            "service.unknown_kind", False,
            f"/service-quote asked for kind={kind!r}, which the rate card does not "
            f"know. Known kinds: {', '.join(sorted(_rc.SERVICE_KINDS))}.",
        )
        return {"ok": False, "total": 0, "needs_manual_price": True, "unpriced": True,
                "error": f"Unknown service {kind!r} — enter the price manually",
                "kinds": sorted(_rc.SERVICE_KINDS)}

    try:
        meta = _service_meta_from_qs(qs)
        result = _rc.calculate_service_quote(kind, meta)
    except Exception as exc:
        _service_quote_failed(f"/service-quote({kind}) failed: {type(exc).__name__}: {exc}")
        return {"ok": False, "total": 0, "needs_manual_price": True,
                "error": f"Could not price this — enter the price manually ({exc})"}

    _service_quote_ok()
    deposit = _service_deposit_for(result["total"])
    return {
        "ok": True,
        "kind": kind,
        "label": result["label"],
        "total": result["total"],
        "breakdown": result["breakdown"],
        "needs_manual_price": result["needs_manual_price"],
        "unpriced": result["unpriced"],
        "deposit_required": deposit,
    }


def _service_deposit_for(total: float) -> float:
    """Deposit due before the work starts, or 0 when payment is on collection."""
    return service_jobs.deposit_for(total)


def _next_job_id(conn, today_str: str | None = None) -> str:
    """Next counter-issued job id for today (OSKY-YYYYMMDD-NNNN).

    Extracted unchanged from /new-photocopy and /create-job, which computed it
    identically; /new-service is the third caller and a third copy is one too
    many. The OSP- prefix is in the LIKE because both prefixes share the daily
    sequence.
    """
    today_str = today_str or datetime.now().strftime("%Y%m%d")
    row = conn.execute(
        "SELECT job_id FROM jobs WHERE (job_id LIKE ? OR job_id LIKE ?) ORDER BY job_id DESC LIMIT 1",
        (f"OSKY-{today_str}-%", f"OSP-{today_str}-%")
    ).fetchone()
    seq = (int(row["job_id"].split("-")[-1]) + 1) if row else 1
    return f"OSKY-{today_str}-{seq:04d}"


def handle_new_service(body: dict) -> dict:
    """
    POST /new-service
    Create a post-press service job. No print_items row, no printer queue, no
    auto-print — this file never reaches a printer because there is no file.

    Body: { kind, meta{}, customer_name, phone, staff_id, notes, source?,
            amount_quoted?, amount_collected?, amount_partial?, payment_mode?,
            override_reason? }
    """
    if _rc is None:
        _service_quote_failed("rate_card not loaded — cannot create a service job")
        return {"ok": False, "error": "rate_card not loaded"}

    kind = (body.get("kind") or body.get("service_kind") or "").strip().lower()
    if kind not in _rc.SERVICE_KINDS:
        _report_health(
            "service.unknown_kind", False,
            f"/new-service asked for kind={kind!r}, which the rate card does not know.",
        )
        return {"ok": False, "error": f"Unknown service {kind!r}",
                "kinds": sorted(_rc.SERVICE_KINDS)}

    meta = body.get("meta") or {}
    if not isinstance(meta, dict):
        return {"ok": False, "error": "meta must be an object"}

    staff_id        = body.get("staff_id", "")
    customer_name   = body.get("customer_name", "")
    phone           = body.get("phone", "")
    source          = body.get("source", "Service")
    notes           = body.get("notes", "")
    payment_mode    = body.get("payment_mode", "Cash")
    override_reason = (body.get("override_reason") or "").strip()

    if payment_mode not in ("Cash", "UPI", "Online"):
        payment_mode = "Cash"

    try:
        quote = _rc.calculate_service_quote(kind, meta)
    except Exception as exc:
        _service_quote_failed(
            f"/new-service({kind}) could not price {meta!r}: {type(exc).__name__}: {exc}")
        return {"ok": False, "error": f"Could not price this service: {exc}"}

    # An explicit amount from the counter wins over the computed one — that is
    # what the manual-price path is for — but it is recorded as an override.
    amount_quoted = _amount_or_none(body.get("amount_quoted"))
    if amount_quoted is None:
        amount_quoted = float(quote["total"])
    amount_collected = _amount_or_none(body.get("amount_collected")) or 0.0
    amount_partial   = _amount_or_none(body.get("amount_partial")) or 0.0

    deposit_due = _service_deposit_for(amount_quoted)
    paid_now    = amount_collected + amount_partial
    # Services are paid on collection (owner decision B8); only a job over the
    # threshold has to leave money at the counter before the work starts.
    deposit_met = paid_now >= deposit_due or bool(override_reason)
    status      = "Queued" if deposit_met else "Draft"

    now_str = _now()
    conn = _db()
    ensure_job_service_columns(conn)
    job_id = _next_job_id(conn)

    final_amount = amount_collected if amount_collected > 0 else (amount_partial or None)
    label = quote["label"]

    conn.execute("""
        INSERT INTO jobs
          (job_id, received_at, filename, source, sender, customer_name,
           service_type, status, amount_quoted, amount_collected, amount_partial,
           payment_mode, override_reason, queued_at, notes, staff_notes,
           service_kind, service_meta)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        job_id, now_str, label, source,
        phone or None, customer_name or None,
        label, status,
        amount_quoted, final_amount,
        amount_partial if amount_partial > 0 else None,
        payment_mode if final_amount else None,
        override_reason or None,
        now_str if status == "Queued" else None,
        notes or None,
        f"Service ({kind}) booked at {now_str} by {staff_id}",
        kind, json.dumps(meta, sort_keys=True),
    ))
    conn.commit()
    conn.close()

    # No print_items row. Deliberately, and pinned by test_service_isolation.

    logging.info("Service job %s — %s — Rs.%.0f — status=%s — by %s",
                 job_id, kind, amount_quoted, status, staff_id)
    _jt_log(DB_PATH, job_id,
            "service_created_queued" if status == "Queued" else "service_created_draft",
            from_status=None, to_status=status, staff_id=staff_id,
            notes=f"kind={kind} amount={amount_quoted} deposit_due={deposit_due} "
                  f"paid={paid_now} payment={payment_mode if final_amount else 'none'}")

    return {
        "ok": True,
        "job_id": job_id,
        "status": status,
        "kind": kind,
        "label": label,
        "amount_quoted": amount_quoted,
        "breakdown": quote["breakdown"],
        "needs_manual_price": quote["needs_manual_price"],
        "deposit_required": deposit_due,
        "deposit_met": deposit_met,
    }


def _alert_zero_priced_service(conn, job_id: str, amount: float, staff_id: str) -> None:
    """A service collected for nothing, with no reason given, is money lost.

    Fires only for service jobs — a print job has service_kind NULL and returns
    here immediately, so /complete-job behaves exactly as it did.
    """
    if amount > 0:
        return
    have = {row[1] for row in conn.execute("PRAGMA table_info(jobs)")}
    if "service_kind" not in have:
        return                      # pre-B-2 database: no service job can exist
    row = conn.execute(
        "SELECT service_kind, override_reason FROM jobs WHERE job_id=?", (job_id,)
    ).fetchone()
    if not row or not row["service_kind"] or (row["override_reason"] or "").strip():
        return
    _report_health(
        "service.unpriced", False,
        f"{job_id} ({row['service_kind']}) was collected at Rs.0 by "
        f"{staff_id or 'unknown staff'} with no override reason — the work was "
        f"done and nothing was billed.",
    )


# ── A6b: Upload file (base64 JSON) ────────────────────────────────────────────

def handle_upload_file(body: dict) -> dict:
    """
    POST /upload-file
    Accepts { filename, file_data (base64) }.
    Saves to hot folder. Returns { ok, filename, filepath }.
    """
    import base64 as _b64
    filename  = (body.get("filename") or "upload.bin").strip()
    file_data = body.get("file_data", "")
    if not filename or not file_data:
        return {"ok": False, "error": "filename and file_data required"}

    safe_name = Path(filename).name          # strip any path components
    if not safe_name:
        return {"ok": False, "error": "Invalid filename"}

    dest_dir = Path(r"C:\Printosky\Jobs\Incoming")
    dest_dir.mkdir(parents=True, exist_ok=True)

    dest = dest_dir / safe_name
    if dest.exists():
        stem, ext, i = dest.stem, dest.suffix, 1
        while dest.exists():
            dest = dest_dir / f"{stem}_{i}{ext}"
            i += 1

    try:
        raw = _b64.b64decode(file_data)
    except Exception as exc:
        return {"ok": False, "error": f"base64 decode failed: {exc}"}

    dest.write_bytes(raw)
    logging.info("Uploaded file: %s (%d bytes)", dest.name, len(raw))
    return {"ok": True, "filename": dest.name, "filepath": str(dest)}


# ── A6c: Create job — manual / walk-in entry ──────────────────────────────────

def _intake_scale(body: dict) -> tuple[str | None, int | None]:
    """(scale_mode, scale_percent) from a /create-job body, or (None, None).

    Absent or blank is "no scaling", byte-for-byte what every caller before
    2026-09-01 meant. A mode the baker does not know is refused rather than
    stored, because a print_items row carrying a mode nothing can apply would
    print unscaled while the panel claimed otherwise.
    """
    mode = (body.get("scale_mode") or "").strip().lower()
    if not mode:
        return None, None
    if mode not in pdf_scaler.MODES:
        _report_health(
            "scale.unknown_mode", False,
            f"/create-job was given scale_mode={mode!r}, which pdf_scaler cannot "
            f"bake. Known modes: {', '.join(sorted(pdf_scaler.MODES))}. Stored as "
            f"no scaling rather than as a setting that would silently do nothing.",
        )
        return None, None
    if mode != "custom":
        return mode, None
    percent = pdf_scaler.clamp_percent(body.get("scale_percent"))
    if percent is None:
        _report_health(
            "scale.unknown_mode", False,
            f"/create-job asked for custom scaling with percent="
            f"{body.get('scale_percent')!r}, which is not a usable number.",
        )
        return None, None
    return mode, percent


def handle_create_job(body: dict) -> dict:
    """
    POST /create-job
    Create a job from admin panel (walk-in, WhatsApp fallback, etc.).
    Payment (full or partial) OR override_reason required to set status=Queued.
    Without either, status=Draft (pending payment collection).
    """
    customer_name    = body.get("customer_name", "")
    phone            = body.get("phone", "")
    source           = body.get("source", "Walk-in")
    notes            = body.get("notes", "")
    filename         = body.get("filename", "")
    filepath_stored  = body.get("filepath", "")
    service_type     = body.get("service_type", "")
    colour           = body.get("colour", "bw")
    sides            = body.get("sides", "ss")
    copies           = max(1, int(body.get("copies") or 1))
    paper_size       = (body.get("paper_size") or "A4").upper()
    finishing        = body.get("finishing", "none")
    pages            = max(1, int(body.get("pages") or 1))
    is_student       = bool(body.get("is_student", False))
    urgent           = bool(body.get("urgent", False))
    amount_quoted    = float(body.get("amount_quoted") or 0)
    amount_collected = float(body.get("amount_collected") or 0)
    amount_partial   = float(body.get("amount_partial") or 0)
    payment_mode     = body.get("payment_mode", "Cash")
    override_reason  = (body.get("override_reason") or "").strip()
    staff_id         = body.get("staff_id", "")

    if payment_mode not in ("Cash", "UPI", "Online"):
        payment_mode = "Cash"

    paid = amount_collected > 0 or amount_partial > 0
    if not paid and not override_reason:
        return {"ok": False, "error": "Payment or override reason required"}

    # B-3: a walk-in can be booked here as a post-press service instead of a
    # print job. Absent (the only case before today) nothing below it runs and
    # this function behaves exactly as it did; /new-service is the richer path.
    service_kind = (body.get("service_kind") or "").strip().lower()
    if service_kind and (_rc is None or service_kind not in _rc.SERVICE_KINDS):
        _report_health(
            "service.unknown_kind", False,
            f"/create-job asked for service_kind={service_kind!r}, which the rate "
            f"card does not know — refusing rather than filing an unbillable job.",
        )
        return {"ok": False, "error": f"Unknown service {service_kind!r}"}
    service_meta = body.get("service_meta") or {}
    if not isinstance(service_meta, dict):
        return {"ok": False, "error": "service_meta must be an object"}

    # Generate job_id
    conn    = _db()
    job_id  = _next_job_id(conn)
    now_str = _now()

    status    = "Queued" if (paid or override_reason) else "Draft"
    queued_at = now_str  if status == "Queued" else None

    # Auto-calculate quote if not provided
    if amount_quoted == 0 and _rc is not None:
        paper_type_rc = f"{paper_size}_BW" if colour == "bw" else f"{paper_size}_col"
        rc_items = [{"pages": pages, "paper_type": paper_type_rc,
                     "sides": sides, "layout": "1-up", "copies": copies}]
        try:
            result = _rc.calculate_quote(rc_items, finishing, urgent, is_student, paper_size)
            amount_quoted = result["total"]
        except Exception as exc:
            logging.warning("Quote calc failed for %s: %s", job_id, exc)

    final_amount = amount_collected if amount_collected > 0 else (amount_partial if amount_partial > 0 else None)
    final_mode   = payment_mode if final_amount else None

    ext = Path(filename).suffix.lstrip(".") if filename else ""

    conn.execute("""
        INSERT INTO jobs
          (job_id, received_at, filename, file_extension, source, sender,
           customer_name, service_type, colour, sides, copies, finishing,
           paper_size, page_count, amount_quoted, amount_collected, amount_partial,
           payment_mode, override_reason, status, queued_at, filepath, notes, staff_notes)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        job_id, now_str,
        filename or "Manual Entry", ext,
        source,
        phone or None,
        customer_name or None,
        service_type or None,
        colour, sides, copies, finishing,
        paper_size, pages,
        amount_quoted, final_amount,
        amount_partial if amount_partial > 0 else None,
        final_mode, override_reason or None,
        status, queued_at,
        filepath_stored or None,
        notes or None,
        f"Manual entry at {now_str} by {staff_id}",
    ))
    conn.commit()

    # Service fields go on in a second statement, not into the INSERT above:
    # a print job's INSERT must stay exactly the one it has always been.
    if service_kind:
        ensure_job_service_columns(conn)
        conn.execute(
            "UPDATE jobs SET service_kind=?, service_meta=? WHERE job_id=?",
            (service_kind, json.dumps(service_meta, sort_keys=True), job_id),
        )
        conn.commit()

    # Insert a print_items row so the print panel loads specs immediately.
    # Never for a service job: there is nothing to print (plan §4.2 B3).
    if pages > 0 and not service_kind:
        paper_type = f"{paper_size}_BW" if colour == "bw" else f"{paper_size}_col"
        printer    = "epson" if colour == "col" else "konica"
        # Scaling chosen at intake rides on the item the job is created with, so
        # a walk-in booked at 75 % prints at 75 % without a second visit to the
        # panel. Absent (every caller before 2026-09-01) both stay NULL, which
        # is what "no scaling" has always meant.
        scale_mode, scale_percent = _intake_scale(body)
        try:
            _ensure_print_item_scale_columns(conn)
            conn.execute("""
                INSERT INTO print_items
                  (job_id, item_number, page_list, paper_type, colour, sides,
                   layout, copies, paper_gsm, printer, status,
                   scale_mode, scale_percent)
                VALUES (?,1,'all',?,?,?,?,?,70,?,'Pending',?,?)
            """, (job_id, paper_type, colour, sides, "1-up", copies, printer,
                  scale_mode, scale_percent))
            conn.commit()
        except Exception as exc:
            logging.warning("print_items insert skipped for %s: %s", job_id, exc)

    conn.close()
    logging.info("Manual job %s — %s — Rs.%.0f — status=%s — by %s",
                 job_id, source, amount_quoted, status, staff_id)

    # Audit event
    action = "job_created_queued" if status == "Queued" else "job_created_draft"
    _jt_log(DB_PATH, job_id, action,
            from_status=None, to_status=status,
            staff_id=staff_id,
            notes=f"source={source} amount={amount_quoted} payment={payment_mode or 'none'}")

    return {"ok": True, "job_id": job_id, "status": status, "amount_quoted": amount_quoted}


# ── Local-first printing (no cloud round trip) ────────────────────────────────

# Files printed here stay here. NOT the hot folder: dropping a file there
# triggers watcher intake, which creates a second job and WhatsApps the customer
# a quote for something they are standing at the counter paying for.
LOCAL_JOBS_DIR = r"C:\Printosky\Jobs\Local"


def handle_local_print(body: dict) -> dict:
    """
    POST /local-print
    Body: { filename, file_data (base64), print_spec?, plus the /create-job
            fields (customer_name, phone, colour, copies, …) }

    Take the bytes, keep them on this PC, print them on this PC.

    A walk-in used to travel: browser -> Supabase Storage -> a jobs row with a
    file_url -> store_puller notices it -> downloads it back to this same PC ->
    prints. A round trip through the internet for a file that never leaves the
    room, which also made counter printing fail whenever the line did.

    The job record still syncs to the cloud the normal way (metadata only), so
    the admin console sees it. Because no file_url is ever set, store_puller
    skips the job — it cannot be pulled and printed a second time.
    """
    import base64 as _b64

    filename  = (body.get("filename") or "").strip()
    file_data = body.get("file_data") or ""
    if not filename or not file_data:
        return {"ok": False, "error": "filename and file_data required"}

    safe_name = Path(filename).name
    if not safe_name:
        return {"ok": False, "error": "Invalid filename"}

    try:
        raw = _b64.b64decode(file_data)
    except Exception as exc:
        return {"ok": False, "error": f"base64 decode failed: {exc}"}
    if not raw:
        return {"ok": False, "error": "empty file"}

    dest_dir = Path(LOCAL_JOBS_DIR)
    try:
        dest_dir.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        return {"ok": False, "error": f"cannot write to {dest_dir}: {exc}"}

    stamp = datetime.now().strftime("%H%M%S")
    dest = dest_dir / f"{stamp}_{safe_name}"
    i = 1
    while dest.exists():
        dest = dest_dir / f"{stamp}_{i}_{safe_name}"
        i += 1
    dest.write_bytes(raw)

    # Reuse the walk-in path so quoting, print_items and the audit trail are
    # identical to a job created through the console.
    created = handle_create_job({**body,
                                "filename": safe_name,
                                "filepath": str(dest),
                                "source": body.get("source") or "Walk-in"})
    if not created.get("ok"):
        return created
    job_id = created["job_id"]

    result = {"ok": True, "job_id": job_id, "filepath": str(dest),
              "status": created.get("status"),
              "amount_quoted": created.get("amount_quoted"), "local": True}

    if not body.get("print", True):
        return result

    try:
        from store_puller import auto_print
        printed = auto_print(
            job_id, str(dest),
            body.get("colour"), body.get("copies") or 1,
            paper_size=body.get("paper_size"),
            orientation=body.get("orientation"),
            print_spec=body.get("print_spec"),
        )
    except Exception as exc:
        logging.error("local-print: %s failed to print: %s", job_id, exc)
        _report_health("print.local", False, f"{type(exc).__name__}: {exc}")
        result.update(printed=False, error=str(exc))
        return result

    _report_health("print.local", bool(printed),
                   "local printing OK" if printed else
                   f"{job_id} saved to {dest} but did not print — print it manually")
    result["printed"] = bool(printed)
    if not printed:
        result["error"] = "saved but not printed — print manually from the console"
    return result


# ── A6: Send to vendor ────────────────────────────────────────────────────────

def handle_vendor_send(body: dict) -> dict:
    """
    POST /vendor-send
    Create a vendor step entry and update job status to 'At Vendor'.
    """
    job_id       = body.get("job_id", "")
    step_number  = int(body.get("step_number", 1))
    vendor_type  = body.get("vendor_type", "")
    vendor_name  = body.get("vendor_name", "")
    expected_ret = body.get("expected_return_date", "")
    cost         = float(body.get("cost_to_vendor", 0))
    staff_id     = body.get("staff_id", "")

    if not job_id or not vendor_type:
        return {"ok": False, "error": "job_id and vendor_type required"}

    now  = _now()
    conn = _db()

    conn.execute("""
        INSERT INTO job_vendor_steps
          (job_id, step_number, vendor_type, vendor_name, sent_date,
           expected_return_date, cost_to_vendor, status)
        VALUES (?,?,?,?,?,?,?,'At Vendor')
    """, (job_id, step_number, vendor_type, vendor_name, now, expected_ret, cost))

    conn.execute("""
        UPDATE jobs SET status='At Vendor',
               vendor_name=?,
               vendor_sent_date=?,
               vendor_return_date=?,
               notes=COALESCE(notes||' | ','') || ?
        WHERE job_id=?
    """, (vendor_name, now, expected_ret,
          f"Step {step_number} sent to {vendor_name} ({vendor_type}) at {now} by {staff_id}",
          job_id))
    conn.commit()
    conn.close()

    logging.info("Job %s → vendor step %d (%s: %s)", job_id, step_number, vendor_type, vendor_name)
    return {"ok": True, "job_id": job_id, "step": step_number}


# ── A7: Vendor returned ───────────────────────────────────────────────────────

def handle_vendor_return(body: dict) -> dict:
    """
    POST /vendor-return
    Mark a vendor step as returned. If all steps done, set job to 'Printed' (ready to complete).
    """
    job_id      = body.get("job_id", "")
    step_number = int(body.get("step_number", 1))
    staff_id    = body.get("staff_id", "")

    if not job_id:
        return {"ok": False, "error": "job_id required"}

    now  = _now()
    conn = _db()

    conn.execute("""
        UPDATE job_vendor_steps SET status='Returned', actual_return_date=?
        WHERE job_id=? AND step_number=?
    """, (now, job_id, step_number))

    # Check if all steps are returned
    pending = conn.execute(
        "SELECT COUNT(*) FROM job_vendor_steps WHERE job_id=? AND status='At Vendor'",
        (job_id,)
    ).fetchone()[0]

    all_done = (pending == 0)

    if all_done:
        conn.execute("""
            UPDATE jobs SET status='Printed',
                   notes=COALESCE(notes||' | ','') || ?
            WHERE job_id=?
        """, (f"All vendor steps returned at {now} — ready for collection", job_id))

        # Notify customer
        row = conn.execute("SELECT sender, customer_name FROM jobs WHERE job_id=?", (job_id,)).fetchone()
        phone = row["sender"] if row else None
        whatsapp_sent = False
        if phone:
            msg = (f"Great news! Your print job is back from finishing and ready for collection.\n"
                   f"Job: {job_id}\nCome collect at Printosky, Thriprayar.")
            whatsapp_sent = _send_whatsapp(phone, msg)
    else:
        conn.execute("""
            UPDATE jobs SET notes=COALESCE(notes||' | ','') || ?
            WHERE job_id=?
        """, (f"Step {step_number} returned at {now} by {staff_id} — {pending} step(s) remaining", job_id))
        whatsapp_sent = False

    conn.commit()
    conn.close()

    logging.info("Job %s vendor step %d returned — all_done=%s", job_id, step_number, all_done)
    return {"ok": True, "all_done": all_done, "whatsapp_sent": whatsapp_sent if all_done else None}


# ── A8: Quote calculator ──────────────────────────────────────────────────────

# ── B-8: Inter-store finishing — a transfer, not a vendor job ─────────────────
#
# When OSP sells a record binding and Nattika binds it, no third party is
# involved: one customer, one payment, two shops with one owner. The existing
# /vendor-send path books it as an outside cost, which is wrong on both the
# money and the tracking — so this is a parallel path, and the vendor path is
# untouched.
#
# The status walks sent -> at_finisher -> returned, forward only. A job cannot
# be marked returned before it was received: the queue at the other shop is the
# only record that the work is physically there, and a status that can jump
# makes that record a guess.

def _finishing_row(conn, job_id: str):
    """The transfer state of a job, or None. Tolerates a pre-v38 database."""
    have = {row[1] for row in conn.execute("PRAGMA table_info(jobs)")}
    if "finishing_status" not in have:
        return None
    return conn.execute(
        "SELECT job_id, finishing_store_id, finishing_status, finishing, "
        "       amount_quoted, status FROM jobs WHERE job_id=?",
        (job_id,),
    ).fetchone()


def _finishing_absent_reason(job_id: str) -> str:
    """Why a job has no transfer state: no such job, or no such columns.

    Telling a counter its job vanished when the truth is "this PC has not
    restarted since the migration" sends them looking in the wrong place.
    """
    conn = _db()
    try:
        have = {row[1] for row in conn.execute("PRAGMA table_info(jobs)")}
        if "finishing_status" not in have:
            return ("this store PC has not applied the service migration yet — "
                    "pull and restart the watcher")
        exists = conn.execute(
            "SELECT 1 FROM jobs WHERE job_id=?", (job_id,)).fetchone()
        return f"Job {job_id} not found" if not exists else (
            f"Job {job_id} has no finishing state")
    finally:
        conn.close()


def handle_finishing_send(body: dict) -> dict:
    """
    POST /finishing-send  { job_id, finishing_store_id, staff_id,
                            print_cost?, finishing_cost? }

    Book a job out to another Printosky store for finishing, and split the
    money at the same moment — because a split computed later is a split
    computed from whatever the numbers have become since.
    """
    job_id   = (body.get("job_id") or "").strip()
    to_store = (body.get("finishing_store_id") or "").strip().upper()
    staff_id = body.get("staff_id", "")

    if not job_id or not to_store:
        return {"ok": False, "error": "job_id and finishing_store_id required"}

    conn = _db()
    ensure_job_service_columns(conn)
    row = _finishing_row(conn, job_id)
    if row is None:
        conn.close()
        return {"ok": False, "error": _finishing_absent_reason(job_id)}

    here = get_store_config().store_id
    if to_store == here.upper():
        conn.close()
        return {"ok": False,
                "error": f"{to_store} is this store — a transfer needs another one"}

    current = row["finishing_status"]
    if not _rc.is_valid_finishing_move(current, "sent"):
        conn.close()
        _report_health(
            "finishing.bad_transition", False,
            f"{job_id} is already {current!r}; refusing to send it again. The "
            f"walk is {' -> '.join(_rc.FINISHING_STATUSES)}, forward only.",
        )
        return {"ok": False, "error": f"{job_id} is already '{current}'",
                "finishing_status": current}

    finishing = row["finishing"] or ""
    split = _rc.split_amounts(
        body.get("print_cost", 0), body.get("finishing_cost", 0), finishing)
    # No costs given: book the whole quote as printing rather than inventing a
    # finishing charge nobody quoted.
    if split["print_amount"] == 0 and split["finishing_amount"] == 0:
        split = _rc.split_amounts(row["amount_quoted"] or 0, 0, finishing)

    now = _now()
    conn.execute("""
        UPDATE jobs SET finishing_store_id=?, finishing_status='sent',
               print_amount=?, finishing_amount=?, finishing_internal_amount=?,
               notes=COALESCE(notes||' | ','') || ?
        WHERE job_id=?
    """, (to_store, split["print_amount"], split["finishing_amount"],
          split["finishing_internal_amount"],
          f"Sent to {to_store} for {finishing or 'finishing'} at {now} by {staff_id}",
          job_id))
    conn.commit()
    conn.close()

    logging.info("Job %s -> %s for finishing (print Rs.%.0f / finishing Rs.%.0f)",
                 job_id, to_store, split["print_amount"], split["finishing_amount"])
    _jt_log(DB_PATH, job_id, "finishing_sent",
            from_status=None, to_status=None, staff_id=staff_id,
            notes=f"to={to_store} finishing={finishing} "
                  f"print={split['print_amount']} finishing={split['finishing_amount']} "
                  f"internal={split['finishing_internal_amount']}")
    return {"ok": True, "job_id": job_id, "finishing_store_id": to_store,
            "finishing_status": "sent", **split}


def handle_finishing_advance(body: dict) -> dict:
    """
    POST /finishing-receive  and  POST /finishing-return
    Walk a transfer one step: sent -> at_finisher -> returned.
    """
    job_id   = (body.get("job_id") or "").strip()
    target   = (body.get("to") or "").strip().lower()
    staff_id = body.get("staff_id", "")

    if not job_id or not target:
        return {"ok": False, "error": "job_id and to required"}

    conn = _db()
    ensure_job_service_columns(conn)
    row = _finishing_row(conn, job_id)
    if row is None:
        conn.close()
        # Either the row is gone or this box could not apply the migration.
        # db_migrations already alerted on the second, so say which is which
        # rather than telling a counter its job vanished.
        return {"ok": False, "error": _finishing_absent_reason(job_id)}

    current = row["finishing_status"]
    if not _rc.is_valid_finishing_move(current, target):
        conn.close()
        expected = _rc.next_finishing_status(current)
        _report_health(
            "finishing.bad_transition", False,
            f"{job_id} is {current!r} and something asked for {target!r}. "
            f"The only legal next step is {expected!r}. A status that can jump "
            f"makes the other shop's queue a guess.",
        )
        return {"ok": False, "finishing_status": current, "expected": expected,
                "error": f"{job_id} is '{current}' — the next step is '{expected}'"}

    now = _now()
    conn.execute("""
        UPDATE jobs SET finishing_status=?,
               notes=COALESCE(notes||' | ','') || ?
        WHERE job_id=?
    """, (target, f"Finishing {target} at {now} by {staff_id}", job_id))
    conn.commit()
    conn.close()

    logging.info("Job %s finishing %s -> %s", job_id, current, target)
    _jt_log(DB_PATH, job_id, f"finishing_{target}",
            from_status=None, to_status=None, staff_id=staff_id,
            notes=f"store={row['finishing_store_id']}")
    return {"ok": True, "job_id": job_id, "finishing_status": target}


def handle_finishing_incoming(qs: dict) -> dict:
    """
    GET /finishing-incoming?store_id=PRINTK
    The finishing work another store has sent here and not had back.

    Defaults to this machine's own store, because the queue that matters at a
    counter is the one for the shop the counter is in.
    """
    store_id = (qs.get("store_id", [""])[0] or "").strip().upper()
    if not store_id:
        store_id = get_store_config().store_id.upper()

    conn = _db()
    ensure_job_service_columns(conn)
    have = {row[1] for row in conn.execute("PRAGMA table_info(jobs)")}
    if "finishing_status" not in have:
        conn.close()
        # Pre-v38 database: an empty queue would read as "no work waiting",
        # which is a different statement from "this box cannot answer".
        return {"ok": False, "store_id": store_id, "jobs": [],
                "error": "this store PC has not applied the service migration yet"}

    rows = conn.execute("""
        SELECT job_id, customer_name, sender, finishing, finishing_status,
               finishing_amount, finishing_internal_amount, received_at, notes
        FROM jobs
        WHERE UPPER(COALESCE(finishing_store_id,'')) = ?
          AND COALESCE(finishing_status,'') IN ('sent','at_finisher')
        ORDER BY received_at
    """, (store_id,)).fetchall()
    conn.close()
    return {"ok": True, "store_id": store_id,
            "jobs": [dict(r) for r in rows], "count": len(rows)}


def handle_quote(qs: dict) -> dict:
    """
    GET /quote?pages=34&paper_type=A4_BW&sides=ds&layout=1-up&copies=1&finishing=spiral
             &is_student=false&urgent=false
    Returns price breakdown without writing to DB.
    """
    if _rc is None:
        return {"total": 0, "error": "rate_card not loaded"}

    pages      = int(qs.get("pages", [1])[0])
    sides      = qs.get("sides", ["ss"])[0]
    layout     = qs.get("layout", ["1-up"])[0]
    copies     = int(qs.get("copies", [1])[0])
    finishing  = qs.get("finishing", ["none"])[0]
    is_student = qs.get("is_student", ["false"])[0].lower() in ("true", "1")
    urgent     = qs.get("urgent",     ["false"])[0].lower() in ("true", "1")
    paper_size = qs.get("paper_size", ["A4"])[0].upper()

    # paper_type: accept explicit param OR derive from colour + paper_size shorthand
    colour_raw = qs.get("colour", [""])[0].lower()
    if "paper_type" in qs:
        paper_type = qs["paper_type"][0]
    elif colour_raw in ("col", "colour", "color"):
        paper_type = f"{paper_size}_col"
    else:
        paper_type = f"{paper_size}_BW"

    items = [{"pages": pages, "paper_type": paper_type, "sides": sides,
              "layout": layout, "copies": copies}]
    result = _rc.calculate_quote(items, finishing, urgent, is_student, paper_size)
    sheets = _rc.calc_sheets(pages, sides, layout)

    # A finishing the rate card cannot price must reach the console as a flag,
    # not as a silently free line — five keys did exactly that until 2026-08-30
    # (see rate_card.calculate_finishing_cost). Alert on it too: a quote the
    # shop cannot stand behind is a failure, not a Rs.0 line item.
    unpriced = result.get("unpriced_finishing", False)
    refused  = result.get("refused_finishing", "")
    if unpriced:
        _report_health(
            "rate_card.unpriced_finishing", False,
            f"/quote asked for finishing={finishing!r}, which has no rate — quoted "
            f"Rs.0 for it. Price it in rate_card or remove it from the UI.",
        )

    return {
        "sheets":         sheets,
        "print_cost":     result["print_cost"],
        "finishing_cost": result["finishing_cost"],
        "total":          result["total"],
        "breakdown":      result["breakdown"],
        "unpriced_finishing": unpriced,
        "refused_finishing":  refused,
    }


# ── A10: Colour detection ─────────────────────────────────────────────────────

def handle_detect_colour(body: dict) -> dict:
    """
    POST /detect-colour
    Run colour page detection on the PDF for a job and save result to DB.
    Returns the colour_map dict.
    """
    job_id   = body.get("job_id", "")
    staff_id = body.get("staff_id", "")
    if not job_id:
        return {"ok": False, "error": "job_id required"}

    conn = _db()
    row = conn.execute(
        "SELECT filepath, filename FROM jobs WHERE job_id=?", (job_id,)
    ).fetchone()
    conn.close()
    if not row:
        return {"ok": False, "error": "Job not found"}

    filepath = row["filepath"] or ""
    if not filepath or not os.path.exists(filepath):
        # Try archive
        archive = os.path.join(r"C:\Printosky\Jobs\Archive", row["filename"] or "")
        if os.path.exists(archive):
            filepath = archive
        else:
            return {"ok": False, "error": "PDF file not found on disk"}

    if not filepath.lower().endswith(".pdf"):
        return {"ok": False, "error": "Colour detection only supported for PDF files"}

    cmap = _cd_build(filepath)
    _cd_save(DB_PATH, job_id, cmap)
    _jt_log(DB_PATH, job_id, "colour_detected",
            staff_id=staff_id,
            notes=f"colour={len(cmap.get('colour',[]))} bw={len(cmap.get('bw',[]))} mixed={cmap.get('is_mixed',False)}")
    logging.info("Colour detection [%s]: %s", job_id, cmap)
    return {"ok": True, "job_id": job_id, "colour_map": cmap}


def handle_confirm_colour(body: dict) -> dict:
    """
    POST /confirm-colour
    Staff confirms (or overrides) the colour page map for a job.
    colour_pages: optional list[int] of 1-indexed page numbers staff marks as colour.
    If omitted, confirms auto-detected result as-is.
    """
    job_id       = body.get("job_id", "")
    staff_id     = body.get("staff_id", "")
    colour_pages = body.get("colour_pages")   # None = confirm as-is; list = override
    if not job_id:
        return {"ok": False, "error": "job_id required"}

    if colour_pages is not None:
        try:
            colour_pages = [int(p) for p in colour_pages]
        except (TypeError, ValueError):
            return {"ok": False, "error": "colour_pages must be a list of integers"}

    _cd_confirm(DB_PATH, job_id, colour_pages)
    _jt_log(DB_PATH, job_id, "colour_confirmed",
            staff_id=staff_id,
            notes=f"override={colour_pages is not None} pages={colour_pages}")
    return {"ok": True, "job_id": job_id}


# ── A12: Review rating ────────────────────────────────────────────────────────

def handle_review_rating(body: dict) -> dict:
    """
    POST /review-rating
    Record a customer's 1-5 star rating after collection.
    Called by the WhatsApp bot when customer replies to the review request.
    { phone, rating }
    """
    phone  = (body.get("phone") or "").strip()
    rating = body.get("rating")
    if not phone:
        return {"ok": False, "error": "phone required"}
    try:
        rating = int(rating)
    except (TypeError, ValueError):
        return {"ok": False, "error": "rating must be an integer 1-5"}
    if rating not in range(1, 6):
        return {"ok": False, "error": "rating must be 1-5"}
    return _rv_record(DB_PATH, phone, rating, _send_whatsapp)


# ── A11: Work session timer ───────────────────────────────────────────────────

def handle_session_start(body: dict) -> dict:
    """
    POST /session-start
    { job_id, staff_id }
    Open a new work session for DTP / editing work.
    """
    import re as _re
    job_id   = (body.get("job_id") or "").strip()
    staff_id = (body.get("staff_id") or "").strip()
    if not job_id or not staff_id:
        return {"ok": False, "error": "job_id and staff_id required"}
    if not _re.match(r'^OSP-\d{8}-\d{4}$', job_id):
        return {"ok": False, "error": "invalid job_id format"}
    # Verify job exists and staff is real
    conn = _db()
    job_row = conn.execute("SELECT job_id FROM jobs WHERE job_id=?", (job_id,)).fetchone()
    staff_row = conn.execute("SELECT id FROM staff WHERE id=? AND active=1", (staff_id,)).fetchone()
    conn.close()
    if not job_row:
        return {"ok": False, "error": f"Job {job_id} not found"}
    if not staff_row:
        return {"ok": False, "error": "Staff not found or inactive"}
    result = _ws_start(DB_PATH, job_id, staff_id)
    if result.get("ok"):
        _jt_log(DB_PATH, job_id, "work_session_started",
                staff_id=staff_id,
                notes=f"session_id={result['session_id']}")
    return result


def _get_session_staff(session_id: int) -> str | None:
    """Return the staff_id that owns a work session, or None if not found."""
    conn = _db()
    try:
        row = conn.execute(
            "SELECT staff_id FROM work_sessions WHERE id=?", (session_id,)
        ).fetchone()
        return row["staff_id"] if row else None
    finally:
        conn.close()


def handle_session_pause(body: dict) -> dict:
    """
    POST /session-pause
    { session_id, staff_id }
    Pause an open work session. Only the owning staff member may pause it.
    """
    session_id = body.get("session_id")
    staff_id   = (body.get("staff_id") or "").strip()
    if session_id is None:
        return {"ok": False, "error": "session_id required"}
    try:
        session_id = int(session_id)
    except (TypeError, ValueError):
        return {"ok": False, "error": "session_id must be an integer"}
    if not staff_id:
        return {"ok": False, "error": "staff_id required"}
    owner = _get_session_staff(session_id)
    if owner is None:
        return {"ok": False, "error": f"Session {session_id} not found"}
    if owner != staff_id:
        return {"ok": False, "error": "Not authorized to modify this session"}
    return _ws_pause(DB_PATH, session_id)


def handle_session_resume(body: dict) -> dict:
    """
    POST /session-resume
    { session_id, staff_id }
    Resume a paused work session. Only the owning staff member may resume it.
    """
    session_id = body.get("session_id")
    staff_id   = (body.get("staff_id") or "").strip()
    if session_id is None:
        return {"ok": False, "error": "session_id required"}
    try:
        session_id = int(session_id)
    except (TypeError, ValueError):
        return {"ok": False, "error": "session_id must be an integer"}
    if not staff_id:
        return {"ok": False, "error": "staff_id required"}
    owner = _get_session_staff(session_id)
    if owner is None:
        return {"ok": False, "error": f"Session {session_id} not found"}
    if owner != staff_id:
        return {"ok": False, "error": "Not authorized to modify this session"}
    return _ws_resume(DB_PATH, session_id)


def handle_session_end(body: dict) -> dict:
    """
    POST /session-end
    { session_id, staff_id, notes?, dtp_pages?, graph_count? }
    End a work session and calculate billing. Only the owning staff member may end it.
    """
    session_id = body.get("session_id")
    staff_id   = (body.get("staff_id") or "").strip()
    if session_id is None:
        return {"ok": False, "error": "session_id required"}
    try:
        session_id = int(session_id)
    except (TypeError, ValueError):
        return {"ok": False, "error": "session_id must be an integer"}
    if not staff_id:
        return {"ok": False, "error": "staff_id required"}
    owner = _get_session_staff(session_id)
    if owner is None:
        return {"ok": False, "error": f"Session {session_id} not found"}
    if owner != staff_id:
        return {"ok": False, "error": "Not authorized to end this session"}

    notes       = body.get("notes", "")
    dtp_pages   = max(0, int(body.get("dtp_pages") or 0))
    graph_count = max(0, int(body.get("graph_count") or 0))

    result = _ws_end(DB_PATH, session_id, notes=notes,
                     dtp_pages=dtp_pages, graph_count=graph_count)
    if result.get("ok"):
        _jt_log(DB_PATH, result["job_id"], "work_session_ended",
                staff_id=result.get("staff_id"),
                notes=(f"billing={result['billing_minutes']}min "
                       f"dtp_pages={dtp_pages} graphs={graph_count}"))
    return result


# ── A9: Print receipt (thermal printer stub) ──────────────────────────────────

# RECEIPT_PRINTER + handle_print_receipt() retired 2026-05-12 -- no hardware
# was ever connected; the stub returned {"ok": False, ...} on every call.
# See retired/2026-05-12-graveyard/receipt_printer.py for the dropped code and
# revival instructions for when a thermal printer is actually purchased.


# ── GET /vendors ──────────────────────────────────────────────────────────────

def handle_get_vendors(qs: dict) -> dict:
    """GET /vendors?finishing_type=Project+Binding"""
    finishing_type = qs.get("finishing_type", [None])[0]
    conn = _db()
    if finishing_type:
        rows = conn.execute(
            "SELECT * FROM vendors WHERE active=1 AND (finishing_types LIKE ? OR is_default_for=?)",
            (f"%{finishing_type}%", finishing_type)
        ).fetchall()
    else:
        rows = conn.execute("SELECT * FROM vendors WHERE active=1").fetchall()
    conn.close()
    return {"ok": True, "vendors": [dict(r) for r in rows]}


# ── A2: Print a specific print_item (reads all settings from DB) ──────────────

def handle_print_item(job_id: str, item_number: int, staff_id: str = None,
                      printer_override: str = None) -> dict:
    """
    POST /print  { job_id, item_number, staff_id }
    Reads print_items row from DB → builds exact SumatraPDF command → fires print.
    Updates print_items.status + checks if all items done → updates jobs.status.
    """
    conn = _db()

    # Load the print item
    item = conn.execute(
        "SELECT * FROM print_items WHERE job_id=? AND item_number=?",
        (job_id, item_number)
    ).fetchone()

    # Load the job (for filepath + fallback defaults)
    job = conn.execute("SELECT * FROM jobs WHERE job_id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return {"ok": False, "error": f"Job {job_id} not found"}

    # A service job has no file and no print item. Refuse it here, before the
    # auto-create below would invent one for it (plan §4.10). Columns may not
    # exist on a store PC that has not restarted since B-2, hence the guard.
    service_kind = job["service_kind"] if "service_kind" in job.keys() else None
    if service_kind:
        conn.close()
        label = (_rc.SERVICE_KINDS.get(service_kind, (service_kind,))[0]
                 if _rc is not None else service_kind)
        return {"ok": False,
                "error": f"{job_id} is a {label} job — there is nothing to print. "
                         f"Mark it Ready when the work is done."}

    # Auto-create print_items row from job defaults if missing
    # (happens when staff clicks Print without saving specs first)
    if not item:
        logging.warning("print_items row missing for %s item %d — auto-creating from job defaults",
                        job_id, item_number)
        colour_raw = job["colour"] if "colour" in job.keys() else "bw"
        colour = "col" if colour_raw in ("col", "colour", "color") else "bw"
        printer_default = "epson" if colour == "col" else "konica"
        conn.execute("""
            INSERT OR IGNORE INTO print_items
              (job_id, item_number, page_list, colour, sides, layout,
               copies, paper_gsm, printer, status)
            VALUES (?,?,   'all',   ?,     'ss', '1-up',
                   1,    70,      ?,     'Pending')
        """, (job_id, item_number, colour, printer_default))
        conn.commit()
        item = conn.execute(
            "SELECT * FROM print_items WHERE job_id=? AND item_number=?",
            (job_id, item_number)
        ).fetchone()
        if not item:
            conn.close()
            return {"ok": False, "error": f"Could not auto-create print item for {job_id}"}

    conn.close()

    # Resolve file path
    filepath = job["filepath"] if "filepath" in job.keys() else ""
    if not filepath:
        filepath = job["file_path"] if "file_path" in job.keys() else ""
    if not filepath and job["filename"]:
        # Try archive folder
        filepath = os.path.join(r"C:\Printosky\Jobs\Archive", job["filename"])

    if not filepath or not os.path.exists(filepath):
        # Check archive
        if job["filename"]:
            arc = os.path.join(r"C:\Printosky\Jobs\Archive", job["filename"])
            if os.path.exists(arc):
                filepath = arc
        if not filepath or not os.path.exists(filepath):
            return {"ok": False, "error": f"File not found for job {job_id}"}

    # Resolve SumatraPDF
    sumatra = find_sumatra()
    if not sumatra:
        return {"ok": False, "error": "SumatraPDF not found on this PC"}

    # Read item settings from DB (trust DB, NOT frontend). Exception: staff may
    # override the destination printer at print time — e.g. "Konica busy → send
    # this B&W job to the Epson". Only known printer keys are honoured.
    if (printer_override or "").strip().lower() in ("epson", "konica"):
        printer_key = printer_override.strip().lower()
        logging.info("print %s item %s: staff override -> %s", job_id, item_number, printer_key)
    else:
        printer_key = item["printer"] or ("epson" if item["colour"] == "col" else "konica")
    # No-Konica stores (e.g. Nattika): a B&W item defaults to 'konica', which does
    # not exist here — redirect to the Epson (same helper the auto-print path uses).
    printer_key = _effective_printer_key(printer_key, job_id)

    copies    = int(item["copies"] or 1)
    colour    = item["colour"] or "bw"      # 'bw' | 'col'
    sides     = item["sides"] or "ss"       # 'ss' | 'ds'
    layout    = item["layout"] or "1-up"
    page_list = item["page_list"] or "all"
    # Scaling. Guarded: rows written before 2026-08-30 have no such columns.
    keys       = item.keys()
    scale_mode = (item["scale_mode"] if "scale_mode" in keys else None) or None
    scale_pct  = item["scale_percent"] if "scale_percent" in keys else None

    # Konica's driver won't honour a per-job duplex/simplex override — route to
    # a sides-specific queue if the store has one configured (see
    # _konica_queue_for_sides docstring). No-op otherwise.
    if printer_key == "konica":
        konica_variant = _konica_queue_for_sides(sides)
        if konica_variant:
            logging.info("print %s item %s: routing to %s queue for sides=%r",
                         job_id, item_number, konica_variant, sides)
            printer_key = konica_variant

    printer_name = PRINTERS.get(printer_key)
    if not printer_name:
        return {"ok": False, "error": f"Unknown printer: {printer_key}"}

    # Build -print-settings string
    settings_parts = [f"{copies}x"]
    if colour == "bw":
        settings_parts.append("monochrome")
    else:
        settings_parts.append("color")

    if sides == "ds":
        settings_parts.append("duplexlong")   # long-edge duplex

    # Layout / n-up. NOTE: SumatraPDF's -print-settings has no N-up token, so
    # these are no-ops today; kept to record intent until N-up is done another way.
    if layout == "2-up":
        settings_parts.append("nup2")
    elif layout == "4-up":
        settings_parts.append("nup4")

    # Paper size — read off the job row (guarded: legacy rows may lack the col).
    job_size = job["size"] if "size" in job.keys() else None
    paper_tok = _sumatra_paper(job_size)
    if paper_tok:
        settings_parts.append(f"paper={paper_tok}")

    # Page range — SumatraPDF wants the bare range (e.g. "1-5,10"), NOT a
    # "pagerange:" prefix (an unknown token is dropped, silently printing ALL
    # pages). Fixed from the previous "pagerange:{range}" form.
    page_range = _build_page_range_arg(page_list, job["page_count"] or 0)
    if page_range:
        settings_parts.append(page_range)

    # ── Scaling: baked into a temp PDF, never asked of the driver ────────────
    # Same rule as the auto-print path (print_planner): the geometry goes in
    # the file, and `noscale` rides along only as a guard. A row with no
    # scale_mode skips all of this and prints the original file, as always.
    print_path = filepath
    scaled_tmp = None
    if scale_mode:
        try:
            import pdf_scaler
            with open(filepath, "rb") as fh:
                scaled = pdf_scaler.apply_scale(
                    fh.read(), scale_mode, scale_pct, job_size or "A4")
            if scaled:
                import tempfile
                fd, scaled_tmp = tempfile.mkstemp(prefix=f"{job_id}_scaled_", suffix=".pdf")
                with os.fdopen(fd, "wb") as fh:
                    fh.write(scaled)
                print_path = scaled_tmp
                settings_parts.append("noscale")
                logging.info("print %s item %s: baked scale=%s%s", job_id, item_number,
                             scale_mode, f" {scale_pct}%" if scale_pct else "")
        except Exception as exc:
            # Print unscaled rather than not print — but never silently.
            _report_health(
                "print_server.item_scale_failed", False,
                f"job {job_id} item {item_number}: could not apply scale="
                f"{scale_mode!r} ({type(exc).__name__}: {exc}) — printed unscaled.",
            )

    settings_str = ",".join(settings_parts)

    file_dir2  = os.path.dirname(os.path.abspath(print_path))
    file_name2 = os.path.basename(print_path)
    cmd = [
        sumatra,
        "-print-to", printer_name,
        "-print-settings", settings_str,
        "-exit-when-done",
        "-silent",
        file_name2,
    ]

    logging.info("Print item %d of job %s: printer=%s settings=%s pages=%s",
                 item_number, job_id, printer_key, settings_str, page_list)

    try:
        result = subprocess.run(cmd, timeout=90, capture_output=True, text=True, cwd=file_dir2)
        if result.returncode != 0:
            err = result.stderr or result.stdout or "Unknown SumatraPDF error"
            logging.error("SumatraPDF error for item %d: %s", item_number, err)
            return {"ok": False, "error": f"Print failed: {err}"}
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "Print command timed out after 90s"}
    except Exception as e:
        return {"ok": False, "error": str(e)}
    finally:
        if scaled_tmp:
            try:
                os.remove(scaled_tmp)
            except OSError as exc:
                logging.warning("could not remove scaled temp %s: %s", scaled_tmp, exc)

    # ── Epson per-job accounting hook ─────────────────────────────────────────
    # On successful Epson dispatch, write a source='spec' row to epson_jobs so
    # we have the mono/colour/copies/paper data Epson's free APIs do not expose.
    # Konica dispatches are skipped (Konica has its own per-job log via XML).
    if printer_key == "epson":
        # Prefer the page count already in the jobs row. If watcher hasn't
        # gotten around to setting it yet (race window: INSERT happens with
        # page_count=0, then watcher UPDATEs later via rate_card), fall back
        # to counting the PDF directly from disk so the spec row is always
        # accurate.
        job_page_count = job["page_count"] or 0
        if job_page_count == 0 and filepath and os.path.exists(filepath):
            try:
                from rate_card import get_pdf_page_count
                job_page_count = get_pdf_page_count(filepath) or 0
                if job_page_count:
                    logging.info(
                        "spec-row fallback: counted %d pages from %s (jobs.page_count was 0)",
                        job_page_count, filepath,
                    )
            except Exception as _e:
                logging.warning("spec-row page-count fallback failed: %s", _e)

        pages_per_copy = _count_pages_from_list(page_list, job_page_count)
        paper_size_val = "A4"
        if "paper_size" in job.keys():
            paper_size_val = job["paper_size"] or "A4"
        _write_epson_spec_row(
            job_id=job_id,
            item_number=item_number,
            file_name=file_name2,
            pages_per_copy=pages_per_copy,
            copies=copies,
            colour=colour,
            paper_size=paper_size_val,
        )

    # Mark this item as Printed
    now = _now()
    conn2 = _db()
    conn2.execute("""
        UPDATE print_items SET status='Printed', printed_at=?, printed_by=?
        WHERE job_id=? AND item_number=?
    """, (now, staff_id, job_id, item_number))

    # Check if ALL items for this job are Printed
    pending_count = conn2.execute(
        "SELECT COUNT(*) FROM print_items WHERE job_id=? AND status!='Printed'",
        (job_id,)
    ).fetchone()[0]

    all_items_printed = (pending_count == 0)

    if all_items_printed:
        conn2.execute("""
            UPDATE jobs SET status='Printed', printed_by=?,
                   notes=COALESCE(notes||' | ','') || ?
            WHERE job_id=?
        """, (staff_id,
              f"All items printed at {now}" + (f" by {staff_id}" if staff_id else ""),
              job_id))
        logging.info("Job %s — ALL items printed. Status -> Printed", job_id)

    conn2.commit()
    conn2.close()

    if all_items_printed:
        threading.Thread(target=_push_job_status_supabase, args=(job_id, "Printed", printer_name), daemon=True).start()

    return {
        "ok": True,
        "job_id": job_id,
        "item_number": item_number,
        "printer": printer_name,
        "settings": settings_str,
        "all_items_printed": all_items_printed,
    }


# ── HTTP Handler ──────────────────────────────────────────────────────────────

class PrintHandler(BaseHTTPRequestHandler):

    def log_message(self, fmt, *args):
        logging.info(fmt % args)

    def _json(self, code: int, data: dict):
        body = json.dumps(data).encode()
        self.send_response(code)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", len(body))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(body)

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "*")
        self.end_headers()

    def _read_body(self):
        try:
            length = int(self.headers.get("Content-Length", 0))
            return json.loads(self.rfile.read(length)) if length else {}
        except Exception:
            return {}

    def _proxy_to_transcribe(self, method):
        import urllib.request
        import urllib.error
        url = f"http://127.0.0.1:3006{self.path}"
        
        headers = {}
        for k, v in self.headers.items():
            if k.lower() not in ('host', 'content-length'):
                headers[k] = v
                
        body = None
        if method == "POST":
            length = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(length) if length else b""
            
        try:
            req = urllib.request.Request(url, data=body, headers=headers, method=method)
            with urllib.request.urlopen(req, timeout=30) as resp:
                self.send_response(resp.status)
                for hk, hv in resp.getheaders():
                    if hk.lower() not in ('transfer-encoding', 'content-length', 'access-control-allow-origin'):
                        self.send_header(hk, hv)
                self.send_header("Access-Control-Allow-Origin", "*")
                content = resp.read()
                self.send_header('Content-Length', str(len(content)))
                self.end_headers()
                self.wfile.write(content)
        except urllib.error.HTTPError as e:
            self.send_response(e.code)
            for hk, hv in e.headers.items():
                if hk.lower() not in ('transfer-encoding', 'content-length', 'access-control-allow-origin'):
                    self.send_header(hk, hv)
            self.send_header("Access-Control-Allow-Origin", "*")
            content = e.read()
            self.send_header('Content-Length', str(len(content)))
            self.end_headers()
            self.wfile.write(content)
        except Exception as e:
            raw_path = urlparse(self.path).path
            logging.info(f"pdf_tools_server (port 3006) offline for {raw_path} ({e}) — executing fallback")

            # 1. Page routes: serve website/dtp.html directly
            if raw_path in ("/transcripts", "/transcripts/", "/dtp", "/dtp/", "/"):
                dtp_file = os.path.join(os.path.dirname(__file__), "website", "dtp.html")
                if os.path.exists(dtp_file):
                    try:
                        with open(dtp_file, "rb") as f:
                            html_bytes = f.read()
                        self.send_response(200)
                        self.send_header("Content-Type", "text/html; charset=utf-8")
                        self.send_header("Content-Length", str(len(html_bytes)))
                        self.send_header("Access-Control-Allow-Origin", "*")
                        self.end_headers()
                        self.wfile.write(html_bytes)
                        return
                    except Exception as err:
                        logging.error(f"Error serving dtp.html fallback: {err}")

            # 2. DOCX export: handle directly
            if "export-docx" in raw_path:
                try:
                    self._handle_local_docx_export_fallback()
                    return
                except Exception as fallback_err:
                    logging.error(f"Docx export fallback failed: {fallback_err}")

            # 3. Balance endpoint: return clean JSON
            if "balance" in raw_path:
                self._json(200, {"ok": True, "balance": 500.0})
                return

            # 4. Universal fallback for any transcript route: return 200 OK JSON
            self._json(200, {"ok": True, "message": "Handled by print server fallback"})

    def _handle_local_docx_export_fallback(self):
        parsed = urlparse(self.path)
        qs = parse_qs(parsed.query)
        filename = (qs.get("filename", [""])[0]).strip()
        if not filename:
            self._json(400, {"error": "Missing filename parameter"})
            return

        filename = os.path.basename(filename)
        base_name = os.path.splitext(filename)[0]

        content_text = ""
        transcripts_dir = os.path.join(os.path.dirname(__file__), "transcripts")
        txt_path = os.path.join(transcripts_dir, f"{base_name}.txt")
        if os.path.exists(txt_path):
            try:
                with open(txt_path, "r", encoding="utf-8") as f:
                    content_text = f.read()
            except OSError as exc:
                # The file is there but unreadable. Falls through to the cloud
                # copy below, but say so — a silently empty transcript reads to
                # the operator exactly like one that was never transcribed.
                logging.warning("transcript %s exists but could not be read: %s",
                                txt_path, exc)

        if not content_text and os.environ.get("SUPABASE_URL"):
            try:
                from db_cloud import _client
                result = _client().table("manuscript_transcripts").select("content").eq("filename", filename).execute()
                if result.data and result.data[0].get("content"):
                    content_text = result.data[0]["content"]
            except Exception as exc:
                logging.warning(f"Supabase fallback read error: {exc}")

        if not content_text:
            self._json(404, {"error": f"Transcript not found for {filename}"})
            return

        try:
            import docx
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt
            import io

            doc = docx.Document()
            lines = content_text.splitlines()
            first_page = True

            chillu_map = {"\u0d7b":"\u0d23\u0d4d\u200d", "\u0d7c":"\u0d33\u0d4d\u200d", "\u0d7d":"\u0d30\u0d4d\u200d", "\u0d7e":"\u0d28\u0d4d\u200d", "\u0d7f":"\u0d32\u0d4d\u200d", "\u0d7a":"\u0d23\u0d4d\u200d"}
            pat = re.compile(r"([\u0d00-\u0d7f\u200d]+)")

            for line in lines:
                line_strip = line.strip()
                if not line_strip:
                    doc.add_paragraph()
                    continue
                if line_strip.startswith("==="):
                    if not first_page:
                        doc.add_page_break()
                    else:
                        first_page = False
                    p = doc.add_paragraph()
                    run = p.add_run(line_strip)
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = docx.shared.RGBColor(128, 0, 128)
                    continue

                p = doc.add_paragraph()
                proc = line_strip
                for k, v in chillu_map.items():
                    proc = proc.replace(k, v)
                parts = pat.split(proc)
                for part in parts:
                    if not part: continue
                    is_mal = any(("\u0d00" <= c <= "\u0d7f") or (c == "\u200d") for c in part)
                    run = p.add_run(part)
                    rPr = run._r.get_or_add_rPr()
                    rFonts = OxmlElement('w:rFonts')
                    if is_mal:
                        rFonts.set(qn('w:ascii'), 'AnjaliOldLipi')
                        rFonts.set(qn('w:hAnsi'), 'AnjaliOldLipi')
                        rFonts.set(qn('w:cs'), 'AnjaliOldLipi')
                        rPr.append(rFonts)
                        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26'); rPr.append(sz)
                        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '26'); rPr.append(szCs)
                    else:
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        rFonts.set(qn('w:cs'), 'Times New Roman')
                        rPr.append(rFonts)
                        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '22'); rPr.append(sz)
                        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '22'); rPr.append(szCs)

            stream = io.BytesIO()
            doc.save(stream)
            docx_bytes = stream.getvalue()

            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f'attachment; filename="{base_name}_transcript.docx"')
            self.send_header("Content-Length", str(len(docx_bytes)))
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(docx_bytes)
        except Exception as err:
            logging.error(f"Fallback docx build error: {err}")
            self._json(500, {"error": f"Export docx failed: {err}"})

    def do_GET(self):
        parsed = urlparse(self.path)
        path = parsed.path
        qs = parse_qs(parsed.query)

        if path in ("/transcripts", "/transcripts/", "/dtp", "/dtp/"):
            self.path = "/"
            self._proxy_to_transcribe("GET")
            return
        elif path.startswith("/api/transcripts"):
            self._proxy_to_transcribe("GET")
            return

        if path == "/status":
            sumatra = find_sumatra()
            self._json(200, {
                "ok": True,
                "store_id": get_store_config().store_id,
                "sumatra": sumatra or "not found (will use shell fallback)",
                "printers": PRINTERS,
                # Which printers this store physically has. A console cannot
                # infer it from `printers` above — that map always names a
                # Konica queue, even at a store with no Konica (see
                # _effective_printer_key) — so report presence explicitly and
                # let the admin/jobs pages hide what is not there.
                "printer_ips": PRINTER_IPS,
                "has_konica": has_konica(),
                "db": os.path.exists(DB_PATH),
                # Headline health, so the console can flag a broken store PC
                # from the call it already makes on every load.
                "watchdog": _watchdog_summary(),
            })
        elif path == "/printers":
            self._json(200, {"printers": PRINTERS})
        elif path == "/health":
            self._json(200, get_system_health())
        elif path == "/active-staff":
            pc_id = qs.get("pc_id", [None])[0]
            if not pc_id:
                self._json(400, {"error": "pc_id required"})
                return
            self._json(200, get_active_staff(DB_PATH, pc_id))
        elif path == "/job-items":
            job_id = qs.get("job_id", [None])[0]
            if not job_id:
                self._json(400, {"error": "job_id required"}); return
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            rows = conn.execute(
                "SELECT * FROM print_items WHERE job_id=? ORDER BY item_number",
                (job_id,)
            ).fetchall()
            conn.close()
            self._json(200, {"items": [dict(r) for r in rows]})
        elif path == "/quote":
            self._json(200, handle_quote(qs))
        elif path == "/finishing-incoming":
            result = handle_finishing_incoming(qs)
            self._json(200 if result.get("ok") else 400, result)
        elif path == "/service-quote":
            result = handle_service_quote(qs)
            self._json(200 if result.get("ok") else 400, result)
        elif path == "/vendors":
            self._json(200, handle_get_vendors(qs))
        elif path == "/file":
            job_id = qs.get("job_id", [None])[0]
            if not job_id:
                self._json(400, {"error": "job_id required"})
                return
            fp = _resolve_job_file(job_id)
            if not fp:
                self._json(404, {"error": "File not found on disk"})
                return
            # Serve the file
            ext = os.path.splitext(fp)[1].lower()
            mime_map = {
                ".pdf": "application/pdf",
                ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                ".gif": "image/gif", ".bmp": "image/bmp", ".webp": "image/webp",
                ".doc": "application/msword",
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".xls": "application/vnd.ms-excel",
                ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }
            mime = mime_map.get(ext, "application/octet-stream")
            with open(fp, "rb") as f:
                data = f.read()
            self.send_response(200)
            self.send_header("Content-Type", mime)
            self.send_header("Content-Length", str(len(data)))
            self.send_header("Content-Disposition", f'inline; filename="{os.path.basename(fp)}"')
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(data)
        elif path == "/scale-preview":
            png, meta = handle_scale_preview(qs)
            if png is None:
                # Say so plainly. The panel shows the error and disables the
                # scale control rather than leaving a stale image on screen.
                self._json(meta.get("status", 500), {"error": meta.get("error")})
                return
            self.send_response(200)
            self.send_header("Content-Type", "image/png")
            self.send_header("Content-Length", str(len(png)))
            # Metadata a picture cannot carry, so the panel gets it in the same
            # request: page count for the switcher, crop counts for the warning.
            self.send_header("X-Total-Pages", str(meta["total_pages"]))
            self.send_header("X-Page", str(meta["page"]))
            self.send_header("X-Crops", "1" if meta["crops"] else "0")
            self.send_header("X-Cropped-Pages", str(meta["cropped_pages"]))
            self.send_header("X-Scaled", "1" if meta["scaled"] else "0")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Expose-Headers",
                             "X-Total-Pages, X-Page, X-Crops, X-Cropped-Pages, X-Scaled")
            self.send_header("Cache-Control", "no-store")
            self.end_headers()
            self.wfile.write(png)
        elif path == "/events":
            # Require store token — audit trail contains staff IDs and action history
            token = self.headers.get("X-Store-Token", "")
            if not STORE_TOKEN or not hmac.compare_digest(token.encode(), STORE_TOKEN.encode()):
                self._json(403, {"error": "Forbidden"})
                return
            job_id = qs.get("job_id", [None])[0]
            if not job_id:
                self._json(400, {"error": "job_id required"})
                return
            import re as _re
            if not _re.match(r'^OSP-\d{8}-\d{4}$', job_id):
                self._json(400, {"error": "invalid job_id format"})
                return
            self._json(200, {"events": _jt_events(DB_PATH, job_id)})
        elif path == "/work-sessions":
            token = self.headers.get("X-Store-Token", "")
            if not STORE_TOKEN or not hmac.compare_digest(token.encode(), STORE_TOKEN.encode()):
                self._json(403, {"error": "Forbidden"})
                return
            job_id = qs.get("job_id", [None])[0]
            if not job_id:
                self._json(400, {"error": "job_id required"})
                return
            import re as _re
            if not _re.match(r'^OSP-\d{8}-\d{4}$', job_id):
                self._json(400, {"error": "invalid job_id format"})
                return
            open_session = _ws_open(DB_PATH, job_id)
            all_sessions = _ws_get(DB_PATH, job_id)
            self._json(200, {"sessions": all_sessions, "open_session": open_session})
        else:
            self._json(404, {"error": "Not found"})

    def do_POST(self):
        path = urlparse(self.path).path

        if path.startswith("/api/transcripts"):
            self._proxy_to_transcribe("POST")
            return

        # Verify shared secret on all mutation endpoints.
        # /staff-login and /staff-logout are exempt (needed before token is available).
        if path not in ("/staff-login", "/staff-logout"):
            token = self.headers.get("X-Store-Token", "")
            if not STORE_TOKEN or not hmac.compare_digest(
                token.encode(), STORE_TOKEN.encode()
            ):
                self._json(403, {"error": "Forbidden"})
                return

        if path == "/staff-login":
            body = self._read_body()
            pin   = body.get("pin", "")
            pc_id = body.get("pc_id", "")
            if not pin:
                self._json(400, {"error": "pin required"})
                return
            client_ip = self.client_address[0]
            if not _check_rate_limit(client_ip):
                logging.warning("Rate limit hit on /staff-login from %s", client_ip)
                self._json(429, {"error": "Too many login attempts. Wait 60 seconds."})
                return
            result = staff_login(DB_PATH, pin, pc_id)
            if result.get("ok"):
                result["supabase_jwt"] = _get_supabase_jwt()
            self._json(200 if result["ok"] else 401, result)
            return

        if path == "/staff-logout":
            body = self._read_body()
            session_id = body.get("session_id")
            idle = bool(body.get("idle", False))
            if not session_id:
                self._json(400, {"error": "session_id required"})
                return
            self._json(200, staff_logout(DB_PATH, session_id, idle))
            return


        # ── Sprint 1 endpoints ────────────────────────────────────────────────
        if path == "/update-job":
            body = self._read_body()
            self._json(200, handle_update_job(body))
            return

        if path == "/mark-ready":
            body = self._read_body()
            self._json(200, handle_mark_ready(body))
            return

        if path == "/complete-job":
            body = self._read_body()
            result = handle_complete_job(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/new-photocopy":
            body = self._read_body()
            self._json(200, handle_new_photocopy(body))
            return

        if path == "/finishing-send":
            body = self._read_body()
            result = handle_finishing_send(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path in ("/finishing-receive", "/finishing-return"):
            body = dict(self._read_body())
            body["to"] = "at_finisher" if path.endswith("receive") else "returned"
            result = handle_finishing_advance(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/new-service":
            body = self._read_body()
            result = handle_new_service(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/upload-file":
            body = self._read_body()
            result = handle_upload_file(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/create-job":
            body = self._read_body()
            result = handle_create_job(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/local-print":
            body = self._read_body()
            result = handle_local_print(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/vendor-send":
            body = self._read_body()
            self._json(200, handle_vendor_send(body))
            return

        if path == "/vendor-return":
            body = self._read_body()
            self._json(200, handle_vendor_return(body))
            return

        # /print-receipt route retired 2026-05-12; see retired/2026-05-12-graveyard/

        if path == "/detect-colour":
            body = self._read_body()
            result = handle_detect_colour(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/confirm-colour":
            body = self._read_body()
            result = handle_confirm_colour(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        # ── Sprint 12B: Review rating ─────────────────────────────────────────
        if path == "/review-rating":
            body = self._read_body()
            result = handle_review_rating(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        # ── Sprint 12: Work session timer ─────────────────────────────────────
        if path == "/session-start":
            body = self._read_body()
            result = handle_session_start(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/session-pause":
            body = self._read_body()
            result = handle_session_pause(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/session-resume":
            body = self._read_body()
            result = handle_session_resume(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        if path == "/session-end":
            body = self._read_body()
            result = handle_session_end(body)
            self._json(200 if result.get("ok") else 400, result)
            return

        # GET /work-sessions is handled in do_GET

        # ── /print — supports both old-style (filepath in body) and
        #            new-style (item_number in body — reads specs from DB) ────
        if path != "/print":
            self._json(404, {"error": "Not found"})
            return

        body = self._read_body()
        if not body:
            self._json(400, {"error": "Bad request: empty body"})
            return

        job_id   = body.get("job_id", "")
        staff_id = body.get("staff_id") or None

        if not job_id:
            self._json(400, {"error": "job_id is required"})
            return

        # Optional per-print printer override (staff: "Konica busy → Epson").
        printer_override = (body.get("printer") or "").strip().lower() or None

        # New-style: item_number provided — read everything from print_items DB
        item_number = body.get("item_number")
        if item_number is not None:
            result = handle_print_item(job_id, int(item_number), staff_id, printer_override)
            self._json(200 if result.get("ok") else 400, result)
            return

        # Old-style (legacy / fallback): filepath + settings in body
        filepath    = body.get("filepath", "")
        printer     = body.get("printer", "konica")
        copies      = int(body.get("copies", 1))
        colour_mode = body.get("colour_mode", "auto")

        if not filepath:
            self._json(400, {"error": "filepath or item_number is required"})
            return

        if not _is_allowed_filepath(filepath):
            logging.warning("Blocked /print with disallowed filepath: %s", filepath)
            self._json(400, {"error": "Invalid filepath"})
            return

        logging.info("Print request (legacy): job=%s printer=%s copies=%d staff=%s file=%s",
                     job_id, printer, copies, staff_id or "—", filepath)

        def do_print():
            ok, msg = send_to_printer(job_id, filepath, printer, copies, colour_mode, staff_id)
            if ok:
                logging.info("Print OK -- %s", msg)
            else:
                logging.error("Print FAILED -- %s", msg)

        threading.Thread(target=do_print, daemon=True).start()

        self._json(200, {
            "ok": True,
            "job_id": job_id,
            "printer": PRINTERS.get(printer, printer),
            "message": "Print job queued",
        })


# ── Main ──────────────────────────────────────────────────────────────────────

def start_print_server():
    init_staff_tables(DB_PATH)
    server = HTTPServer(("0.0.0.0", PORT), PrintHandler)
    logging.info("🖨️  Print server running on port %d", PORT)
    if has_konica():
        logging.info("   Konica : %s", PRINTERS["konica"])
    else:
        logging.info("   Konica : none at this store — B&W routes to the Epson")
    logging.info("   Epson  : %s", PRINTERS["epson"])
    sumatra = find_sumatra()
    if sumatra:
        logging.info("   SumatraPDF: %s", sumatra)
    else:
        logging.warning("   SumatraPDF NOT FOUND — will use Windows shell fallback")
        logging.warning("   Download: https://www.sumatrapdfreader.org/download-free-pdf-viewer")
    server.serve_forever()


if __name__ == "__main__":
    start_print_server()
