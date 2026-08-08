# -*- coding: utf-8 -*-
"""
pdf_tools_server.py — Local tools server for staff PDF utilities and Manuscript Transcription.
Runs on store PC at port 3006.

Endpoints:
  GET  /                 — HTML Transcription Dashboard
  GET  /status           — health check
  POST /bw-convert       — black and white PDF conversion
  GET  /api/transcripts  — get status of all files in watch folder
  POST /api/transcripts/set-mode - set mode (urgent vs standard)
  GET  /api/transcripts/view - view transcript text
  GET  /api/transcripts/logs - get live console logs
  POST /api/transcripts/trigger - manually trigger a transcription job
"""

import io
import logging
import os
import sys
import tempfile
import time
import re
import json
import threading
from pathlib import Path
from flask import Flask, jsonify, request, send_file
import fitz  # PyMuPDF
from PIL import Image
from dotenv import load_dotenv
from google import genai
from google.genai.errors import APIError

# Ensure UTF-8 output encoding for console
sys.stdout.reconfigure(encoding="utf-8")

# Pull in the conversion engine
_here = Path(__file__).parent
sys.path.insert(0, str(_here))
from pdf_bw import convert  # noqa: E402

load_dotenv(str(_here.parent / ".env"), override=True)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("pdf-tools")

app = Flask(__name__)

# ---------------------------------------------------------------------------
# Global State for Transcription Watcher
# ---------------------------------------------------------------------------
WATCH_DIR = r"D:\Divya teacher\Preeksha sahayi"
TEMP_DIR = r"d:\PY\printosky\_tmp_watcher"

def find_transcript_file(filename):
    r"""
    Search for a transcript text file by filename.
    Checks in:
    1. Legacy WATCH_DIR
    2. C:\DTP and its subdirectories
    """
    base_name = os.path.splitext(os.path.basename(filename))[0]
    target_name = f"{base_name}_transcript.txt"
    
    # 1. Try legacy WATCH_DIR
    legacy_path = os.path.join(WATCH_DIR, target_name)
    if os.path.exists(legacy_path):
        return legacy_path
        
    # 2. Try C:\DTP subfolders
    dtp_base = r"C:\DTP"
    if os.path.exists(dtp_base):
        for root, dirs, files in os.walk(dtp_base):
            if target_name in files:
                return os.path.join(root, target_name)
                
    return None

def find_pdf_file(filename):
    r"""
    Search for the original PDF file.
    Checks in:
    1. Legacy WATCH_DIR
    2. C:\DTP and its subdirectories
    """
    base_name = os.path.basename(filename)
    
    # 1. Try legacy WATCH_DIR
    legacy_path = os.path.join(WATCH_DIR, base_name)
    if os.path.exists(legacy_path):
        return legacy_path
        
    # 2. Try C:\DTP subfolders (space and case-insensitive check to match legacy behavior)
    dtp_base = r"C:\DTP"
    if os.path.exists(dtp_base):
        for root, dirs, files in os.walk(dtp_base):
            for f in files:
                if f.lower() == base_name.lower() or f.replace(" ", "") == base_name.replace(" ", ""):
                    return os.path.join(root, f)
                    
    return None

GLOSSARY = u'''
Use this vocabulary glossary of common terms in this manuscript to verify spelling and resolve hard-to-read handwriting:
- "ഹൈസ്കൂൾ" (High School) - often written in flowy cursive.
- "മലയാളം" (Malayalam)
- "പരീക്ഷാസഹായി" (Exam Helper)
- "ഖണ്ഡികയും ചോദ്യോത്തരങ്ങളും" (Paragraph and Question-Answers)
- "മൂന്നിലൊന്നായി സംഗ്രഹിക്കുക" (Summarize into one-third)
- "ആസ്വാദനക്കുറിപ്പ്" (Appreciation note)
- "ഉപന്യാസം" (Essay)
- "പ്രസംഗം" (Speech)
- "താരതമ്യക്കുറിപ്പ്" (Comparison note)
- "വിശദീകരണങ്ങൾ" (Explanations)
- "കവിതാശകലങ്ങളും പഴഞ്ചൊല്ലുകളും" (Poetic lines and proverbs)
- "പ്രതീക്ഷിക്കാവുന്ന മാതൃകാ ചോദ്യങ്ങൾ" (Expected model questions)
- "ദിവ്യനാദം" (Divyanadham)
- "തൃശ്ശൂർ" / "TSR" (Thrissur)
- "വിവേകോദയം" / "VBHSS" (Vivekodayam School)
'''

PROMPT = u'''You are an expert Malayalam manuscript transcriber.
Your task is to transcribe the handwritten text in the image line-by-line.

Here is a GLOSSARY of terms known to appear in this manuscript. Use it to verify spelling and resolve hard-to-read handwriting:
"""
%s
"""

Follow these rules:
1. Do not translate the Malayalam text. Transcribe it exactly in Malayalam script.
2. Transcribe any English text in English.
3. Do NOT create a new line for every physical line of text written on the page. Instead, flow the sentences continuously. Only create a new line/paragraph when there is a clear paragraph break (separated by a blank line) or a change in question section. Use full stops (.), question marks (?), exclamation marks (!), and other common elements to determine sentence completion and continue writing on the same line.
4. If a word or character is completely illegible, write '[illegible]' instead of guessing.
5. Output ONLY the transcribed text of the image. Do not include any intro, outro, explanations, or meta-comments.
''' % GLOSSARY

console_logs = []
active_file = None
active_page = 0
active_total_pages = 0

def log_to_dashboard(msg):
    line = f"[{time.strftime('%H:%M:%S')}] {msg}"
    console_logs.append(line)
    if len(console_logs) > 200:
        console_logs.pop(0)
    log.info(msg)

# Load/Save config modes (urgent vs standard)
def get_config_modes():
    config_path = os.path.join(WATCH_DIR, "transcribe_config.json")
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return {}

def save_config_mode(filename, mode):
    config_path = os.path.join(WATCH_DIR, "transcribe_config.json")
    cfg = get_config_modes()
    cfg[filename] = mode
    try:
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2, ensure_ascii=False)
    except Exception as e:
        log.error(f"Failed to save transcribe config: {e}")

def get_actual_costs():
    costs_path = os.path.join(WATCH_DIR, "transcribe_costs.json")
    if os.path.exists(costs_path):
        try:
            with open(costs_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return {}

def add_actual_cost(filename, cost):
    costs_path = os.path.join(WATCH_DIR, "transcribe_costs.json")
    costs = get_actual_costs()
    costs[filename] = costs.get(filename, 0.0) + cost
    try:
        with open(costs_path, "w", encoding="utf-8") as f:
            json.dump(costs, f, indent=2, ensure_ascii=False)
    except Exception as e:
        log.error(f"Failed to save transcribe costs: {e}")

def get_current_balance():
    balance_path = os.path.join(WATCH_DIR, "transcribe_balance.json")
    if os.path.exists(balance_path):
        try:
            with open(balance_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("current_balance", 0.0)
        except:
            pass
    return 0.0

def update_current_balance(balance_value):
    balance_path = os.path.join(WATCH_DIR, "transcribe_balance.json")
    data = {
        "current_balance": balance_value,
        "last_updated": time.time()
    }
    try:
        with open(balance_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        log.error(f"Failed to save transcribe balance: {e}")

def subtract_balance_cost(cost):
    bal = get_current_balance()
    update_current_balance(max(0.0, bal - cost))

# Helpers for page counts
def get_transcribed_page_count(transcript_path):
    if not os.path.exists(transcript_path):
        return 0
    try:
        with open(transcript_path, "r", encoding="utf-8") as f:
            content = f.read()
            markers = re.findall(r"=== PAGE (\d+) ===", content)
            return len(markers)
    except:
        return 0

def get_last_transcribed_page(out_path):
    if not os.path.exists(out_path):
        return -1
    last_page = -1
    with open(out_path, "r", encoding="utf-8") as f:
        content = f.read()
        markers = re.findall(r"=== PAGE (\d+) ===", content)
        if markers:
            # markers are 1-indexed (e.g. === PAGE 1 ===), so 0-indexed index is max() - 1
            last_page = max(int(m) for m in markers) - 1
    return last_page

def is_transcription_complete(pdf_path, transcript_path):
    if not os.path.exists(transcript_path) or os.path.getsize(transcript_path) < 100:
        return False
    try:
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        doc.close()
    except:
        return False
    
    # 1-indexed last marker
    last_marker = f"=== PAGE {total_pages} ==="
    with open(transcript_path, "r", encoding="utf-8") as f:
        content = f.read()
        return last_marker in content

def find_reference_pdf(folder):
    for f in os.listdir(folder):
        if 'By DM' in f and 'പരീക്ഷാ' in f and f.lower().endswith('.pdf'):
            return os.path.join(folder, f)
    return None

def transcribe_pdf_job(pdf_path, client, mode):
    global active_file, active_page, active_total_pages
    pdf_name = os.path.basename(pdf_path)
    base_name = os.path.splitext(pdf_name)[0]
    out_path = os.path.join(WATCH_DIR, f"{base_name}_transcript.txt")
    
    active_file = pdf_name
    
    # Select model based on mode
    # Urgent = gemini-3.5-flash (₹20.00/pg)
    # Standard = gemini-3.1-flash-lite (₹5.00/pg)
    model_name = "models/gemini-3.5-flash" if mode == "urgent" else "models/gemini-3.1-flash-lite"
    
    log_to_dashboard(f"Starting {mode.upper()} transcription for {pdf_name} using {model_name}...")
    
    try:
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        active_total_pages = total_pages
        
        last_done = get_last_transcribed_page(out_path)
        start_page = last_done + 1
        active_page = start_page
        
        file_mode = "a" if start_page > 0 else "w"
        with open(out_path, file_mode, encoding="utf-8") as out_file:
            for idx in range(start_page, total_pages):
                active_page = idx
                log_to_dashboard(f"Transcribing page {idx + 1} of {total_pages}...")
                
                page = doc[idx]
                pix = page.get_pixmap(dpi=150)
                img_path = os.path.join(TEMP_DIR, f"temp_page_{idx}.png")
                pix.save(img_path)
                
                img_target = Image.open(img_path)
                img_target.thumbnail((768, 768))
                
                success = False
                retries = 3
                backoff = 10
                
                while not success and retries > 0:
                    try:
                        response = client.models.generate_content(
                            model=model_name,
                            contents=[img_target, PROMPT]
                        )
                        text = response.text
                        success = True
                    except APIError as e:
                        if e.code == 429:
                            log_to_dashboard(f"Rate limit hit. Retrying in {backoff}s...")
                            time.sleep(backoff)
                            backoff *= 2
                            retries -= 1
                        else:
                            log_to_dashboard(f"API Error: {e}. Retrying in {backoff}s...")
                            time.sleep(backoff)
                            retries -= 1
                    except Exception as e:
                        log_to_dashboard(f"Unexpected error: {e}. Retrying...")
                        time.sleep(backoff)
                        retries -= 1
                
                if not success:
                    log_to_dashboard(f"CRITICAL: Failed to transcribe page {idx + 1}. Marking page as illegible to avoid infinite loop.")
                    text = "[safety blocked / illegible]"
                
                text = text.strip()
                out_file.write(f"\n\n=== PAGE {idx + 1} ===\n\n")
                out_file.write(text)
                out_file.write("\n")
                out_file.flush()
                
                # Add to actual costs
                page_cost = 20.0 if mode == "urgent" else 5.0
                add_actual_cost(pdf_name, page_cost)
                subtract_balance_cost(page_cost)
                
                try:
                    os.remove(img_path)
                except:
                    pass
                
                # Dynamic delay to avoid rate limits
                time.sleep(5)
                
        log_to_dashboard(f"SUCCESS: Finished transcribing {pdf_name}")
    except Exception as e:
        log_to_dashboard(f"Error transcribing PDF {pdf_name}: {e}")
    finally:
        active_file = None
        active_page = 0
        active_total_pages = 0

def watcher_thread_fn():
    log_to_dashboard("Watcher thread initialized and scanning watch directory...")
    os.makedirs(TEMP_DIR, exist_ok=True)
    
    while True:
        try:
            api_key = os.environ.get("GEMINI_API_KEY")
            if not api_key or not api_key.strip():
                time.sleep(10)
                continue
                
            client = genai.Client(api_key=api_key)
            
            ref_pdf = find_reference_pdf(WATCH_DIR)
            if not ref_pdf:
                time.sleep(10)
                continue
                
            cfg = get_config_modes()
            
            # Scan WATCH_DIR for PDFs
            for f in os.listdir(WATCH_DIR):
                if f.lower().endswith(".pdf"):
                    pdf_path = os.path.join(WATCH_DIR, f)
                    base_name = os.path.splitext(f)[0]
                    transcript_path = os.path.join(WATCH_DIR, f"{base_name}_transcript.txt")
                    
                    if not is_transcription_complete(pdf_path, transcript_path):
                        # Get mode: default to 'standard' (Flash-Lite)
                        mode = cfg.get(f, "standard")
                        transcribe_pdf_job(pdf_path, client, mode)
                        break # process one PDF at a time
                        
        except Exception as e:
            log_to_dashboard(f"Watcher background loop error: {e}")
            
        time.sleep(10)

# Launch Watcher Thread
watcher_thread = threading.Thread(target=watcher_thread_fn, daemon=True)
watcher_thread.start()


# ---------------------------------------------------------------------------
# CORS — allow calls from Netlify and localhost (no flask-cors dependency)
# ---------------------------------------------------------------------------
ALLOWED_ORIGINS = {
    "https://printosky.com",
    "https://printosky.netlify.app",
}

@app.after_request
def add_cors(resp):
    origin = request.headers.get("Origin", "")
    if origin in ALLOWED_ORIGINS or origin.startswith("http://localhost"):
        resp.headers["Access-Control-Allow-Origin"] = origin
    resp.headers["Access-Control-Allow-Origin"] = "*" # Allow all local dashboard requests
    resp.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
    return resp

@app.route("/bw-convert", methods=["OPTIONS"])
@app.route("/status", methods=["OPTIONS"])
def preflight():
    return "", 200

# ---------------------------------------------------------------------------
# Dashboard Endpoints
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    dashboard_path = _here / "transcribe_dashboard.html"
    if os.path.exists(dashboard_path):
        with open(dashboard_path, "r", encoding="utf-8") as f:
            return f.read()
    return "Dashboard template not found in tools directory.", 404

@app.route("/status")
def status():
    return jsonify({"ok": True, "service": "pdf-tools", "version": "2.0"})

@app.route("/api/transcripts")
def get_transcripts_status():
    if not os.path.exists(WATCH_DIR):
        return jsonify({"error": f"Watch directory not found: {WATCH_DIR}"}), 400
        
    cfg = get_config_modes()
    actual_costs = get_actual_costs()
    files_status = []
    
    total_docs = 0
    total_pages = 0
    total_transcribed_pages = 0
    total_cost = 0.0
    total_actual_cost = 0.0
    
    for f in os.listdir(WATCH_DIR):
        if f.lower().endswith(".pdf"):
            pdf_path = os.path.join(WATCH_DIR, f)
            base_name = os.path.splitext(f)[0]
            transcript_path = os.path.join(WATCH_DIR, f"{base_name}_transcript.txt")
            
            # Get total pages
            try:
                doc = fitz.open(pdf_path)
                pdf_pages = len(doc)
                doc.close()
            except:
                pdf_pages = 0
                
            transcribed = get_transcribed_page_count(transcript_path)
            mode = cfg.get(f, "standard")
            actual_cost = actual_costs.get(f, 0.0)
            
            # Calculate cost (Standard = ₹5.00/pg, Urgent = ₹20.00/pg)
            rate = 20.0 if mode == "urgent" else 5.0
            est_cost = (pdf_pages - transcribed) * rate # projected cost for remaining pages
            
            # Determine status
            status = "pending"
            if is_transcription_complete(pdf_path, transcript_path):
                status = "completed"
            elif active_file == f:
                status = "transcribing"
            elif transcribed > 0:
                status = "paused"
                
            # Get modification time (transcription activity date or upload date)
            pdf_mtime = os.path.getmtime(pdf_path)
            tx_mtime = os.path.getmtime(transcript_path) if os.path.exists(transcript_path) else 0
            mtime = max(pdf_mtime, tx_mtime)
                
            files_status.append({
                "filename": f,
                "total_pages": pdf_pages,
                "transcribed_pages": transcribed,
                "status": status,
                "mode": mode,
                "est_cost": est_cost,
                "actual_cost": actual_cost,
                "mtime": mtime
            })
            
            total_docs += 1
            total_pages += pdf_pages
            total_transcribed_pages += transcribed
            total_cost += est_cost
            total_actual_cost += actual_cost
            
    # Sort files by transcription/activity date in descending order (most recently active first)
    files_status.sort(key=lambda x: x["mtime"], reverse=True)
    
    return jsonify({
        "total_docs": total_docs,
        "total_pages": total_pages,
        "total_transcribed_pages": total_transcribed_pages,
        "total_cost": total_cost,
        "total_actual_cost": total_actual_cost,
        "current_balance": get_current_balance(),
        "watcher_active": watcher_thread.is_alive(),
        "files": files_status
    })

@app.route("/api/transcripts/set-mode", methods=["POST"])
def set_mode():
    data = request.json
    filename = data.get("filename")
    mode = data.get("mode") # "urgent" or "standard"
    if not filename or mode not in ["urgent", "standard"]:
        return jsonify({"error": "Invalid arguments"}), 400
        
    save_config_mode(filename, mode)
    log_to_dashboard(f"Priority mode for {filename} changed to {mode.upper()}")
    return jsonify({"ok": True})

@app.route("/api/transcripts/view")
def view_transcript():
    filename = request.args.get("filename")
    if not filename:
        return jsonify({"error": "Missing filename"}), 400
        
    transcript_path = find_transcript_file(filename)
    
    if transcript_path and os.path.exists(transcript_path):
        try:
            with open(transcript_path, "r", encoding="utf-8") as f:
                text = f.read()
            return jsonify({"ok": True, "text": text})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500
    return jsonify({"ok": False, "error": "Transcript file not found"}), 404

@app.route("/api/transcripts/pdf")
def get_pdf():
    filename = request.args.get("filename")
    if not filename:
        return "Missing filename", 400
    
    # Clean the filename to prevent directory traversal
    filename = os.path.basename(filename)
    
    target_path = find_pdf_file(filename)
                 
    if target_path and os.path.exists(target_path):
        return send_file(target_path, mimetype="application/pdf")
    return "File not found", 404

@app.route("/api/transcripts/page-image")
def get_page_image():
    filename = request.args.get("filename")
    page_num = request.args.get("page", type=int)
    if not filename or page_num is None:
        return "Missing arguments", 400
        
    filename = os.path.basename(filename)
    
    target_path = find_pdf_file(filename)
                
    if not target_path or not os.path.exists(target_path):
        return "File not found", 404
        
    try:
        doc = fitz.open(target_path)
        if page_num < 0 or page_num >= len(doc):
            doc.close()
            return "Invalid page number", 400
            
        page = doc[page_num]
        pix = page.get_pixmap(dpi=150)
        img_data = pix.tobytes("png")
        doc.close()
        
        return send_file(
            io.BytesIO(img_data),
            mimetype="image/png"
        )
    except Exception as e:
        return f"Error: {e}", 500

CHILLU_MAP = {
    "\u0d7b": "\u0d23\u0d4d\u200d", # ൺ
    "\u0d7c": "\u0d33\u0d4d\u200d", # ൾ
    "\u0d7d": "\u0d30\u0d4d\u200d", # ർ
    "\u0d7e": "\u0d28\u0d4d\u200d", # ൻ
    "\u0d7f": "\u0d32\u0d4d\u200d", # ൽ
    "\u0d7a": "\u0d23\u0d4d\u200d", # ൺ
}

def replace_chillus(text):
    for chillu, replacement in CHILLU_MAP.items():
        text = text.replace(chillu, replacement)
    return text

def split_malayalam_english(text):
    # Regex to find blocks of Malayalam characters (Unicode block 0D00-0D7F and ZWJ 200D)
    pattern = re.compile(r"([\u0d00-\u0d7f\u200d]+)")
    parts = pattern.split(text)
    
    segments = []
    for part in parts:
        if not part:
            continue
        is_mal = any(("\u0d00" <= char <= "\u0d7f") or (char == "\u200d") for char in part)
        segments.append((part, is_mal))
    return segments

@app.route("/api/transcripts/export-docx")
def export_docx():
    filename = request.args.get("filename")
    if not filename:
        return "Missing filename", 400
        
    filename = os.path.basename(filename)
    base_name = os.path.splitext(filename)[0]
    transcript_path = find_transcript_file(filename)
    
    if not transcript_path or not os.path.exists(transcript_path):
        return "Transcript not found", 404
        
    try:
        import docx
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
        
        doc = docx.Document()
        
        # Read lines
        with open(transcript_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        first_page = True
        p = None
        
        for line in lines:
            line_strip = line.strip()
            if not line_strip:
                # Add an empty paragraph to represent spacing
                p = doc.add_paragraph()
                continue
                
            # Page breaks
            if line_strip.startswith("==="):
                if not first_page:
                    doc.add_page_break()
                else:
                    first_page = False
                
                # Add page header
                p = doc.add_paragraph()
                run = p.add_run(line_strip)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = docx.shared.RGBColor(128, 0, 128) # purple
                continue
                
            # Normal text line
            p = doc.add_paragraph()
            processed_line = replace_chillus(line_strip)
            segments = split_malayalam_english(processed_line)
            
            for part_text, is_mal in segments:
                run = p.add_run(part_text)
                rPr = run._r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                
                if is_mal:
                    # Malayalam Unicode
                    rFonts.set(qn('w:ascii'), 'AnjaliOldLipi')
                    rFonts.set(qn('w:hAnsi'), 'AnjaliOldLipi')
                    rFonts.set(qn('w:cs'), 'AnjaliOldLipi')
                    rPr.append(rFonts)
                    
                    sz = OxmlElement('w:sz')
                    sz.set(qn('w:val'), '26') # 13 pt (half-points)
                    rPr.append(sz)
                    
                    szCs = OxmlElement('w:szCs')
                    szCs.set(qn('w:val'), '26') # 13 pt
                    rPr.append(szCs)
                else:
                    # English
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rFonts.set(qn('w:cs'), 'Times New Roman')
                    rPr.append(rFonts)
                    
                    sz = OxmlElement('w:sz')
                    sz.set(qn('w:val'), '22') # 11 pt
                    rPr.append(sz)
                    
                    szCs = OxmlElement('w:szCs')
                    szCs.set(qn('w:val'), '22') # 11 pt
                    rPr.append(szCs)
                
        # Save to memory stream
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        return send_file(
            stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{base_name}_transcript.docx"
        )
    except Exception as e:
        return f"Error: {e}", 500

@app.route("/api/transcripts/logs")
def get_logs():
    cloud_log_path = os.path.join(os.path.dirname(__file__), "cloud_worker.log")
    combined_logs = list(console_logs)
    if os.path.exists(cloud_log_path):
        try:
            with open(cloud_log_path, "r", encoding="utf-8") as f:
                lines = f.readlines()
                last_lines = [l.strip() for l in lines[-100:] if l.strip()]
                combined_logs.extend(last_lines)
        except Exception:
            pass
    return jsonify({"logs": combined_logs})

@app.route("/api/transcripts/trigger", methods=["POST"])
def trigger_transcription():
    data = request.json
    filename = data.get("filename")
    if not filename:
        return jsonify({"error": "Missing filename"}), 400
        
    pdf_path = os.path.join(WATCH_DIR, filename)
    if os.path.exists(pdf_path):
        # We delete the transcript file to force restart it from scratch!
        base_name = os.path.splitext(filename)[0]
        tx_path = os.path.join(WATCH_DIR, f"{base_name}_transcript.txt")
        try:
            if os.path.exists(tx_path):
                os.remove(tx_path)
            log_to_dashboard(f"Forcing transcription restart for: {filename}")
            return jsonify({"ok": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500
            
    return jsonify({"error": "File not found"}), 404

@app.route("/api/transcripts/save", methods=["POST"])
def save_transcript():
    data = request.json
    filename = data.get("filename")
    text = data.get("text")
    if not filename or text is None:
        return jsonify({"error": "Missing arguments"}), 400
        
    base_name = os.path.splitext(filename)[0]
    transcript_path = os.path.join(WATCH_DIR, f"{base_name}_transcript.txt")
    
    try:
        with open(transcript_path, "w", encoding="utf-8") as f:
            f.write(text)
        log_to_dashboard(f"Operator manually updated and saved transcript for: {filename}")
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/transcripts/balance", methods=["GET", "POST"])
def manage_balance():
    if request.method == "POST":
        data = request.json or {}
        try:
            new_balance = float(data.get("balance", 0.0))
            update_current_balance(new_balance)
            log_to_dashboard(f"Operator manually synchronized prepaid balance to: ₹{new_balance:.2f}")
            return jsonify({"ok": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    else:
        try:
            return jsonify({"ok": True, "balance": get_current_balance()})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

@app.route("/api/transcripts/upload", methods=["POST"])
def upload_pdf():
    if "file" not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Only PDF files are allowed"}), 400
        
    filename = os.path.basename(file.filename)
    dest_path = os.path.join(WATCH_DIR, filename)
    try:
        file.save(dest_path)
        log_to_dashboard(f"User uploaded new manuscript: {filename}")
        return jsonify({"ok": True, "filename": filename})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/transcripts/delete", methods=["POST"])
def delete_pdf():
    data = request.json
    filename = data.get("filename")
    if not filename:
        return jsonify({"error": "Missing filename"}), 400
        
    filename = os.path.basename(filename)
    pdf_path = find_pdf_file(filename) or os.path.join(WATCH_DIR, filename)
    transcript_path = find_transcript_file(filename) or os.path.join(WATCH_DIR, f"{base_name}_transcript.txt")
    
    try:
        deleted = False
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
            deleted = True
        if os.path.exists(transcript_path):
            os.remove(transcript_path)
            deleted = True
            
        # Remove from actual costs
        costs = get_actual_costs()
        if filename in costs:
            del costs[filename]
            costs_path = os.path.join(WATCH_DIR, "transcribe_costs.json")
            try:
                with open(costs_path, "w", encoding="utf-8") as f:
                    json.dump(costs, f, indent=2, ensure_ascii=False)
            except:
                pass
                
        if deleted:
            log_to_dashboard(f"Operator deleted manuscript and transcript for: {filename}")
            return jsonify({"ok": True})
        return jsonify({"error": "File not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ---------------------------------------------------------------------------
# Original Black-and-White Convert Endpoint (Preserved)
# ---------------------------------------------------------------------------
@app.route("/bw-convert", methods=["POST"])
def bw_convert():
    if "pdf" not in request.files:
        return jsonify({"error": "No PDF file in request (field name must be 'pdf')"}), 400

    pdf_file = request.files["pdf"]
    fname = pdf_file.filename or "upload.pdf"
    if not fname.lower().endswith(".pdf"):
        return jsonify({"error": "Uploaded file must be a .pdf"}), 400

    try:
        dpi         = max(72,   min(600,  int(request.form.get("dpi",         200))))
        sensitivity = max(0.01, min(0.50, float(request.form.get("sensitivity", 0.10))))
        window      = max(5,    min(200,  int(request.form.get("window",       40))))
        tone_split  = max(0.05, min(0.90, float(request.form.get("tone_split",  0.30))))
        white_point = max(150,  min(255,  int(request.form.get("white_point",  220))))
    except (ValueError, TypeError) as e:
        return jsonify({"error": f"Invalid parameter: {e}"}), 400

    log.info(
        "BW convert  file=%r  dpi=%d  sens=%.2f  win=%d  ts=%.2f  wp=%d",
        fname, dpi, sensitivity, window, tone_split, white_point,
    )

    with tempfile.TemporaryDirectory(prefix="bw_") as tmpdir:
        src  = os.path.join(tmpdir, fname)
        stem = Path(fname).stem
        dst  = os.path.join(tmpdir, f"{stem}_BW.pdf")

        pdf_file.save(src)
        input_size = os.path.getsize(src)

        buf = io.StringIO()
        old_stdout = sys.stdout
        sys.stdout = buf
        try:
            out_path = convert(
                src=src, dst=dst,
                dpi=dpi, sensitivity=sensitivity,
                window=window, tone_split=tone_split,
                white_point=white_point,
            )
        except Exception as exc:
            sys.stdout = old_stdout
            log.exception("Conversion failed")
            return jsonify({"error": str(exc)}), 500
        finally:
            sys.stdout = old_stdout

        conv_log    = buf.getvalue()
        text_pages  = _parse_stat(conv_log, "Text pages")
        image_pages = _parse_stat(conv_log, "Image pages")
        elapsed     = _parse_elapsed(conv_log)
        output_size = os.path.getsize(out_path)

        log.info(
            "Done  text=%s  image=%s  %.0fs  %.1fMB→%.1fMB",
            text_pages, image_pages, elapsed,
            input_size / 1e6, output_size / 1e6,
        )

        with open(out_path, "rb") as f:
            pdf_bytes = f.read()

    out_name = f"{stem}_BW.pdf"
    resp = send_file(
        io.BytesIO(pdf_bytes),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=out_name,
    )
    resp.headers["X-Text-Pages"]  = str(text_pages)
    resp.headers["X-Image-Pages"] = str(image_pages)
    resp.headers["X-Elapsed-Sec"] = str(round(elapsed))
    resp.headers["X-Input-MB"]    = f"{input_size / 1e6:.1f}"
    resp.headers["X-Output-MB"]   = f"{output_size / 1e6:.1f}"
    resp.headers["Access-Control-Expose-Headers"] = (
        "X-Text-Pages, X-Image-Pages, X-Elapsed-Sec, X-Input-MB, X-Output-MB"
    )
    return resp

def _parse_stat(log_text: str, label: str) -> str:
    for line in log_text.splitlines():
        if label in line and ":" in line:
            return line.split(":")[-1].strip().split()[0]
    return "?"

def _parse_elapsed(log_text: str) -> float:
    for line in log_text.splitlines():
        if "Time" in line and ":" in line:
            try:
                return float(line.split(":")[-1].strip().rstrip("s"))
            except ValueError:
                pass
    return 0.0

if __name__ == "__main__":
    port = int(os.environ.get("TOOLS_PORT", 3006))
    log.info("PDF Tools Server  →  http://localhost:%d", port)
    log.info("Endpoints:  GET /  |  GET /status  |  POST /bw-convert")
    app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
