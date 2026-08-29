import os
import re
import threading
import time
import logging
import fitz  # PyMuPDF
from PIL import Image
from dotenv import load_dotenv
from google import genai
from google.genai.errors import APIError

# Initialize logging to console and file
import sys
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

log_file = os.path.join(os.path.dirname(__file__), "cloud_worker.log")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)s  %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(log_file, mode="a", encoding="utf-8")
    ]
)
log = logging.getLogger("cloud_worker")

# Load environment and configurations
import sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
load_dotenv(r"d:\PY\printosky\.env")
from store_config import get_store_config
from ops_watchdog import report as _report_health

try:
    from supabase import create_client
    url = os.environ["SUPABASE_URL"]
    key = os.environ.get("SUPABASE_SERVICE_KEY") or os.environ["SUPABASE_KEY"]
    sb = create_client(url, key)
    log.info("Supabase client initialized.")
except Exception as e:
    log.error(f"Failed to initialize Supabase client: {e}")
    sb = None

# Gemini Client
api_key = os.environ.get("GEMINI_API_KEY")
if api_key:
    client = genai.Client(api_key=api_key)
    log.info("Gemini client initialized.")
else:
    client = None
    log.error("GEMINI_API_KEY not found in environment!")

# Configuration defaults
cfg = get_store_config()

def download_pdf_from_storage(filename):
    """Download PDF from storage, supporting unicode/special characters in URLs."""
    import urllib.parse
    import httpx
    
    encoded_filename = urllib.parse.quote(filename)
    public_url = f"{url}/storage/v1/object/public/manuscripts/{encoded_filename}"
    
    try:
        with httpx.Client(http2=True) as client_http:
            resp = client_http.get(public_url)
            if resp.status_code == 200:
                return resp.content
            else:
                log.warning(f"Public URL download returned status {resp.status_code}, falling back to supabase client")
    except Exception as e:
        log.warning(f"Public URL download failed: {e}, falling back to supabase client")
        
    return sb.storage.from_("manuscripts").download(filename)
MY_STORE_ID = cfg.store_id

# Pickup is event-driven: a Supabase Realtime subscription on
# manuscript_transcripts wakes the loop the instant a job goes pending,
# rather than waiting for the next scheduled cycle. POLL_SECONDS is now a
# fallback safety net only, in case the realtime connection drops — it used
# to be the primary pickup path at 60s, costing ~8.6k round trips a day
# against the Supabase egress quota for a low-volume (tens of jobs) queue.
# Override with TRANSCRIBE_POLL_SECONDS for a tighter fallback. Set
# TRANSCRIBE_REALTIME=0 to fall back to poll-only.
POLL_SECONDS = int(os.environ.get("TRANSCRIBE_POLL_SECONDS", "900"))
REALTIME_ENABLED = os.environ.get("TRANSCRIBE_REALTIME", "1").lower() not in ("0", "false", "no")

# Set by the realtime callback (a different thread); the main loop waits on
# it instead of a flat sleep so a change wakes it immediately.
_wake_event = threading.Event()

WATCH_DIR = r"D:\Divya teacher\Preeksha sahayi"
TEMP_DIR = r"d:\PY\printosky\_tmp_watcher"

if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR)

GLOSSARY = u'''
Use this vocabulary glossary of common terms in this manuscript to verify spelling and resolve hard-to-read handwriting:
- "ഹൈസ്കൂൾ" (High School) - often written in flowy cursive.
- "മലയാളം" (Malayalam)
- "പരീക്ഷാസഹായി" (Exam Helper)
- "ഖണ്ഡികയും ചോദ്യോത്തരങ്ങളും" (Paragraph and Question-Answers)
- "മൂന്നിലൊന്നായി സംഗ്രഹിക്കുക" (Summarize into one-third)
- "ആസ്വാദനക്കുറിപ്പ്" (Appreciation note)
- "ഉപന്യാസം" (Essay)
- "പ്രസംഗം" (Speech)
- "താരതമ്യക്കുറിപ്പ്" (Comparison note)
- "വിശദീകരണങ്ങൾ" (Explanations)
- "കവിതാശകലങ്ങളും പഴഞ്ചൊല്ലുകളും" (Poetic lines and proverbs)
- "പ്രതീക്ഷിക്കാവുന്ന മാതൃകാ ചോദ്യങ്ങൾ" (Expected model questions)
- "ദിവ്യനാദം" (Divyanadham)
- "തൃശ്ശൂർ" / "TSR" (Thrissur)
- "വിവേകോദയം" / "VBHSS" (Vivekodayam School)
'''

PROMPT_TEMPLATE = u'''You are an expert Malayalam manuscript transcriber.
Your task is to transcribe the handwritten text in the image line-by-line.

Here is a GLOSSARY of terms known to appear in this manuscript. Use it to verify spelling and resolve hard-to-read handwriting:
"""
%s
"""

Follow these rules:
1. Do not translate the Malayalam text. Transcribe it exactly in Malayalam script.
2. Transcribe any English text in English.
3. Do NOT create a new line for every physical line of text written on the page. Instead, flow the sentences continuously. Only create a new line/paragraph when there is a clear paragraph break (separated by a blank line) or a change in question section. Use full stops (.), question marks (?), exclamation marks (!), and other common elements to determine sentence completion and continue writing on the same line.
4. If a word or character is completely illegible, write '[illegible]' instead of guessing.
5. Output ONLY the transcribed text of the image. Do not include any intro, outro, explanations, or meta-comments.
''' % GLOSSARY

CHILLU_MAP = {
    "\u0d7b": "\u0d23\u0d4d\u200d", # ൺ
    "\u0d7c": "\u0d33\u0d4d\u200d", # ൾ
    "\u0d7d": "\u0d30\u0d4d\u200d", # ർ
    "\u0d7e": "\u0d28\u0d4d\u200d", # ൻ
    "\u0d7f": "\u0d32\u0d4d\u200d", # ൽ
    "\u0d7a": "\u0d23\u0d4d\u200d", # ൺ
}

def replace_chillus(text):
    for chillu, replacement in CHILLU_MAP.items():
        text = text.replace(chillu, replacement)
    return text

def _fail_job(job_id, filename, reason):
    """Mark a transcription job failed AND tell a human.

    A job that dies in here is invisible otherwise: nothing retries it (the
    poll loop claims `pending` and resumes `transcribing`, never `failed`), and
    the dtp console shows the row as failed only to whoever happens to look.
    That is how two manuscripts sat dead for ten days after the OCR-confidence
    change started writing a `confidence_data` column that did not exist in
    Supabase. docs/FAIL_LOUD.md: an empty table is not an alert.
    """
    log.error(f"Job {job_id} ({filename}) failed: {reason}")
    try:
        sb.table("manuscript_transcripts").update({"status": "failed"}).eq("id", job_id).execute()
    except Exception as exc:
        log.error(f"could not mark job {job_id} as failed: {exc}")
    _report_health("transcription_worker.job", False, f"{filename}: {reason}")


def process_transcription_job(job):
    job_id = job["id"]
    filename = job["filename"]
    mode = job.get("mode", "standard")
    
    log.info(f"Starting job {job_id} for file {filename} in {mode.upper()} mode")
    
    # 1. Download PDF from storage
    pdf_temp_path = os.path.join(TEMP_DIR, filename)
    try:
        log.info(f"Downloading {filename} from Supabase Storage...")
        res = download_pdf_from_storage(filename)
        with open(pdf_temp_path, "wb") as f:
            f.write(res)
        log.info("Download completed.")
    except Exception as e:
        _fail_job(job_id, filename, f"download from storage failed: {e}")
        return

    # 2. Open PDF
    try:
        doc = fitz.open(pdf_temp_path)
        total_pages = len(doc)
    except Exception as e:
        _fail_job(job_id, filename, f"PDF could not be opened: {e}")
        return

    # 3. Transcribe page by page
    # Urgent = gemini-3.5-flash
    # Standard = gemini-3.1-flash-lite
    model_name = "models/gemini-3.5-flash" if mode == "urgent" else "models/gemini-3.1-flash-lite"
    
    # Fetch current state from DB
    current_job = sb.table("manuscript_transcripts").select("content,transcribed_pages").eq("id", job_id).execute().data[0]
    start_page = current_job.get("transcribed_pages", 0)
    current_content = current_job.get("content") or ""
    
    log.info(f"Resuming/starting transcription from page {start_page + 1} of {total_pages}...")
    
    try:
        for idx in range(start_page, total_pages):
            log.info(f"Transcribing page {idx + 1}/{total_pages}...")
            
            # Extract target page image
            page = doc[idx]
            pix = page.get_pixmap(dpi=150)
            target_img_path = os.path.join(TEMP_DIR, f"temp_page_{idx}.png")
            pix.save(target_img_path)
            
            img_target = Image.open(target_img_path)
            img_target.thumbnail((768, 768))
            
            success = False
            retries = 3
            backoff = 10
            text = ""
            
            while not success and retries > 0:
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=[img_target, PROMPT_TEMPLATE]
                    )
                    text = response.text or ""
                    success = True
                except APIError as e:
                    if e.code == 429:
                        log.warning(f"Rate limit hit. Retrying in {backoff}s...")
                        time.sleep(backoff)
                        backoff *= 2
                        retries -= 1
                    else:
                        log.error(f"API Error: {e}")
                        retries -= 1
                except Exception as e:
                    log.error(f"Unexpected error: {e}")
                    retries -= 1
            
            if not success:
                log.error(f"Failed to transcribe page {idx + 1} after retries. Saving placeholder.")
                text = "[safety blocked / illegible]"
            
            page_word_confidence = []

            # Format and save updates back to DB
            text = text.strip()
            page_header = f"\n\n=== PAGE {idx + 1} ===\n\n"
            current_content += page_header + text
            
            # Fetch existing confidence data from DB and append
            existing_confidence = []
            try:
                curr_rec = sb.table("manuscript_transcripts").select("confidence_data").eq("id", job_id).execute().data[0]
                existing_confidence = curr_rec.get("confidence_data") or []
            except Exception as conf_err:
                log.warning(f"could not read existing confidence_data: {conf_err}")
            if page_word_confidence:
                existing_confidence.extend(page_word_confidence)

            # Push live update to Supabase
            sb.table("manuscript_transcripts").update({
                "transcribed_pages": idx + 1,
                "content": current_content,
                "confidence_data": existing_confidence,
                "total_pages": total_pages
            }).eq("id", job_id).execute()
            
            # Cleanup temp page image
            try:
                os.remove(target_img_path)
            except:
                pass
                
            time.sleep(4)  # Anti rate-limit delay
            
        # Complete job
        sb.table("manuscript_transcripts").update({"status": "completed"}).eq("id", job_id).execute()
        log.info(f"Job {job_id} successfully completed.")
        _report_health("transcription_worker.job", True,
                       f"{filename}: {total_pages} page(s) transcribed")

    except Exception as e:
        _fail_job(job_id, filename, f"transcription loop aborted: {e}")
    finally:
        doc.close()
        try:
            os.remove(pdf_temp_path)
        except:
            pass

def split_malayalam_english(text):
    # Regex to find blocks of Malayalam characters (Unicode block 0D00-0D7F and ZWJ 200D)
    pattern = re.compile(r"([\u0d00-\u0d7f\u200d]+)")
    parts = pattern.split(text)
    
    segments = []
    for part in parts:
        if not part:
            continue
        is_mal = any(("\u0d00" <= char <= "\u0d7f") or (char == "\u200d") for char in part)
        segments.append((part, is_mal))
    return segments

def _fetch_transcript_content(job_id):
    """Fetch the ``content`` blob for a single job, on demand.

    Split out of the sync probe so the transcript text — by far the largest
    column in the table — leaves the database only when a file actually has to
    be written. Returns "" when the row is gone or has no content yet.
    """
    res = (
        sb.table("manuscript_transcripts")
        .select("content")
        .eq("id", job_id)
        .limit(1)
        .execute()
    )
    rows = res.data or []
    return (rows[0].get("content") or "") if rows else ""


def _synced_filenames(dtp_root):
    r"""Every filename already sitting in any day folder under ``dtp_root``.

    The sync destination is day-stamped (``C:\DTP\DDMMYY``), so "is this job
    already on disk?" cannot be answered by looking in today's folder alone.
    On the first cycle after midnight every completed job in the store's
    history is missing from the new folder, so the worker re-fetched the
    transcript, rebuilt the .docx and re-downloaded the PDF for all of them —
    every day, growing with the archive. Realtime made it plainly visible: the
    loop now wakes the moment a job is added, so adding one manuscript
    re-downloaded the lot.

    A file that exists in any day's folder has been synced. It does not need to
    be materialised into today's folder as well. Anything genuinely new — or a
    job that completed while this PC was off — is still absent everywhere and
    syncs into today's folder as before.
    """
    found = set()
    try:
        entries = list(os.scandir(dtp_root))
    except FileNotFoundError:
        return found                      # first run on this box
    except OSError as exc:
        log.warning(f"could not scan {dtp_root} for already-synced files: {exc}")
        return found
    for entry in entries:
        try:
            if entry.is_dir():
                found.update(os.listdir(entry.path))
        except OSError as exc:
            log.warning(f"could not list {entry.path}: {exc}")
    return found


def sync_completed_jobs():
    # Sync files that were completed and belong to this store.
    #
    # Egress: this runs on every poll cycle, so it must stay cheap. It selects
    # only the two columns needed to work out what is missing on disk — NOT
    # ``content``, which holds the full OCR transcript. Pulling ``select("*")``
    # here re-downloaded every completed transcript on every cycle (~8.7k
    # requests/day), which is what blew the Supabase egress quota. The content is
    # now fetched per-job, only when a local file is actually missing, so a
    # steady state where everything is already synced costs one small query.
    try:
        res = (
            sb.table("manuscript_transcripts")
            .select("id,filename")
            .eq("status", "completed")
            .eq("uploaded_by_store", MY_STORE_ID)
            .execute()
        )
        jobs = res.data or []

        from datetime import datetime
        today_str = datetime.now().strftime("%d%m%y")
        local_sync_dir = os.path.join(r"C:\DTP", today_str)

        # What is already on disk anywhere under C:\DTP, not just in today's
        # folder — see _synced_filenames. dirname() rather than the literal
        # root so the day folder and the scan root always agree.
        synced = _synced_filenames(os.path.dirname(local_sync_dir))

        for job in jobs:
            filename = job["filename"]
            base_name = os.path.splitext(filename)[0]
            local_txt_path = os.path.join(local_sync_dir, f"{base_name}_transcript.txt")
            local_docx_path = os.path.join(local_sync_dir, f"{base_name}_transcript.docx")
            local_pdf_path = os.path.join(local_sync_dir, filename)

            have_txt = f"{base_name}_transcript.txt" in synced
            have_docx = f"{base_name}_transcript.docx" in synced
            have_pdf = filename in synced

            # Nothing missing for this job → skip without fetching the
            # transcript at all. This is the common case once a job has synced,
            # on that day and on every day after it.
            if have_txt and have_docx and have_pdf:
                continue

            # Something has to be written, so pull the transcript text once and
            # reuse it for both the .txt and the .docx below.
            content = None
            if not (have_txt and have_docx):
                try:
                    content = _fetch_transcript_content(job["id"])
                except Exception as e:
                    log.error(f"Failed to fetch transcript content for {filename}: {e}")
                    continue
                if not content:
                    log.warning(f"No transcript content yet for {filename}; skipping this cycle")
                    continue

            # Ensure folder is created before writing/downloading
            if not os.path.exists(local_sync_dir):
                os.makedirs(local_sync_dir, exist_ok=True)

            # 1. Download/Write the completed transcript locally
            if not have_txt:
                log.info(f"Downloading finished transcript locally: {local_txt_path}")
                with open(local_txt_path, "w", encoding="utf-8") as f:
                    f.write(content)
                    
            # 2. Generate and write the Word Docx locally
            if not have_docx:
                log.info(f"Generating Word document locally: {local_docx_path}")
                try:
                    import docx
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    from docx.shared import Pt
                    
                    doc = docx.Document()
                    lines = content.split("\n")
                    first_page = True
                    p = None
                    
                    for line in lines:
                        line_strip = line.strip()
                        if not line_strip:
                            p = doc.add_paragraph()
                            continue
                            
                        if line_strip.startswith("==="):
                            if not first_page:
                                doc.add_page_break()
                            else:
                                first_page = False
                            p = doc.add_paragraph()
                            run = p.add_run(line_strip)
                            run.bold = True
                            continue
                            
                        p = doc.add_paragraph()
                        processed_line = replace_chillus(line_strip)
                        segments = split_malayalam_english(processed_line)
                        
                        for part_text, is_mal in segments:
                            run = p.add_run(part_text)
                            rPr = run._r.get_or_add_rPr()
                            rFonts = OxmlElement('w:rFonts')
                            
                            if is_mal:
                                rFonts.set(qn('w:ascii'), 'AnjaliOldLipi')
                                rFonts.set(qn('w:hAnsi'), 'AnjaliOldLipi')
                                rFonts.set(qn('w:cs'), 'AnjaliOldLipi')
                                rPr.append(rFonts)
                                
                                sz = OxmlElement('w:sz')
                                sz.set(qn('w:val'), '26') # 13 pt
                                rPr.append(sz)
                                
                                szCs = OxmlElement('w:szCs')
                                szCs.set(qn('w:val'), '26') # 13 pt
                                rPr.append(szCs)
                            else:
                                rFonts.set(qn('w:ascii'), 'Times New Roman')
                                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                                rFonts.set(qn('w:cs'), 'Times New Roman')
                                rPr.append(rFonts)
                                
                                sz = OxmlElement('w:sz')
                                sz.set(qn('w:val'), '22') # 11 pt
                                rPr.append(sz)
                                
                                szCs = OxmlElement('w:szCs')
                                szCs.set(qn('w:val'), '22') # 11 pt
                                rPr.append(szCs)
                                
                    doc.save(local_docx_path)
                    log.info("Word document saved successfully.")
                except Exception as ex:
                    log.error(f"Failed to generate Word document during sync: {ex}")
                    
            # 3. Download the original PDF locally if missing
            if not have_pdf:
                log.info(f"Downloading completed PDF locally: {local_pdf_path}")
                try:
                    pdf_bytes = download_pdf_from_storage(filename)
                    with open(local_pdf_path, "wb") as f:
                        f.write(pdf_bytes)
                except Exception as e:
                    log.error(f"Failed to sync completed PDF: {e}")
    except Exception as e:
        log.error(f"Error syncing completed jobs: {e}")

def _realtime_thread(stop: threading.Event) -> None:
    """Run the Supabase Realtime subscription on its own asyncio loop.

    supabase-py 2.15 exposes Realtime through the **async** client only — the
    sync client's ``channel.subscribe()`` raises ``NotImplementedError`` ("use
    the realtime feature in the async client only"), which is what produced the
    "realtime subscription failed" line at every startup here and left
    transcription_worker.realtime red for nine days while the worker ran on its
    15-minute fallback poll. So the subscription lives on an asyncio event loop
    in a dedicated daemon thread; the main loop stays a plain blocking poll and
    this thread's only job is to set ``_wake_event`` when a
    manuscript_transcripts row changes.

    Mirrors store_puller._realtime_thread and the academic worker's copy rather
    than sharing them, matching the rest of this codebase where each poller
    stays self-contained.

    Best-effort: any failure reports to ops_watchdog and returns, leaving the
    fallback poll as the safety net. A socket that dies later is rebuilt by
    realtime_liveness.hold, which also reports both edges — the client's own
    auto-reconnect does not cover every way a connection can go, and a dead
    socket used to leave this check green.
    """
    import asyncio

    import realtime_liveness

    if not (url and key):
        log.warning("realtime disabled — SUPABASE_URL / key not set")
        _report_health("transcription_worker.realtime", False, "SUPABASE_URL / key not set")
        return

    def _on_change(_payload):
        # Runs on the asyncio thread; threading.Event.set is thread-safe.
        _wake_event.set()

    async def _run() -> None:
        from supabase import create_async_client

        # create_async_client passes `key` to the realtime socket as its token,
        # so the connection authorises as service_role and RLS lets changes
        # through — no separate set_auth needed.
        client_async = await create_async_client(url, key)

        async def _subscribe() -> None:
            channel = client_async.channel("manuscript-transcripts")
            channel.on_postgres_changes(
                "*", callback=_on_change, table="manuscript_transcripts", schema="public",
            )
            await client_async.realtime.connect()
            await channel.subscribe()

        await _subscribe()
        log.info("realtime subscription active for manuscript_transcripts")
        _report_health("transcription_worker.realtime", True, "subscribed")

        # Hold the loop open so the client's background listen + heartbeat
        # tasks keep pumping, and rebuild the socket if it dies — the client's
        # own auto-reconnect does not cover every way it can go (see
        # realtime_liveness). On return, asyncio.run cancels those tasks and
        # closes the socket.
        await realtime_liveness.hold(
            client_async.realtime, stop,
            resubscribe=_subscribe,
            on_status=lambda ok, detail: _report_health("transcription_worker.realtime", ok, detail),
            label="transcription_worker",
        )

    try:
        asyncio.run(_run())
    except Exception as exc:
        log.warning(f"realtime subscription failed ({exc}) — falling back to "
                    f"polling every {POLL_SECONDS}s")
        _report_health("transcription_worker.realtime", False, f"{type(exc).__name__}: {exc}")


def start_realtime() -> threading.Event:
    """Start the Realtime subscription in a background daemon thread. Returns a
    stop Event (callers that run forever can ignore it). Non-blocking — the
    subscription is a latency optimisation over the fallback poll, never the
    source of truth, so it must never delay or block startup.
    """
    stop = threading.Event()
    threading.Thread(
        target=_realtime_thread, args=(stop,),
        name="transcription-worker-realtime", daemon=True,
    ).start()
    return stop


def _check_realtime_delivery(claimed_new_job: bool, woken_by_realtime: bool) -> None:
    """Secondary signal: is the subscription actually delivering, not just
    connected? See store_puller._check_realtime_delivery for the full
    rationale — mirrored here rather than shared, matching the rest of this
    codebase where each poller stays self-contained.

    Only ``claimed_new_job`` (a fresh pending job picked up) counts — the
    crash-recovery "transcribing" resume below is unrelated to realtime
    timing and would otherwise show as a false positive on every restart.

    Deliberately not gated on store hours (docs/FAIL_LOUD.md rejects an hours
    gate). It only ever fires when a job was actually claimed, so it is
    silent by construction outside business activity.
    """
    if not claimed_new_job:
        return
    if woken_by_realtime:
        _report_health("transcription_worker.realtime_delivery", True, "claimed job via realtime wake")
        return
    _report_health(
        "transcription_worker.realtime_delivery", False,
        f"claimed a job via the {POLL_SECONDS}s fallback poll with no prior realtime wake — "
        "subscription is connected but not delivering events (check Realtime is enabled on "
        "the `manuscript_transcripts` table in Supabase)",
    )


def main_loop():
    log.info(f"Starting Printosky Cloud Worker for Store ID: {MY_STORE_ID}")
    if REALTIME_ENABLED:
        start_realtime()

    # None on the first cycle: there is no preceding wait to judge yet.
    woken_by_realtime = None
    while True:
        claimed_new_job = False
        try:
            # 1. Check for pending jobs to execute.
            # Only the three fields process_transcription_job reads are selected;
            # it re-fetches its own resume state, so pulling ``content`` here
            # would ship a partial transcript on every cycle for nothing.
            res = sb.table("manuscript_transcripts").select("id,filename,mode").eq("status", "pending").limit(1).execute()
            if res.data:
                job = res.data[0]
                # Claim the job to avoid race conditions
                claim = sb.table("manuscript_transcripts").update({"status": "transcribing"}).eq("id", job["id"]).eq("status", "pending").execute()
                if claim.data:
                    claimed_new_job = True
                    process_transcription_job(job)

            # 2. Check for transcribing jobs that might have crashed/interrupted
            res = sb.table("manuscript_transcripts").select("id,filename,mode").eq("status", "transcribing").limit(1).execute()
            if res.data:
                job = res.data[0]
                # If we are the worker that crashed, or we want to resume
                # For simplicity, we can resume any job that is currently transcribing
                # Let's claim it by updating the timestamp or just continue processing it
                process_transcription_job(job)

            # 3. Synchronize completed jobs to this store's folder
            sync_completed_jobs()

        except Exception as e:
            log.error(f"Error in main worker loop: {e}")

        if REALTIME_ENABLED and woken_by_realtime is not None:
            _check_realtime_delivery(claimed_new_job, woken_by_realtime)

        # Woken immediately by the realtime callback on a matching row change;
        # otherwise falls through to the next scheduled cycle after POLL_SECONDS.
        woken_by_realtime = _wake_event.wait(POLL_SECONDS)
        _wake_event.clear()

if __name__ == "__main__":
    if sb and client:
        main_loop()
    else:
        log.critical("Missing configurations or credentials. Cloud worker cannot start.")
