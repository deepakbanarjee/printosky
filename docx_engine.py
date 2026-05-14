"""
docx_engine.py — Document generation and formatting engine for Printosky Project Builder.

Handles:
- Loading university formatting configs from university_configs/*.json
- Extracting text from uploaded .docx and .pdf files
- Parsing document structure via Claude Haiku
- Generating formatted .docx output for all three product tiers:
    - generate_free_template()  → blank template with placeholders
    - format_fix()              → re-format pasted or extracted text
    - generate_from_form()      → full report from structured form data
"""

import io
import json
import logging
import os
import re
from pathlib import Path

logger = logging.getLogger(__name__)

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from dotenv import load_dotenv

load_dotenv()

_CONFIG_DIR = Path(__file__).parent / "university_configs"


# ---------------------------------------------------------------------------
# Public exceptions
# ---------------------------------------------------------------------------

class StructureDetectionError(Exception):
    """Raised when the Claude structure parser cannot produce a publishable
    structure. The API layer MUST catch this and decide whether to surface
    an upsell (free-preview phase) or route the order to the operator queue
    (post-payment phase). Never silently swallow — that path produced the
    'FORMATTED DOCUMENT' embarrassment we explicitly killed.
    """

    def __init__(
        self,
        message: str,
        *,
        errors: list[str] | None = None,
        model_used: str = "",
        partial_structure: dict | None = None,
        phase: str = "preview",
    ) -> None:
        super().__init__(message)
        self.errors            = errors or []
        self.model_used        = model_used
        self.partial_structure = partial_structure or {}
        self.phase             = phase  # "preview" | "post_payment"

    def to_dict(self) -> dict:
        """Serializable form for API responses + operator queue records."""
        return {
            "message":           str(self),
            "errors":            list(self.errors),
            "model_used":        self.model_used,
            "partial_structure": self.partial_structure,
            "phase":             self.phase,
        }


# ---------------------------------------------------------------------------
# Config helpers
# ---------------------------------------------------------------------------

def load_university_config(university_id: str) -> dict:
    """Load and return JSON config for a university by ID."""
    config_path = _CONFIG_DIR / f"{university_id}.json"
    if not config_path.exists():
        logger.warning("Unknown university ID %r — falling back to ktu", university_id)
        config_path = _CONFIG_DIR / "ktu.json"
    with open(config_path, "r", encoding="utf-8") as f:
        return json.load(f)


def list_universities() -> list:
    """Return [{id, name, short_name, copies_required, cover_color}] for all configs."""
    result = []
    for config_file in sorted(_CONFIG_DIR.glob("*.json")):
        with open(config_file, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        result.append({
            "id":              cfg["id"],
            "name":            cfg["name"],
            "short_name":      cfg["short_name"],
            "copies_required": cfg.get("copies_required", 3),
            "cover_color":     cfg.get("cover_color", ""),
        })
    return result


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file supplied as bytes."""
    doc = Document(io.BytesIO(file_bytes))
    parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return "\n\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF supplied as bytes.

    Returns empty string if the PDF cannot be parsed (corrupt, encrypted, etc.)
    rather than raising — callers should handle empty text gracefully.
    """
    import pdfplumber
    try:
        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text.strip())
        return "\n\n".join(parts)
    except Exception as exc:
        logger.warning("extract_text_from_pdf failed: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# PDF block-based extraction (positional, no LLM)
# ---------------------------------------------------------------------------

_PDF_BULLET_RE   = re.compile(r"^[●•▪◆►]\s*")
_PDF_NUM_LIST_RE = re.compile(r"^(\d{1,2})[\.)]\s+")
_PDF_PAGE_NUM_RE      = re.compile(r"^\(\d+\)$")
_PDF_BODY_PAGENUM_RE  = re.compile(r"^\d{1,3}$")          # bare "1", "12"
_PDF_ALL_CAPS_RE      = re.compile(r"^[A-Z][A-Z \-,&\.\?:!\d]{2,}$")
# Numbered sub-headings like "1.1 INTRODUCTION", "2.3.1 RESEARCH GAP"
_PDF_NUM_HEADING_RE   = re.compile(
    r"^\d+(?:\.\d+){0,2}\s+[A-Z][A-Z \-,&\.\?:!&/]{2,}$"
)
_PDF_KV_RE       = re.compile(r"^([A-Z][^:\n]{2,40}?)\s*:\s*(.+)$")
_PDF_NOISE_TOKENS = ("Map Camera", "Plus Code", " Lat ", " Long ", "GPS")


def _pdf_is_noise(line: str) -> bool:
    return any(tok in line for tok in _PDF_NOISE_TOKENS)


def _pdf_merge_sentences(lines: list[str]) -> list[str]:
    """Join Canva-broken visual lines back into full sentences."""
    out: list[str] = []
    for ln in lines:
        s = ln.rstrip()
        if not s:
            continue
        looks_independent = (
            _PDF_BULLET_RE.match(s)
            or _PDF_NUM_LIST_RE.match(s)
            or (_PDF_ALL_CAPS_RE.match(s) and len(s) < 60)
            or _PDF_PAGE_NUM_RE.match(s)
            or _PDF_KV_RE.match(s)
        )
        if (out
                and not out[-1].rstrip().endswith((".", "?", "!", ":", ";"))
                and not looks_independent
                and not (_PDF_ALL_CAPS_RE.match(out[-1]) and len(out[-1]) < 60)):
            out[-1] = out[-1].rstrip() + " " + s.lstrip()
        else:
            out.append(s)
    return out


def _pdf_block_alignment(x0: float, x1: float, page_width: float) -> str:
    """Classify a block's horizontal alignment from its bbox.

    'center' if the block midpoint is within 8% of page midpoint.
    'right'  if the block left edge is past 35% and right edge past 75%.
    'left'   otherwise (typical body text).
    """
    if page_width <= 0:
        return "left"
    block_mid = (x0 + x1) / 2
    page_mid  = page_width / 2
    if abs(block_mid - page_mid) < page_width * 0.08:
        return "center"
    if x0 > page_width * 0.35 and x1 > page_width * 0.75:
        return "right"
    return "left"


def _pdf_blocks_with_font(pdf, page_no_0based: int) -> list[tuple]:
    """Extract blocks with font metadata.

    Returns 5-tuples: (text, max_size, dominant_size, bold, alignment),
    sorted by Y then X. Text per block is reflowed (Canva line breaks merged).
    """
    page = pdf[page_no_0based]
    page_width = page.rect.width
    d = page.get_text("dict")
    items: list[tuple] = []
    for block in d.get("blocks", []):
        if "lines" not in block:
            continue
        x0, y0, x1, y1 = block["bbox"]
        align = _pdf_block_alignment(x0, x1, page_width)
        max_size = 0.0
        size_chars: dict[float, int] = {}
        any_bold = False
        line_strs: list[str] = []
        for line in block["lines"]:
            line_text = ""
            for span in line["spans"]:
                t = span["text"]
                if not t:
                    continue
                sz = round(span.get("size", 12.0), 1)
                line_text += t
                max_size = max(max_size, sz)
                size_chars[sz] = size_chars.get(sz, 0) + len(t)
                if span.get("flags", 0) & 16:  # bold flag
                    any_bold = True
            if line_text.strip():
                line_strs.append(line_text.strip())
        if not line_strs:
            continue
        dominant = (max(size_chars.items(), key=lambda kv: kv[1])[0]
                    if size_chars else 12.0)
        # Reflow lines within block (Canva visual-line breaks)
        merged: list[str] = []
        buf = ""
        for ln in line_strs:
            is_break = (
                _PDF_BULLET_RE.match(ln)
                or _PDF_NUM_LIST_RE.match(ln)
                or (_PDF_ALL_CAPS_RE.match(ln) and len(ln) < 60)
                or _PDF_PAGE_NUM_RE.match(ln)
            )
            if is_break:
                if buf:
                    merged.append(buf)
                    buf = ""
                merged.append(ln)
            else:
                buf = (buf + " " + ln).strip() if buf else ln
        if buf:
            merged.append(buf)
        for m in merged:
            if not _pdf_is_noise(m):
                items.append((y0, x0, m, max_size, dominant, any_bold, align))
    items.sort(key=lambda t: (round(t[0], 1), round(t[1], 1)))
    return [(text, max_sz, dom_sz, bold, align)
            for _y, _x, text, max_sz, dom_sz, bold, align in items]


def _pdf_estimate_body_size(pdf) -> float:
    """Median font size across the doc, used as the 'body' baseline."""
    sizes: list[float] = []
    for page_no in range(len(pdf)):
        blocks = _pdf_blocks_with_font(pdf, page_no)
        for _text, _maxsz, dom_sz, _bold, _align in blocks:
            sizes.append(dom_sz)
    if not sizes:
        return 12.0
    sizes.sort()
    return sizes[len(sizes) // 2]


def _pdf_page_blocks(pdf, page_no_0based: int) -> list[str]:
    """Extract text blocks from a PDF page in true visual reading order."""
    page = pdf[page_no_0based]
    raw_blocks = [b for b in page.get_text("blocks") if b[6] == 0]
    raw_blocks.sort(key=lambda b: (round(b[1], 1), round(b[0], 1)))
    out: list[str] = []
    for b in raw_blocks:
        text = b[4].strip()
        if not text:
            continue
        block_lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        merged: list[str] = []
        buf = ""
        for ln in block_lines:
            is_break = (
                _PDF_BULLET_RE.match(ln)
                or _PDF_NUM_LIST_RE.match(ln)
                or (_PDF_ALL_CAPS_RE.match(ln) and len(ln) < 60)
                or _PDF_PAGE_NUM_RE.match(ln)
            )
            if is_break:
                if buf:
                    merged.append(buf)
                    buf = ""
                merged.append(ln)
            else:
                buf = (buf + " " + ln).strip() if buf else ln
        if buf:
            merged.append(buf)
        for m in merged:
            if not _pdf_is_noise(m):
                out.append(m)
    return out


def _pdf_pixmap_avg_luma(pix) -> int:
    """Mean luminance (0-255) of a Pixmap. Used to drop pure-black/white tiles."""
    try:
        # Sample up to 20k bytes to avoid scanning huge pixmaps
        sample = pix.samples[:20_000]
        if not sample:
            return 128
        return sum(sample) // len(sample)
    except Exception:
        return 128


def _pdf_is_decorative_pix(pix, byte_size: int) -> bool:
    """Universal decoration heuristics, applied per-image regardless of hash."""
    # Tiny icons / chart axis bits
    if byte_size < 300:
        return True
    if pix.width < 50 or pix.height < 50:
        return True
    luma = _pdf_pixmap_avg_luma(pix)
    # Pure black or pure white = page background / blank tile
    if luma < 8 or luma > 248:
        return True
    return False


def _pdf_build_decorative_xrefs(pdf) -> set[int]:
    """xrefs to drop: hash-on->=3-pages OR per-image decoration heuristics."""
    import hashlib
    import fitz
    hash_pages: dict[str, set[int]] = {}
    xref_hash: dict[int, str] = {}
    universally_decorative: set[int] = set()
    for page_no in range(len(pdf)):
        page = pdf[page_no]
        for info in page.get_images(full=True):
            xref = info[0]
            if xref in xref_hash:
                hash_pages.setdefault(xref_hash[xref], set()).add(page_no)
                continue
            try:
                pix = fitz.Pixmap(pdf, xref)
                if pix.alpha or pix.n > 4:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                png_bytes = pix.tobytes("png")
                if _pdf_is_decorative_pix(pix, len(png_bytes)):
                    universally_decorative.add(xref)
                h = hashlib.md5(png_bytes).hexdigest()
                pix = None
                xref_hash[xref] = h
                hash_pages.setdefault(h, set()).add(page_no)
            except Exception:
                continue
    deco_hashes = {h for h, pgs in hash_pages.items() if len(pgs) >= 3}
    duplicate_decorative = {xref for xref, h in xref_hash.items()
                              if h in deco_hashes}
    return universally_decorative | duplicate_decorative


def extract_text_from_pdf_blocks(file_bytes: bytes) -> str:
    """PDF text extraction with positional reading order + sentence merging."""
    import fitz
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            page_texts: list[str] = []
            for page_no in range(len(pdf)):
                lines = _pdf_page_blocks(pdf, page_no)
                lines = _pdf_merge_sentences(lines)
                if lines:
                    page_texts.append("\n".join(lines))
        return "\n\n".join(page_texts)
    except Exception as exc:
        logger.warning("extract_text_from_pdf_blocks failed: %s", exc)
        return ""


_PDF_KNOWN_SECTION_NAMES: set[str] = {
    "ACKNOWLEDGEMENT", "ACKNOWLEDGEMENTS",
    "DECLARATION", "ABSTRACT",
    "TABLE OF CONTENTS", "CONTENTS",
    "LIST OF FIGURES", "LIST OF TABLES", "LIST OF ABBREVIATIONS",
    "INTRODUCTION", "REVIEW OF LITERATURE",
    "METHODOLOGY", "METHODS",
    "RESULTS", "DISCUSSION", "ANALYSIS",
    "CONCLUSION", "CONCLUSIONS",
    "BIBLIOGRAPHY", "REFERENCES", "APPENDIX", "APPENDICES",
}


def _pdf_classify_block(text: str, max_size: float, body_size: float,
                          page_no: int, is_toc_page: bool = False) -> str:
    """Return 'h1' | 'h2' | 'body' | 'skip' for a PDF block."""
    s = text.strip()
    if not s or _PDF_PAGE_NUM_RE.match(s):
        return "skip"
    # Bare body-page numbers ("1", "12", "76") on their own — Word PDFs
    # often emit these as small isolated blocks at top/bottom of the page.
    # Drop only when the block is at body font size or smaller (so we don't
    # accidentally drop "1.1" headings or numerals used in real content).
    if _PDF_BODY_PAGENUM_RE.match(s) and max_size <= body_size + 1:
        return "skip"
    if page_no == 0 or is_toc_page:
        return "body"
    # h1 requires either an explicit "Chapter N" pattern or a large font.
    # Reject lines that are pure digits/symbols, or that are mostly
    # lowercase (sentence text rendered at heading size, e.g. captions).
    is_alpha_heading = (
        s[:1].isupper()
        and sum(1 for c in s if c.isalpha() and c.isupper())
            >= sum(1 for c in s if c.isalpha() and c.islower())
    )
    # Form labels like "Topic: ..." / "Venue: ..." / "Date: ..." are body
    # even when rendered at heading-sized font in Canva forms.
    looks_like_form_label = bool(re.match(
        r"^(Topic|Venue|Duration|Date|Place|Subject|Title)\b\s*[:.]",
        s, re.IGNORECASE,
    ))
    if re.match(r"^CHAPTER\s+\d+", s, re.IGNORECASE):
        return "h1"
    if (max_size > body_size + 8 and is_alpha_heading
            and 4 <= len(s) <= 80
            and not looks_like_form_label):
        return "h1"
    if (_PDF_ALL_CAPS_RE.match(s) and 4 <= len(s) <= 60
            and max_size >= body_size + 2):
        return "h2"
    # Numbered sub-headings: "1.1 INTRODUCTION", "2.3.1 RESEARCH GAP"
    if (_PDF_NUM_HEADING_RE.match(s)
            and max_size >= body_size + 1
            and len(s) <= 80):
        return "h2"
    return "body"


def _merge_body_blocks(blocks: list[tuple], body_size: float) -> list[tuple]:
    """Merge consecutive body-sized blocks into paragraphs.

    Word PDFs put each visual line of a paragraph as a separate text block.
    Without merging, every line becomes its own paragraph in the output.
    Rules:
      - Only merge body-sized blocks (max font close to body_size)
      - Stop at sentence-terminal punctuation (. ! ?)
      - Stop at heading-sized blocks, bullet lines, numbered headings,
        ALL_CAPS short lines, page-numbers, KV lines
    """
    out: list[tuple] = []
    for tup in blocks:
        text, max_sz, dom_sz, bold, align = tup
        s = text.strip()
        if not s:
            continue
        # Use dominant size (most chars by count), NOT max size, for break
        # detection. PyMuPDF reports max as the LARGEST glyph in the line —
        # a single tall character drags max to ~14.7pt for what is plainly
        # a 12pt body line, and fragmenting paragraphs on that is wrong.
        effective_sz = dom_sz if dom_sz > 0 else max_sz
        prev_align = out[-1][4] if out else align
        # Only treat alignment as a break when it's a clear layout shift:
        # short trailing body fragments often misclassify as "center" because
        # their bbox midpoint sits closer to page-mid than the long
        # preceding line.
        align_break = (
            align != prev_align
            and align in ("center", "right")
            and prev_align in ("center", "right")
        )
        is_break = (
            effective_sz > body_size + 1
            or _PDF_BULLET_RE.match(s)
            or _PDF_NUM_LIST_RE.match(s)
            or (_PDF_ALL_CAPS_RE.match(s) and len(s) < 60)
            or _PDF_NUM_HEADING_RE.match(s)
            or _PDF_PAGE_NUM_RE.match(s)
            or _PDF_BODY_PAGENUM_RE.match(s)
            or _PDF_KV_RE.match(s)
            or align_break
        )
        # Bold-flag equality is too strict for body text: PyMuPDF marks the
        # whole block bold if any single span is bold, fragmenting paragraphs
        # that contain one bold word. Only enforce bold-match above body size
        # (i.e. for headings, where bold contrast is meaningful).
        prev_dom = out[-1][2] if out else dom_sz
        size_match = abs(prev_dom - effective_sz) < 1.0
        if (out
                and not is_break
                and not out[-1][0].rstrip().endswith((".", "?", "!", ":", ";"))
                and size_match
                and (effective_sz <= body_size + 1 or out[-1][3] == bold)):
            prev = out[-1]
            merged = (prev[0].rstrip() + " " + s.lstrip()).strip()
            out[-1] = (merged, prev[1], prev[2], prev[3], prev[4])
        else:
            out.append(tup)
    return out


def _dedupe_empty_chapters(chapters: list) -> list:
    """Collapse consecutive chapters with the same heading where the earlier
    one is empty (chapter divider page followed by content page repeats)."""
    if not chapters:
        return chapters
    out: list = [chapters[0]]
    for ch in chapters[1:]:
        prev = out[-1]
        prev_words = len(prev.get("content", "").split())
        same_name  = (prev["heading"].strip().upper()
                      == ch["heading"].strip().upper())
        if same_name and prev_words == 0:
            # Replace empty divider with content version
            ch["number"] = prev["number"]
            out[-1] = ch
        else:
            ch["number"] = len(out) + 1
            out.append(ch)
    return out


def _pdf_is_toc_page(blocks: list[tuple]) -> bool:
    """Detect a Table of Contents page by its content signature.

    A ToC page typically contains "CONTENTS" / "TABLE OF CONTENTS" as a
    title and lists multiple known section names in succession.
    """
    has_toc_title = False
    section_hits = 0
    for text, _max_sz, _dom_sz, _bold, _align in blocks:
        u = text.strip(" ?:.").upper()
        if u in ("CONTENTS", "TABLE OF CONTENTS"):
            has_toc_title = True
        if u in _PDF_KNOWN_SECTION_NAMES:
            section_hits += 1
    return has_toc_title and section_hits >= 3


def detect_structure_from_pdf(pdf_bytes: bytes) -> dict:
    """Detect chapter/heading structure from a PDF deterministically.

    Returns the same shape as detect_structure_from_docx. No LLM in path.
    Uses font-size-aware classification so cover-page text fragments and ToC
    column headers are not mistaken for chapter headings.
    """
    import fitz
    try:
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        logger.warning("detect_structure_from_pdf open failed: %s", exc)
        return {"error": "pdf_open_failed"}
    try:
        body_size = _pdf_estimate_body_size(pdf)
        title = ""
        chapters: list = []
        current_chapter: dict | None = None

        for page_no in range(len(pdf)):
            blocks = _pdf_blocks_with_font(pdf, page_no)
            is_toc = _pdf_is_toc_page(blocks)
            for text, max_sz, _dom_sz, _bold, block_align in blocks:
                # Title: largest text on cover page
                if (page_no == 0 and not title
                        and max_sz >= body_size + 4
                        and 8 <= len(text) <= 120):
                    title = text.strip(" ?:.")
                    continue

                cls = _pdf_classify_block(text, max_sz, body_size, page_no,
                                           is_toc_page=is_toc)
                if cls == "skip":
                    continue
                if cls == "h1":
                    current_chapter = {
                        "number":   len(chapters) + 1,
                        "heading":  text.strip(" ?:."),
                        "sections": [],
                        "content":  "",
                    }
                    chapters.append(current_chapter)
                    continue
                if current_chapter is not None:
                    current_chapter["content"] = (
                        current_chapter["content"] + " " + text
                    ).strip()

        chapters = _dedupe_empty_chapters(chapters)
        if not chapters and not title:
            return {"error": "no_structure_found"}
        return {
            "title":      title or "PROJECT REPORT",
            "chapters":   chapters,
            "references": [],
        }
    finally:
        pdf.close()


# ---------------------------------------------------------------------------
# Claude Haiku — structure parser
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Claude structure parser — Sonnet 4.6 with Opus 4.5 escalation
# ---------------------------------------------------------------------------
# Strategy:
#   Pass 1 — Sonnet 4.6 + 8K extended thinking, strict JSON via tool_use.
#   Validate the structure deterministically (see _validate_structure).
#   If validation fails → caller decides to escalate to Opus or fail.
#   Pass 2 — Opus 4.5 + 16K extended thinking, same schema, given Sonnet's
#            attempt + diagnostic as a hint.
#   If Opus also fails → caller routes to operator queue (never silent fall-
#   through to a "FORMATTED DOCUMENT" placeholder).
# ---------------------------------------------------------------------------

# Model identifiers — env-var overridable. If Anthropic deprecates an alias
# or you want to pin a specific dated snapshot, set ANTHROPIC_MODEL_SONNET
# or ANTHROPIC_MODEL_OPUS in the Vercel dashboard (no redeploy needed).
_MODEL_SONNET = os.environ.get("ANTHROPIC_MODEL_SONNET", "claude-sonnet-4-6")
_MODEL_OPUS   = os.environ.get("ANTHROPIC_MODEL_OPUS",   "claude-opus-4-5")
_THINKING_SONNET_TOKENS = 8000
_THINKING_OPUS_TOKENS   = 16000


def verify_models_available() -> dict:
    """Ping each configured model with a 1-token message. Returns a dict
    mapping model_id -> "ok" or "FAIL: <reason>". Use after every deploy
    via the GET /admin/health/models endpoint to surface 404s/auth issues
    BEFORE the first paying customer hits them.

    Cost: ~₹0.01 per call (one trivial message per model). Safe to call
    on a manual admin route; do not put on a cron unless you want spam.
    """
    import anthropic

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return {"_status": "no_api_key"}

    client = anthropic.Anthropic(api_key=api_key)
    out: dict[str, str] = {"_status": "checked"}
    for mid in (_MODEL_SONNET, _MODEL_OPUS):
        try:
            client.messages.create(
                model=mid,
                max_tokens=1,
                messages=[{"role": "user", "content": "1"}],
            )
            out[mid] = "ok"
        except Exception as exc:
            out[mid] = f"FAIL: {exc}"
            logger.error("Model unavailable at health-check: %s -> %s", mid, exc)
    return out

# Strict tool schema — forcing tool_use guarantees the model returns this
# exact JSON shape (or errors), eliminating the markdown-fences / freeform
# parsing fragility of the old Haiku flow.
_STRUCTURE_TOOL_SCHEMA: dict = {
    "name": "submit_structure",
    "description": (
        "Submit the parsed structure of an academic project report. "
        "Call this exactly once with all detected fields populated."
    ),
    "input_schema": {
        "type": "object",
        "required": [
            "title", "chapters", "references",
            "confidence_score", "missing_fields",
        ],
        "properties": {
            "title": {
                "type": "string",
                "description": "Full report title as written by the author.",
            },
            "subtitle": {
                "type": "string",
                "description": "Optional subtitle / dissertation line.",
            },
            "author": {
                "type": "string",
                "description": "Student author full name (empty string if not detected).",
            },
            "guide": {
                "type": "string",
                "description": "Project guide / supervisor name (empty if missing).",
            },
            "college": {
                "type": "string",
                "description": "College / department name.",
            },
            "university": {
                "type": "string",
                "description": "University name.",
            },
            "year": {
                "type": "string",
                "description": "Academic year (e.g. '2025-26').",
            },
            "abstract": {
                "type": "string",
                "description": "Abstract / executive summary if present.",
            },
            "declaration": {
                "type": "string",
                "description": (
                    "Declaration text if present in the source (typically begins "
                    "'I hereby declare...'). Leave empty if not detected — the "
                    "renderer will emit a red-marked template the user must edit."
                ),
            },
            "acknowledgement": {
                "type": "string",
                "description": (
                    "Acknowledgement section if present in the source. Leave "
                    "empty if not detected — the renderer will emit a red template."
                ),
            },
            "keywords": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Up to 10 keywords describing the project.",
            },
            "chapters": {
                "type": "array",
                "description": "Detected chapters in document order.",
                "items": {
                    "type": "object",
                    "required": ["number", "heading", "content"],
                    "properties": {
                        "number":  {"type": "integer", "description": "1-based chapter index."},
                        "heading": {"type": "string"},
                        "content": {
                            "type": "string",
                            "description": "All prose under this chapter, joined.",
                        },
                        "sections": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "required": ["number", "heading", "content"],
                                "properties": {
                                    "number":  {"type": "string", "description": "e.g. '1.2'."},
                                    "heading": {"type": "string"},
                                    "content": {"type": "string"},
                                },
                            },
                        },
                    },
                },
            },
            "references": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Reference entries (numbered or unnumbered).",
            },
            "confidence_score": {
                "type": "number",
                "description": (
                    "Self-rated confidence 0.0-1.0 that this structure faithfully "
                    "reflects the author's intent. Score below 0.6 if the input "
                    "lacks clear chapter markers or you had to guess heavily."
                ),
            },
            "missing_fields": {
                "type": "array",
                "items": {"type": "string"},
                "description": (
                    "Names of fields you could not detect (e.g. 'guide', 'abstract'). "
                    "Empty array if every field was confidently filled."
                ),
            },
        },
    },
}


def _structure_system_prompt() -> str:
    """System prompt for the structure parser. Cacheable across calls."""
    return (
        "You are an expert academic editor. You parse student project reports "
        "and return their structure using the submit_structure tool. Rules:\n"
        "1. ALWAYS call submit_structure exactly once. Never reply in prose.\n"
        "2. Detect every chapter present. Do not invent chapters that are not "
        "in the source. Do not merge two distinct chapters into one.\n"
        "3. Preserve every word of body content under its chapter. The sum of "
        "chapter content lengths should approximate the total input length.\n"
        "4. Identify front-matter metadata (title, author, guide, college, "
        "university, year, abstract, keywords) only when actually present. "
        "Use an empty string for fields you cannot detect — never fabricate.\n"
        "5. Set confidence_score honestly. Below 0.6 if the input lacks clear "
        "chapter markers and you had to guess. Below 0.4 if the input is "
        "mostly unstructured prose with no chapter boundaries.\n"
        "6. List every field you could not confidently detect in missing_fields."
    )


def _call_claude_structure(
    model: str,
    thinking_budget: int,
    user_text: str,
    prior_attempt: dict | None = None,
) -> dict:
    """Low-level call to a Claude model with the structure tool forced.

    Returns the tool input dict on success, or {"error": "<reason>"} on
    any failure (network, API, tool-use not invoked, etc.).
    """
    import anthropic

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return {"error": "no_api_key"}

    client = anthropic.Anthropic(api_key=api_key)

    # Cap input at ~150k chars (~37k tokens) to leave room for thinking +
    # output within the 200k context window.
    truncated = user_text[:150_000]
    truncated_note = (
        "\n\n[NOTE: input was truncated to fit context]"
        if len(user_text) > 150_000 else ""
    )

    user_blocks: list[dict] = [
        {"type": "text", "text": f"Document text:\n{truncated}{truncated_note}"},
    ]
    if prior_attempt:
        diag = prior_attempt.get("validation_errors") or ["unspecified"]
        diag_str = "; ".join(diag)
        user_blocks.append({
            "type": "text",
            "text": (
                "A previous attempt by a smaller model failed validation: "
                f"{diag_str}. Their partial JSON is below — use it as a hint "
                "but rebuild correctly:\n```json\n"
                f"{json.dumps(prior_attempt.get('structure', {}), indent=2)}\n```"
            ),
        })

    try:
        # Note: Anthropic API does not allow combining `thinking` with a
        # forced tool_choice ({"type":"tool"|"any"}). We use tool_choice="auto"
        # and rely on the system prompt to ensure the model calls the tool.
        # If it ever skips the call, the caller sees a "no_tool_use_in_response"
        # validation error which escalates to Opus / the operator queue.
        message = client.messages.create(
            model=model,
            max_tokens=16_000,
            thinking={"type": "enabled", "budget_tokens": thinking_budget},
            system=[{
                "type": "text",
                "text": _structure_system_prompt(),
                "cache_control": {"type": "ephemeral"},
            }],
            tools=[_STRUCTURE_TOOL_SCHEMA],
            tool_choice={"type": "auto"},
            messages=[{"role": "user", "content": user_blocks}],
        )
    except Exception as exc:
        logger.error("claude structure call (%s) failed: %s", model, exc)
        return {"error": f"api_error: {exc}"}

    # Find the tool_use block in the response
    for block in message.content:
        if getattr(block, "type", None) == "tool_use":
            data = block.input if isinstance(block.input, dict) else {}
            data["model_used"] = model
            return data
    return {"error": "no_tool_use_in_response"}


def _coverage_ratio(structure: dict, input_text: str) -> float:
    """Fraction of input words that survived into chapter content."""
    input_words = max(1, len(input_text.split()))
    chapter_words = 0
    for ch in structure.get("chapters", []) or []:
        chapter_words += len((ch.get("content") or "").split())
        for sec in ch.get("sections", []) or []:
            chapter_words += len((sec.get("content") or "").split())
    return min(1.0, chapter_words / input_words)


def _validate_structure(structure: dict, input_text: str) -> list[str]:
    """Return list of human-readable validation failure reasons.

    Empty list = passes. Non-empty = caller should escalate or fail loudly.
    """
    errs: list[str] = []
    if not isinstance(structure, dict):
        return ["not_a_dict"]
    if "error" in structure:
        return [f"upstream_error: {structure['error']}"]

    title = (structure.get("title") or "").strip()
    if not title or len(title) < 5:
        errs.append("title_too_short")
    if len(title) > 400:
        errs.append("title_too_long")

    chapters = structure.get("chapters") or []
    if not isinstance(chapters, list):
        errs.append("chapters_not_a_list")
        return errs
    if len(chapters) < 3:
        errs.append(f"too_few_chapters ({len(chapters)} < 3)")

    for i, ch in enumerate(chapters):
        if not isinstance(ch, dict):
            errs.append(f"chapter[{i}]_not_a_dict")
            continue
        heading = (ch.get("heading") or "").strip()
        content = (ch.get("content") or "").strip()
        if not heading:
            errs.append(f"chapter[{i}]_missing_heading")
        if not content and not (ch.get("sections") or []):
            errs.append(f"chapter[{i}]_empty")

    conf = structure.get("confidence_score")
    try:
        conf_f = float(conf) if conf is not None else 0.0
    except (TypeError, ValueError):
        conf_f = 0.0
    if conf_f < 0.75:
        errs.append(f"confidence_low ({conf_f:.2f} < 0.75)")

    coverage = _coverage_ratio(structure, input_text)
    if coverage < 0.80:
        errs.append(f"coverage_low ({coverage:.2f} < 0.80)")

    return errs


def _parse_structure_sonnet(text: str) -> dict:
    """First-pass parser: Sonnet 4.6 + 8K thinking."""
    return _call_claude_structure(
        model=_MODEL_SONNET,
        thinking_budget=_THINKING_SONNET_TOKENS,
        user_text=text,
    )


def _parse_structure_opus(text: str, prior_attempt: dict) -> dict:
    """Second-pass parser: Opus 4.5 + 16K thinking, with prior hint."""
    return _call_claude_structure(
        model=_MODEL_OPUS,
        thinking_budget=_THINKING_OPUS_TOKENS,
        user_text=text,
        prior_attempt=prior_attempt,
    )


def _parse_structure_with_claude(text: str, allow_escalation: bool = True) -> dict:
    """Adaptive structure parser. Tries Sonnet first, escalates to Opus.

    Returns either a validated structure dict (with extra metadata
    fields ``model_used``, ``confidence_score``, ``coverage_ratio``,
    ``validation_errors``) or ``{"error": "<reason>", "validation_errors": [...]}``
    so the caller can fail loudly or route to the operator queue.

    ``allow_escalation=False`` is used for the free preview pass — we only
    pay for Opus when the customer has paid.
    """
    sonnet_result = _parse_structure_sonnet(text)
    sonnet_errs   = _validate_structure(sonnet_result, text)

    if not sonnet_errs:
        sonnet_result["coverage_ratio"]    = _coverage_ratio(sonnet_result, text)
        sonnet_result["validation_errors"] = []
        return sonnet_result

    logger.info(
        "Sonnet structure validation failed: %s — escalate=%s",
        sonnet_errs, allow_escalation,
    )

    if not allow_escalation:
        return {
            "error": "validation_failed_preview",
            "validation_errors": sonnet_errs,
            "model_used":        _MODEL_SONNET,
            "structure":         sonnet_result,
        }

    # Escalate to Opus with Sonnet's attempt + diagnostic as a hint.
    opus_result = _parse_structure_opus(text, prior_attempt={
        "structure":         sonnet_result,
        "validation_errors": sonnet_errs,
    })
    opus_errs = _validate_structure(opus_result, text)
    if not opus_errs:
        opus_result["coverage_ratio"]    = _coverage_ratio(opus_result, text)
        opus_result["validation_errors"] = []
        return opus_result

    logger.warning("Opus structure validation also failed: %s", opus_errs)
    return {
        "error": "validation_failed_both",
        "validation_errors": opus_errs,
        "sonnet_errors":     sonnet_errs,
        "model_used":        _MODEL_OPUS,
        "sonnet_structure":  sonnet_result,
        "opus_structure":    opus_result,
    }


# ---------------------------------------------------------------------------
# Style application
# ---------------------------------------------------------------------------

_ALIGN_MAP = {
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ---------------------------------------------------------------------------
# Red-text convention — AI-generated content the user must review
# ---------------------------------------------------------------------------
# Visual contract for the output DOCX:
#   * Red text (#C00000)   = AI-generated / template boilerplate / [edit me]
#   * Black text           = user's own content
# The publish-ready pass below stamps every run to black UNLESS the run is
# already red — that is, the colour is preserved through the discipline pass.

_RED_RGB = RGBColor(0xC0, 0x00, 0x00)


def _is_red_run(run) -> bool:
    """True if this run is explicitly marked with the red review sentinel.

    RGBColor is a tuple subclass in python-docx with no ``__int__`` defined,
    so we compare via its hex-string representation ("C00000").
    """
    try:
        c = run.font.color.rgb
        if c is None:
            return False
        return str(c).upper() == "C00000"
    except Exception:
        return False


def _add_red_run(para, text: str):
    """Append text as a red run, signalling AI-generated content to review.
    The publish-ready pass exempts these runs from black-stamping so the
    colour survives all the way to the saved DOCX.
    """
    run = para.add_run(text)
    run.font.color.rgb = _RED_RGB
    return run


# ---------------------------------------------------------------------------
# Publish-ready typography pass
# ---------------------------------------------------------------------------
# Runs at the END of every generation path. Enforces, per user spec:
#   * Main title (Title style)   16pt bold centered
#   * Heading 1 (chapter)        per config (e.g. 14pt bold), keep_with_next,
#                                page_break_before (except the first one)
#   * Heading 2 (section)        per config (e.g. 12pt bold), keep_with_next
#   * Heading 3 (sub-section)    per config (12pt bold italic), keep_with_next
#   * Normal (body)              body font/size, justified, line-spacing 1.5,
#                                space_before == space_after (even block spacing)
#   * Red runs preserved (AI-generated content user must review)
# Stamps run-level rFonts so inherited theme fonts cannot override.

_PUBLISH_TITLE_SIZE_PT = 16  # cover page title, per user spec


def _publish_ready_pass(doc: Document, config: dict) -> None:
    body_font = config.get("body_font", "Times New Roman")
    body_size = int(config.get("body_size_pt", 12))
    h1_size   = int((config.get("heading1") or {}).get("size_pt", body_size + 2))
    h2_size   = int((config.get("heading2") or {}).get("size_pt", body_size))
    h3_size   = int((config.get("heading3") or {}).get("size_pt", body_size))
    line_spc  = float(config.get("line_spacing", 1.5))
    sp_after  = int(config.get("para_space_after_pt", 6))
    # User's "even space above/below text block" rule -> equal both sides.
    sp_before = sp_after

    style_sizes = {
        "Title":     _PUBLISH_TITLE_SIZE_PT,
        "Heading 1": h1_size,
        "Heading 2": h2_size,
        "Heading 3": h3_size,
    }

    state = {"first_h1_seen": False}

    def _stamp_run(run, *, size_pt: int, bold: bool | None,
                    italic: bool | None = None) -> None:
        is_red = _is_red_run(run)
        run.font.name = body_font
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        if italic is not None and not is_red:
            run.font.italic = italic
        if not is_red:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        # Stamp explicit rFonts to defeat inherited theme fonts
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rFonts.set(qn(f"w:{attr}"), body_font)

    def _stamp_para(para) -> None:
        sn = (para.style.name if para.style else "Normal") or "Normal"
        pf = para.paragraph_format

        if sn == "Title":
            pf.alignment      = WD_ALIGN_PARAGRAPH.CENTER
            pf.keep_with_next = True
            for r in para.runs:
                _stamp_run(r, size_pt=style_sizes["Title"], bold=True)
            return

        if sn.startswith("Heading"):
            pf.keep_with_next = True
            if sn == "Heading 1":
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if state["first_h1_seen"]:
                    pf.page_break_before = True
                state["first_h1_seen"] = True
            else:
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            size   = style_sizes.get(sn, h2_size)
            italic = (sn == "Heading 3")
            for r in para.runs:
                _stamp_run(r, size_pt=size, bold=True, italic=italic)
            return

        # Body / lists / anything else.
        # Lists keep their own indent + alignment; only body gets justified.
        if sn not in ("List Bullet", "List Number"):
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = line_spc
        pf.space_before      = Pt(sp_before)
        pf.space_after       = Pt(sp_after)
        for r in para.runs:
            _stamp_run(r, size_pt=body_size, bold=None)

    for p in doc.paragraphs:
        _stamp_para(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _stamp_para(p)


def _apply_university_styles(doc: Document, config: dict) -> None:
    """Apply margins, body font, and heading styles from a university config."""
    m = config["margins"]
    for section in doc.sections:
        section.left_margin   = Cm(m["left_cm"])
        section.right_margin  = Cm(m["right_cm"])
        section.top_margin    = Cm(m["top_cm"])
        section.bottom_margin = Cm(m["bottom_cm"])

    body_font = config["body_font"]
    body_size = config["body_size_pt"]
    spacing   = config["line_spacing"]
    sp_before = config["para_space_before_pt"]
    sp_after  = config["para_space_after_pt"]

    normal = doc.styles["Normal"]
    normal.font.name = body_font
    normal.font.size = Pt(body_size)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing      = spacing
    normal.paragraph_format.space_before      = Pt(sp_before)
    normal.paragraph_format.space_after       = Pt(sp_after)

    heading_specs = {
        "Heading 1": config["heading1"],
        "Heading 2": config["heading2"],
        "Heading 3": config["heading3"],
    }
    for style_name, hcfg in heading_specs.items():
        try:
            hstyle = doc.styles[style_name]
        except KeyError:
            continue
        hstyle.font.name     = body_font
        hstyle.font.size     = Pt(hcfg["size_pt"])
        hstyle.font.bold     = hcfg["bold"]
        hstyle.font.all_caps = bool(hcfg.get("all_caps"))
        hstyle.paragraph_format.alignment = _ALIGN_MAP.get(
            hcfg.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT
        )
        hstyle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        hstyle.paragraph_format.line_spacing      = spacing
        hstyle.paragraph_format.space_before      = Pt(12)
        hstyle.paragraph_format.space_after       = Pt(6)


def _set_section_page_numbering(section, fmt: str = "decimal", start: int = 1) -> None:
    """Set page number format on a section via XML manipulation."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def _apply_margins_to_section(section, config: dict) -> None:
    m = config["margins"]
    section.left_margin   = Cm(m["left_cm"])
    section.right_margin  = Cm(m["right_cm"])
    section.top_margin    = Cm(m["top_cm"])
    section.bottom_margin = Cm(m["bottom_cm"])


# ---------------------------------------------------------------------------
# Low-level paragraph helpers
# ---------------------------------------------------------------------------

def _add_body_para(doc: Document, text: str, config: dict) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing      = config["line_spacing"]
    para.paragraph_format.space_before      = Pt(config["para_space_before_pt"])
    para.paragraph_format.space_after       = Pt(config["para_space_after_pt"])
    run = para.add_run(text)
    run.font.name = config["body_font"]
    run.font.size = Pt(config["body_size_pt"])


def _add_centered_bold(
    doc: Document, text: str, size_pt: int = 14, space_before: int = 0
) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(space_before)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)


# ---------------------------------------------------------------------------
# TOC field
# ---------------------------------------------------------------------------

def _add_toc_field(doc: Document) -> None:
    """Insert a Word auto-updating TOC field (update on open in Word)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _append_elem(elem):
        run = para.add_run()
        run._r.append(elem)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    _append_elem(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr = para.add_run()
    run_instr._r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    _append_elem(sep)

    placeholder = para.add_run()
    placeholder.text = "[Right-click → Update Field → Update entire table]"
    placeholder.font.italic = True

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    _append_elem(end)


# ---------------------------------------------------------------------------
# Front matter page builders
# ---------------------------------------------------------------------------

def _build_title_page(doc: Document, config: dict, meta: dict) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    college     = meta.get("college_name", "[COLLEGE NAME]")
    department  = meta.get("department", "[DEPARTMENT]")
    title       = meta.get("title", "[PROJECT TITLE]")
    report_type = meta.get("report_type", "B.Tech Final Year Project")
    acad_year   = meta.get("academic_year", "[ACADEMIC YEAR]")
    degree      = meta.get("degree", "Bachelor of Technology")

    _add_centered_bold(doc, college.upper(), size_pt=16)
    _add_centered_bold(doc, f"DEPARTMENT OF {department.upper()}", size_pt=13)

    doc.add_paragraph()
    doc.add_paragraph()
    _add_centered_bold(doc, title.upper(), size_pt=16)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"A {report_type.upper()} REPORT").bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Submitted to")

    _add_centered_bold(doc, config["name"].upper(), size_pt=13)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(
        "in partial fulfillment of the requirements for the award of the Degree of"
    )

    _add_centered_bold(
        doc, f"{degree.upper()}\nIN {department.upper()}", size_pt=13
    )

    doc.add_paragraph()
    _add_centered_bold(doc, "Submitted by", size_pt=12)
    for s in meta.get("students", []):
        name = s.get("name", "")
        roll = s.get("roll_no") or s.get("roll", "")
        if name:
            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.add_run(f"{name}  ({roll})")

    doc.add_paragraph()
    _add_centered_bold(doc, acad_year, size_pt=12)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        f"[Print on {config.get('cover_color', 'Navy Blue')} hard cover "
        f"with {config.get('cover_text_color', 'Gold')} text]"
    )
    nr.italic = True
    nr.font.size = Pt(10)


def _build_certificate_page(doc: Document, config: dict, meta: dict) -> None:
    doc.add_heading("CERTIFICATE", level=1)

    students      = meta.get("students", [])
    student_names = ", ".join(s.get("name", "") for s in students if s.get("name"))
    guide         = meta.get("guide", {})
    hod           = meta.get("hod", {})
    co_guide      = meta.get("co_guide", {})

    cert = config.get("certificate_text", "")
    cert = (
        cert
        .replace("[TITLE]",         meta.get("title", "[TITLE]"))
        .replace("[STUDENT_NAMES]", student_names or "[STUDENT NAMES]")
        .replace("[DEPARTMENT]",    meta.get("department", "[DEPARTMENT]"))
        .replace("[DEGREE]",        meta.get("degree", "Bachelor of Technology"))
        .replace("[GUIDE_NAME]",    guide.get("name", "[GUIDE NAME]"))
    )

    _add_body_para(doc, cert, config)
    doc.add_paragraph()

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.style = "Table Grid"
    cells = sig_table.rows[0].cells
    cells[0].text = (
        f"Guide:\n{guide.get('name', '[Guide Name]')}\n"
        f"{guide.get('designation', 'Assistant Professor')}"
    )
    cells[1].text = (
        f"Head of Department:\n{hod.get('name', '[HOD Name]')}\n"
        f"{hod.get('designation', 'Professor & Head')}"
    )

    if co_guide and co_guide.get("name"):
        doc.add_paragraph()
        _add_body_para(
            doc,
            f"Co-Guide: {co_guide['name']}, {co_guide.get('designation', '')}",
            config,
        )

    doc.add_paragraph()
    _add_body_para(doc, "Place: _______________", config)
    _add_body_para(doc, "Date:  _______________", config)


def _build_declaration_page(doc: Document, config: dict, meta: dict) -> None:
    doc.add_heading("DECLARATION", level=1)

    students      = meta.get("students", [])
    student_names = ", ".join(s.get("name", "") for s in students if s.get("name"))
    guide         = meta.get("guide", {})

    decl = config.get("declaration_text", "")
    decl = (
        decl
        .replace("[TITLE]",         meta.get("title", "[TITLE]"))
        .replace("[STUDENT_NAMES]", student_names or "[STUDENT NAMES]")
        .replace("[DEPARTMENT]",    meta.get("department", "[DEPARTMENT]"))
        .replace("[DEGREE]",        meta.get("degree", "Bachelor of Technology"))
        .replace("[GUIDE_NAME]",    guide.get("name", "[GUIDE NAME]"))
    )

    _add_body_para(doc, decl, config)
    doc.add_paragraph()
    _add_body_para(doc, "Place: _______________", config)
    _add_body_para(doc, "Date:  _______________", config)
    doc.add_paragraph()
    _add_body_para(doc, "Signature(s):", config)
    for s in students:
        name = s.get("name", "")
        roll = s.get("roll_no") or s.get("roll", "")
        if name:
            _add_body_para(doc, f"\n_______________\n{name} ({roll})", config)


def _build_acknowledgement_page(doc: Document, config: dict) -> None:
    doc.add_heading("ACKNOWLEDGEMENT", level=1)
    text = (
        "We express our sincere gratitude to our guide for the invaluable guidance, "
        "encouragement, and support extended to us throughout this project. We also thank "
        "the Head of the Department and all faculty members for their help and inspiration. "
        "We are grateful to the Management and Principal of our college for providing the "
        "necessary facilities. We express our deep gratitude to our family and friends for "
        "their continuous support and motivation.\n\n"
        "[Add your personal acknowledgements here]"
    )
    _add_body_para(doc, text, config)


def _build_abstract_page(doc: Document, config: dict, meta: dict) -> None:
    doc.add_heading("ABSTRACT", level=1)
    abstract = meta.get(
        "abstract",
        "[Insert abstract here — 100 to 300 words summarising your project, "
        "methodology, and key findings.]",
    )
    _add_body_para(doc, abstract, config)

    keywords = meta.get("keywords", [])
    if keywords:
        doc.add_paragraph()
        kp = doc.add_paragraph()
        kr = kp.add_run("Keywords: ")
        kr.bold = True
        kr.font.size = Pt(config["body_size_pt"])
        kp.add_run(", ".join(keywords))


def _build_toc_page(doc: Document) -> None:
    doc.add_heading("TABLE OF CONTENTS", level=1)
    _add_toc_field(doc)
    note = doc.add_paragraph()
    nr = note.add_run(
        "Note: Open in Word → right-click the table above → Update Field → "
        "Update entire table."
    )
    nr.italic = True
    nr.font.size = Pt(9)


def _build_list_of_figures_page(doc: Document) -> None:
    doc.add_heading("LIST OF FIGURES", level=1)
    p = doc.add_paragraph()
    p.add_run("[Right-click → Update Field after inserting all figures]").italic = True


def _build_list_of_tables_page(doc: Document) -> None:
    doc.add_heading("LIST OF TABLES", level=1)
    p = doc.add_paragraph()
    p.add_run("[Right-click → Update Field after inserting all tables]").italic = True


def _build_abbreviations_page(doc: Document, config: dict) -> None:
    doc.add_heading("LIST OF ABBREVIATIONS", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Abbreviation"
    hdr[1].text = "Full Form"
    for _ in range(8):
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = ""


# ---------------------------------------------------------------------------
# Chapter builder
# ---------------------------------------------------------------------------

def _join_text_lines(content: str) -> list[str]:
    """Convert a wrapped-line text blob into proper paragraphs.

    PDF-extracted text and Claude-Haiku-parsed content typically arrive
    with one ``\\n`` between every visual line of a single paragraph (no
    blank-line separators). Splitting on ``\\n`` and emitting each line
    as a Word paragraph produces mid-sentence breaks.

    Rule set (same heuristic as the PDF parser's merge_body_blocks):
      * blank line                                       -> boundary
      * line starts with bullet glyph or numbered marker -> new paragraph
      * line is ALL CAPS and short (<60 chars)           -> new paragraph
      * previous line ends with sentence-terminal punct  -> new paragraph
      * otherwise                                        -> append with space
    """
    if not content:
        return []
    bullet_re    = _PDF_BULLET_RE
    num_list_re  = _PDF_NUM_LIST_RE
    all_caps_re  = _PDF_ALL_CAPS_RE
    sentence_end = (".", "!", "?", ":", ";")

    paragraphs: list[str] = []
    for raw in content.split("\n"):
        line = raw.strip()
        if not line:
            if paragraphs and paragraphs[-1] != "":
                paragraphs.append("")
            continue

        is_new_paragraph = (
            not paragraphs
            or paragraphs[-1] == ""
            or bullet_re.match(line)
            or num_list_re.match(line)
            or (all_caps_re.match(line) and len(line) < 60)
            or paragraphs[-1].rstrip().endswith(sentence_end)
        )
        if is_new_paragraph:
            if paragraphs and paragraphs[-1] == "":
                paragraphs[-1] = line
            else:
                paragraphs.append(line)
        else:
            paragraphs[-1] = (paragraphs[-1].rstrip() + " " + line).strip()

    return [p for p in paragraphs if p]


def _build_chapter(doc: Document, chapter: dict, chapter_num: int, config: dict) -> None:
    heading_text = f"CHAPTER {chapter_num}: {chapter.get('heading', '').upper()}"
    doc.add_heading(heading_text, level=1)

    sections = chapter.get("sections", [])
    if sections:
        for sec in sections:
            sec_num     = sec.get("number", "")
            sec_heading = sec.get("heading", "")
            if sec_heading:
                doc.add_heading(f"{sec_num} {sec_heading}", level=2)
            content = sec.get("content", "")
            for para in _join_text_lines(content):
                _add_body_para(doc, para, config)
    else:
        content = chapter.get("content", "[Chapter content goes here]")
        for para in _join_text_lines(content):
            _add_body_para(doc, para, config)


def _build_references_page(doc: Document, references: list, config: dict) -> None:
    doc.add_heading("REFERENCES", level=1)
    if references:
        for i, ref in enumerate(references, 1):
            ref = ref.strip()
            if ref:
                if not ref.startswith("["):
                    ref = f"[{i}] {ref}"
                _add_body_para(doc, ref, config)
    else:
        _add_body_para(
            doc,
            "[Add references here in IEEE/APA format, one per line]",
            config,
        )


def _build_appendices_page(doc: Document, appendices: str, config: dict) -> None:
    doc.add_heading("APPENDICES", level=1)
    content = appendices.strip() if appendices else ""
    if content:
        for para in _join_text_lines(content):
            _add_body_para(doc, para, config)
    else:
        _add_body_para(
            doc,
            "[Appendix content — code listings, data tables, circuit diagrams, etc.]",
            config,
        )


# ---------------------------------------------------------------------------
# Page numbering and section management
# ---------------------------------------------------------------------------

def _add_chapter_section(doc: Document, config: dict) -> None:
    """Insert a new page section for chapters; set Arabic numbering from 1."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_margins_to_section(section, config)
    _set_section_page_numbering(section, fmt="decimal", start=1)


def _set_front_matter_numbering(doc: Document) -> None:
    """Set Roman numeral page numbering on section 0 (front matter)."""
    _set_section_page_numbering(doc.sections[0], fmt="upperRoman", start=1)


# ---------------------------------------------------------------------------
# Footer CTA
# ---------------------------------------------------------------------------

def _add_footer_cta(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(
            "Once your report is ready — Print, bind & deliver with Printosky, Thrissur. "
            "WhatsApp: 9495706405"
        )
        fr.font.size = Pt(9)
        fr.italic = True


# ---------------------------------------------------------------------------
# Shared front matter dispatch table
# ---------------------------------------------------------------------------

_DEFAULT_FRONT_MATTER_ORDER = [
    "title_page", "certificate", "declaration", "acknowledgement",
    "abstract", "toc", "list_of_figures", "list_of_tables", "list_of_abbreviations",
]


def _build_front_matter(doc: Document, config: dict, meta: dict) -> None:
    """Build all front matter pages in university-specified order."""
    order = config.get("front_matter_order", _DEFAULT_FRONT_MATTER_ORDER)

    builders = {
        "title_page":            lambda: _build_title_page(doc, config, meta),
        "certificate":           lambda: _build_certificate_page(doc, config, meta),
        "declaration":           lambda: _build_declaration_page(doc, config, meta),
        "acknowledgement":       lambda: _build_acknowledgement_page(doc, config),
        "abstract":              lambda: _build_abstract_page(doc, config, meta),
        "toc":                   lambda: _build_toc_page(doc),
        "list_of_figures":       lambda: _build_list_of_figures_page(doc),
        "list_of_tables":        lambda: _build_list_of_tables_page(doc),
        "list_of_abbreviations": lambda: _build_abbreviations_page(doc, config),
    }

    for i, item in enumerate(order):
        if i > 0:
            doc.add_page_break()
        builder = builders.get(item)
        if builder:
            builder()


# ---------------------------------------------------------------------------
# Public API — three product tiers
# ---------------------------------------------------------------------------

def generate_free_template(university_id: str) -> bytes:
    """
    Build a blank .docx template with placeholder text and correct university formatting.
    Returns .docx as bytes.
    """
    config = load_university_config(university_id)
    doc = Document()
    _apply_university_styles(doc, config)
    _set_front_matter_numbering(doc)

    placeholder_meta = {
        "title":        "[YOUR PROJECT TITLE]",
        "college_name": "[YOUR COLLEGE NAME]",
        "department":   "[YOUR DEPARTMENT]",
        "academic_year": "[ACADEMIC YEAR e.g. 2025-26]",
        "report_type":  "B.Tech Final Year Project",
        "degree":       "Bachelor of Technology",
        "students": [
            {"name": "[Student Name 1]", "roll_no": "[Roll No]"},
            {"name": "[Student Name 2]", "roll_no": "[Roll No]"},
        ],
        "guide":    {"name": "[Guide Name]",  "designation": "Assistant Professor"},
        "co_guide": {"name": "",              "designation": ""},
        "hod":      {"name": "[HOD Name]",    "designation": "Professor & Head"},
        "abstract": (
            "[Write your abstract here — 100 to 300 words summarising the problem, "
            "methodology, and results of your project.]"
        ),
        "keywords": ["keyword1", "keyword2", "keyword3", "keyword4", "keyword5"],
    }

    _build_front_matter(doc, config, placeholder_meta)

    # Chapters section (Arabic numbering from page 1)
    _add_chapter_section(doc, config)

    placeholder_chapters = [
        ("INTRODUCTION", [
            ("1.1", "Background",  "[Background content here]"),
            ("1.2", "Objectives",  "[Project objectives here]"),
            ("1.3", "Scope",       "[Scope of the project here]"),
        ]),
        ("LITERATURE REVIEW",       []),
        ("METHODOLOGY",             []),
        ("IMPLEMENTATION",          []),
        ("RESULTS AND DISCUSSION",  []),
        ("CONCLUSION AND FUTURE WORK", []),
    ]

    for i, (ch_heading, sections) in enumerate(placeholder_chapters, 1):
        if i > 1:
            doc.add_page_break()
        doc.add_heading(f"CHAPTER {i}: {ch_heading}", level=1)
        if sections:
            for sec_num, sec_heading, sec_content in sections:
                doc.add_heading(f"{sec_num} {sec_heading}", level=2)
                _add_body_para(doc, sec_content, config)
        else:
            _add_body_para(doc, f"[{ch_heading.title()} content goes here]", config)

    doc.add_page_break()
    _build_references_page(doc, [], config)
    doc.add_page_break()
    _build_appendices_page(doc, "", config)

    # Publish-ready typography pass: stamps every run/paragraph to spec
    # (16/14/12pt, justified, 1.5LS, even block spacing, keep_with_next on
    # headings, page_break_before on chapter H1s, red-run preservation).
    _publish_ready_pass(doc, config)

    _add_footer_cta(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_format_fix_front_matter(
    doc: Document, config: dict, structure: dict,
) -> None:
    """Build cover + declaration + acknowledgement + abstract + TOC from
    a Sonnet-extracted structure dict. Missing fields surface as red text
    so the user can spot exactly what to edit before submission.
    """
    title    = (structure.get("title")    or "[PROJECT TITLE]").strip()
    author   = (structure.get("author")   or "").strip()
    guide    = (structure.get("guide")    or "").strip()
    college  = (structure.get("college")  or "").strip()
    year     = (structure.get("year")     or "").strip()
    abstract = (structure.get("abstract") or "").strip()
    declaration_text = (structure.get("declaration") or "").strip()
    ack_text         = (structure.get("acknowledgement") or "").strip()

    def _centered(text_or_red: str, *, bold: bool = False, red: bool = False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if red:
            run = _add_red_run(p, text_or_red)
        else:
            run = p.add_run(text_or_red)
        if bold:
            run.bold = True
        return p

    # --- Cover page -------------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph()

    # Title in Title style — publish-ready pass will make it 16pt bold center
    title_para = doc.add_paragraph(style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run(title.upper()).bold = True

    doc.add_paragraph()

    _centered("Submitted to")
    _centered(config.get("name", "[UNIVERSITY]"), bold=True)

    doc.add_paragraph()

    if college:
        _centered(college, bold=True)
    else:
        _centered("[COLLEGE NAME — please fill in]", bold=True, red=True)

    doc.add_paragraph()
    doc.add_paragraph()

    _centered("Submitted by")
    if author:
        _centered(author, bold=True)
    else:
        _centered("[STUDENT NAME — please fill in]", bold=True, red=True)

    doc.add_paragraph()

    _centered("Under the guidance of")
    if guide:
        _centered(guide, bold=True)
    else:
        _centered("[GUIDE NAME — please fill in]", bold=True, red=True)

    doc.add_paragraph()
    doc.add_paragraph()

    if year:
        _centered(year, bold=True)
    else:
        _centered("[ACADEMIC YEAR — please fill in]", bold=True, red=True)

    doc.add_page_break()

    # --- Declaration ------------------------------------------------------
    doc.add_heading("DECLARATION", level=1)
    if declaration_text:
        _add_body_para(doc, declaration_text, config)
    else:
        boilerplate = (
            f"I hereby declare that the project report titled \"{title}\" "
            f"submitted to {config.get('name', '[UNIVERSITY]')} is a record "
            "of original work done by me. The content of this project has not "
            "been submitted by me for the award of any degree, diploma, or "
            "any other title to any other University or Institution."
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_red_run(p, boilerplate)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Place: ")
    _add_red_run(p, "_______________")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Date:  ")
    _add_red_run(p, "_______________")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if author:
        p.add_run(f"Signature: {author}")
    else:
        p.add_run("Signature: ")
        _add_red_run(p, "_______________")

    doc.add_page_break()

    # --- Acknowledgement --------------------------------------------------
    doc.add_heading("ACKNOWLEDGEMENT", level=1)
    if ack_text:
        _add_body_para(doc, ack_text, config)
    else:
        boilerplate = (
            "I take this opportunity to express my profound gratitude to "
            f"{guide or '[GUIDE NAME]'} for invaluable guidance, constant "
            "encouragement, and constructive suggestions throughout the "
            "course of this project. I also extend my sincere thanks to the "
            "Head of the Department and all faculty members for their "
            "support. I am grateful to "
            f"{college or '[COLLEGE NAME]'} for providing the resources and "
            "infrastructure necessary for this work. Finally, I thank my "
            "family and friends for their unwavering support."
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_red_run(p, boilerplate)

    doc.add_page_break()

    # --- Abstract ---------------------------------------------------------
    doc.add_heading("ABSTRACT", level=1)
    if abstract:
        _add_body_para(doc, abstract, config)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_red_run(p, (
            "[ABSTRACT — Please add a 150–250 word summary of your project "
            "here, covering the problem statement, methodology, and key "
            "findings.]"
        ))

    doc.add_page_break()

    # --- Table of Contents (auto-field) -----------------------------------
    _build_toc_page(doc)


def format_fix(
    text: str,
    university_id: str,
    *,
    allow_escalation: bool = True,
) -> bytes:
    """Re-format a pasted/extracted text document into a publish-ready DOCX.

    Parsing strategy is adaptive: Sonnet 4.6 + 8K thinking first; if the
    detected structure fails the deterministic validation gate AND
    ``allow_escalation=True`` (post-payment phase), escalate to Opus 4.5 +
    16K thinking. If both passes still fail → raise StructureDetectionError
    so the API layer can route to the operator queue (post-payment) or
    surface an upsell (free preview).

    Pass ``allow_escalation=False`` from the free-preview endpoint to
    avoid paying for Opus before the customer commits.

    Returns the DOCX bytes on success.
    """
    config = load_university_config(university_id)
    doc = Document()
    _apply_university_styles(doc, config)
    _set_front_matter_numbering(doc)

    structure = _parse_structure_with_claude(text, allow_escalation=allow_escalation)

    # Fail loudly. The API layer turns these into:
    #   preview phase     -> 4xx with upsell prompt
    #   post-payment phase -> route to operator queue + WhatsApp the customer
    # NEVER fall back to a silent "FORMATTED DOCUMENT" placeholder.
    if "error" in structure:
        err_code = structure["error"]
        phase = "preview" if not allow_escalation else "post_payment"
        partial = (
            structure.get("structure")
            or structure.get("opus_structure")
            or structure.get("sonnet_structure")
            or {}
        )
        raise StructureDetectionError(
            f"Structure detection failed ({err_code})",
            errors=structure.get("validation_errors", []),
            model_used=structure.get("model_used", ""),
            partial_structure=partial,
            phase=phase,
        )

    chapters   = structure.get("chapters", [])
    references = structure.get("references", [])

    # Front matter: cover -> declaration -> acknowledgement -> abstract -> TOC.
    # Missing fields render in red so the user knows what to edit.
    _build_format_fix_front_matter(doc, config, structure)

    # Body: chapters (page-break-before is enforced by _publish_ready_pass
    # on every H1 after the first, so we don't need explicit breaks here).
    _add_chapter_section(doc, config)
    for i, chapter in enumerate(chapters, 1):
        _build_chapter(doc, chapter, i, config)

    if references:
        doc.add_page_break()
        _build_references_page(doc, references, config)

    # Publish-ready typography pass: stamps every run/paragraph to spec
    # (16/14/12pt, justified, 1.5LS, even block spacing, keep_with_next on
    # headings, page_break_before on chapter H1s, red-run preservation).
    _publish_ready_pass(doc, config)

    _add_footer_cta(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_from_form(form_data: dict, university_id: str) -> bytes:
    """
    Build a complete project report from structured form data (Tier 3).
    Returns .docx as bytes.

    Accepts both the nested API format and the flat web-form format emitted
    by project-builder.html's collectFormData():
      - college / college_name
      - year / academic_year
      - guide_name + guide_desig  vs  guide: {name, designation}
      - coguide_name + coguide_desig  vs  co_guide: {name, designation}
      - hod_name + hod_desig  vs  hod: {name, designation}
      - chapter.heading  vs  chapter.title
    """
    def _person(dict_key: str, name_key: str, desig_key: str) -> dict:
        """Return a {name, designation} dict from either nested or flat fields."""
        p = form_data.get(dict_key)
        if isinstance(p, dict):
            return p
        return {
            "name":        form_data.get(name_key, ""),
            "designation": form_data.get(desig_key, ""),
        }

    config = load_university_config(university_id)
    doc = Document()
    _apply_university_styles(doc, config)
    _set_front_matter_numbering(doc)

    keywords_raw = form_data.get("keywords", "")
    if isinstance(keywords_raw, str):
        keywords = [k.strip() for k in keywords_raw.split(",") if k.strip()]
    else:
        keywords = list(keywords_raw)

    meta = {
        "title":        form_data.get("title", "[PROJECT TITLE]"),
        "college_name": form_data.get("college_name") or form_data.get("college", "[COLLEGE NAME]"),
        "department":   form_data.get("department", "[DEPARTMENT]"),
        "academic_year": form_data.get("academic_year") or form_data.get("year", "[ACADEMIC YEAR]"),
        "report_type":  form_data.get("report_type", "B.Tech Final Year Project"),
        "degree":       form_data.get("degree", "Bachelor of Technology"),
        "students":     form_data.get("students", []),
        "guide":        _person("guide",    "guide_name",   "guide_desig"),
        "co_guide":     _person("co_guide", "coguide_name", "coguide_desig"),
        "hod":          _person("hod",      "hod_name",     "hod_desig"),
        "abstract":     form_data.get("abstract", ""),
        "keywords":     keywords,
    }

    _build_front_matter(doc, config, meta)

    # Chapters section (Arabic numbering from page 1)
    _add_chapter_section(doc, config)

    chapters = form_data.get("chapters", [])
    for i, chapter in enumerate(chapters, 1):
        if i > 1:
            doc.add_page_break()
        ch_dict = {
            "heading":  chapter.get("title") or chapter.get("heading", f"Chapter {i}"),
            "sections": [],
            "content":  chapter.get("content", ""),
        }
        _build_chapter(doc, ch_dict, i, config)

    # References
    refs_raw   = form_data.get("references", "")
    references = [r.strip() for r in refs_raw.split("\n") if r.strip()] if refs_raw else []
    doc.add_page_break()
    _build_references_page(doc, references, config)

    # Appendices (optional)
    appendices = form_data.get("appendices", "")
    if appendices and appendices.strip():
        doc.add_page_break()
        _build_appendices_page(doc, appendices, config)

    # Publish-ready typography pass: stamps every run/paragraph to spec
    # (16/14/12pt, justified, 1.5LS, even block spacing, keep_with_next on
    # headings, page_break_before on chapter H1s, red-run preservation).
    _publish_ready_pass(doc, config)

    _add_footer_cta(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# In-place DOCX reformatter — preserves images, tables, all existing content
# ---------------------------------------------------------------------------

def _classify_with_claude_metadata(paras: list) -> list:
    """Ask Claude Haiku to classify paragraphs from metadata only — no full text sent.

    Only called when neither Word heading styles nor heuristics find any structure.
    Sends a compact list: index, first-60-chars snippet, font size, bold, length.
    Returns a list of strings: 'title'|'h1'|'h2'|'h3'|'body'|'blank'.
    """
    import anthropic

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return ["blank" if p["is_blank"] else "body" for p in paras]

    lines = []
    for i, p in enumerate(paras):
        if p["is_blank"]:
            lines.append(f"{i}: BLANK")
        else:
            snippet = p["text"][:60].replace("\n", " ")
            lines.append(
                f"{i}: text={repr(snippet)}, size={p['max_size']:.0f}pt, "
                f"bold={p['bold']}, len={p['length']}"
            )

    prompt = (
        "Classify each paragraph in an Indian university student project report.\n"
        "Output ONLY a JSON array of strings, one per line index, using these labels:\n"
        "  title = project/report title\n"
        "  h1    = chapter heading (e.g. CHAPTER 1, INTRODUCTION)\n"
        "  h2    = section heading (e.g. 1.1 Background)\n"
        "  h3    = sub-section heading\n"
        "  body  = regular paragraph\n"
        "  blank = empty line\n\n"
        "Paragraph metadata:\n"
        + "\n".join(lines)
        + "\n\nReturn ONLY the JSON array, no markdown, no explanation."
    )

    try:
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        result = json.loads(raw)
        if len(result) == len(paras):
            return result
    except Exception:
        pass

    return ["blank" if p["is_blank"] else "body" for p in paras]


def _xml_adjacent_paragraphs(prev_elem, curr_elem) -> bool:
    """True iff no ``<w:tbl>`` element sits between two paragraph elements
    in the body XML. Two paragraphs separated by a table belong to different
    logical regions and must never be stitched into one paragraph.
    """
    sib = prev_elem.getnext()
    while sib is not None:
        if sib is curr_elem:
            return True
        tag = sib.tag.rsplit("}", 1)[-1] if "}" in sib.tag else sib.tag
        if tag == "tbl":
            return False
        sib = sib.getnext()
    return False


def _stitch_broken_paragraphs(doc: Document, body_size: float = 12.0) -> int:
    """Merge body paragraphs that were split mid-sentence.

    DOCX files converted from PDFs (or hand-typed with hard line-breaks)
    frequently end up with one Word paragraph per visual line, breaking
    sentences across paragraph boundaries. This pass walks the document
    body, finds consecutive ``body``-classified paragraphs whose previous
    one does not end with sentence-terminal punctuation, and appends the
    current text into the previous paragraph (then removes the current
    paragraph's XML).

    Returns the number of paragraphs merged.
    """
    classifications = _classify_paragraphs(doc, body_size=body_size)
    paras = list(doc.paragraphs)
    sentence_end = (".", "!", "?", ":", ";")
    merges = 0

    prev_idx: int | None = None
    for i, (para, cls) in enumerate(zip(paras, classifications)):
        text = para.text.strip()
        if cls != "body" or not text:
            prev_idx = None
            continue

        # Lines starting with a bullet, numbered marker, or ALL-CAPS short
        # header look like the start of a new logical paragraph.
        if (_PDF_BULLET_RE.match(text)
                or _PDF_NUM_LIST_RE.match(text)
                or (_PDF_ALL_CAPS_RE.match(text) and len(text) < 60)):
            prev_idx = i
            continue

        if prev_idx is None:
            prev_idx = i
            continue

        prev_para = paras[prev_idx]
        prev_text = prev_para.text.rstrip()
        if prev_text.endswith(sentence_end):
            prev_idx = i
            continue

        if not _xml_adjacent_paragraphs(prev_para._element, para._element):
            prev_idx = i
            continue

        # Merge: append current text into previous, remove current's XML.
        prev_para.add_run(" " + text)
        curr_elem = para._element
        parent = curr_elem.getparent()
        if parent is not None:
            parent.remove(curr_elem)
        merges += 1
        # prev_idx unchanged — the next body paragraph may also merge into it

    return merges


def _classify_paragraphs(doc: Document, body_size: float = 12.0) -> list:
    """Classify every paragraph in a Document as title/h1/h2/h3/body/blank.

    Three-pass strategy (stops at first success):
      1. Word heading styles — most accurate, zero cost
      2. Heuristics — bold + font size + ALL_CAPS patterns
      3. Claude Haiku on paragraph metadata — fallback when doc has no styling at all
    """
    paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        sname = (para.style.name or "Normal") if para.style else "Normal"
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        run_sizes = [run.font.size.pt for run in para.runs if run.font.size]
        max_size = max(run_sizes) if run_sizes else body_size
        paras.append({
            "text":     text,
            "style":    sname,
            "bold":     is_bold,
            "max_size": max_size,
            "length":   len(text),
            "is_upper": (text == text.upper() and len(text) > 3) if text else False,
            "is_blank": not text,
        })

    # Pass 1 — Word heading styles
    has_word_headings = any(
        p["style"].startswith("Heading") or p["style"] == "Title"
        for p in paras if not p["is_blank"]
    )
    if has_word_headings:
        result = []
        for p in paras:
            s = p["style"]
            if p["is_blank"]:      result.append("blank")
            elif s == "Title":     result.append("title")
            elif s == "Heading 1": result.append("h1")
            elif s == "Heading 2": result.append("h2")
            elif s == "Heading 3": result.append("h3")
            else:                  result.append("body")
        return result

    # Pass 2 — heuristics
    heuristic = []
    for p in paras:
        if p["is_blank"]:
            heuristic.append("blank")
        elif p["bold"] and p["length"] < 120 and (p["max_size"] >= 14 or p["is_upper"]):
            heuristic.append("h1")
        elif p["bold"] and p["length"] < 120 and p["max_size"] >= 13:
            heuristic.append("h2")
        elif p["bold"] and p["length"] < 80:
            heuristic.append("h3")
        else:
            heuristic.append("body")

    if any(c in ("h1", "h2", "h3", "title") for c in heuristic):
        return heuristic

    # Pass 3 — Claude Haiku on paragraph metadata (last resort)
    return _classify_with_claude_metadata(paras)


def detect_structure_from_docx(docx_bytes: bytes) -> dict:
    """Detect document structure using smart 3-pass classification.

    Pass 1: reads Word heading styles (zero cost, most accurate).
    Pass 2: heuristics (bold + font size + ALL_CAPS).
    Pass 3: Claude Haiku on paragraph metadata (when doc has no styling).
    Returns the same shape as _parse_structure_with_claude.
    """
    doc = Document(io.BytesIO(docx_bytes))
    classifications = _classify_paragraphs(doc, body_size=12.0)

    title = ""
    chapters: list = []
    current_chapter: dict | None = None
    current_section: dict | None = None

    for para, cls in zip(doc.paragraphs, classifications):
        text = para.text.strip()
        if not text or cls == "blank":
            continue

        if cls == "title":
            title = text
        elif cls == "h1":
            current_section = None
            current_chapter = {
                "number":   len(chapters) + 1,
                "heading":  text,
                "sections": [],
                "content":  "",
            }
            chapters.append(current_chapter)
        elif cls == "h2":
            if current_chapter is None:
                current_chapter = {
                    "number":   len(chapters) + 1,
                    "heading":  text,
                    "sections": [],
                    "content":  "",
                }
                chapters.append(current_chapter)
                current_section = None
            else:
                current_section = {
                    "number":  f"{len(chapters)}.{len(current_chapter['sections']) + 1}",
                    "heading": text,
                    "content": "",
                }
                current_chapter["sections"].append(current_section)
        elif cls == "h3":
            pass  # don't model deeply
        else:  # body
            if current_section is not None:
                current_section["content"] = (
                    current_section["content"] + " " + text
                ).strip()
            elif current_chapter is not None:
                current_chapter["content"] = (
                    current_chapter["content"] + " " + text
                ).strip()

    if not chapters and not title:
        return {"error": "no_structure_found"}

    return {
        "title":      title or "PROJECT REPORT",
        "chapters":   chapters,
        "references": [],
    }


def format_fix_docx_inplace(docx_bytes: bytes, university_id: str) -> bytes:
    """Reformat an existing DOCX in place — preserves images, tables, all content.

    Uses smart 3-pass heading classification so even unstyled student documents
    get correct Heading 1/2/3 styles applied.  Page margins and numbering are
    fixed to university spec.
    """
    config    = load_university_config(university_id)
    doc       = Document(io.BytesIO(docx_bytes))
    body_font = config["body_font"]
    body_size = config["body_size_pt"]
    spacing   = config["line_spacing"]
    sp_before = config["para_space_before_pt"]
    sp_after  = config["para_space_after_pt"]

    # Fix style definitions first (Normal, Heading 1/2/3 get university spec)
    _apply_university_styles(doc, config)

    # Stitch broken paragraphs BEFORE classification. PDF-converted DOCX
    # files commonly arrive with one Word paragraph per visual line,
    # which produces mid-sentence breaks in the output. The stitch pass
    # merges those back into proper paragraphs, then we classify.
    try:
        merged = _stitch_broken_paragraphs(doc, body_size=body_size)
        if merged:
            logger.info("format_fix_docx_inplace: stitched %d broken paragraphs", merged)
    except Exception as exc:
        logger.warning("format_fix_docx_inplace stitch pass failed: %s", exc)

    # Classify all paragraphs in one smart pass
    classifications = _classify_paragraphs(doc, body_size=body_size)

    style_map = {
        "title": "Heading 1",
        "h1":    "Heading 1",
        "h2":    "Heading 2",
        "h3":    "Heading 3",
    }

    for para, cls in zip(doc.paragraphs, classifications):
        if cls == "blank":
            continue
        if cls in style_map:
            para.style = doc.styles[style_map[cls]]
        else:  # body
            para.style = doc.styles["Normal"]
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing      = spacing
            pf.space_before      = Pt(sp_before)
            pf.space_after       = Pt(sp_after)
            for run in para.runs:
                run.font.name = body_font
                run.font.size = Pt(body_size)

    # Fix table cell fonts — preserve structure, just fix typography
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = body_font
                        if run.font.size:
                            run.font.size = Pt(body_size)

    # Fix page margins and numbering on every section
    for section in doc.sections:
        _apply_margins_to_section(section, config)
        _set_section_page_numbering(section, fmt="decimal", start=1)

    # Publish-ready typography pass: stamps every run/paragraph to spec
    # (16/14/12pt, justified, 1.5LS, even block spacing, keep_with_next on
    # headings, page_break_before on chapter H1s, red-run preservation).
    _publish_ready_pass(doc, config)

    _add_footer_cta(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Survey table parsing + chart regeneration
# ---------------------------------------------------------------------------

_PDF_SURVEY_TABLE_RE = re.compile(
    r"Table\s+(\d+)\s*:\s*(.+?)(?=\n)"
    r".*?TRUE"
    r"(?:\n+\s*(\d+)\s+(\d+)|\n+\s*(\d+)\s*\n+\s*(\d+))"
    r".*?FALSE"
    r"(?:\n+\s*(\d+)\s+(\d+)|\n+\s*(\d+)\s*\n+\s*(\d+))"
    r".*?TOTAL"
    r"(?:\n+\s*(\d+)\s+(\d+)|\n+\s*(\d+)\s*\n+\s*(\d+))",
    re.DOTALL | re.IGNORECASE,
)


def _pdf_parse_survey_tables(text_block: str) -> list[dict]:
    """Find every 'Table N : <stmt>' block and parse TRUE/FALSE/TOTAL counts."""
    out: list[dict] = []
    for m in _PDF_SURVEY_TABLE_RE.finditer(text_block):
        g = m.groups()

        def first_pair(a, b, c, d) -> tuple[int, int]:
            if a is not None and b is not None:
                return int(a), int(b)
            return int(c or 0), int(d or 0)

        true_n,  true_pct  = first_pair(*g[2:6])
        false_n, false_pct = first_pair(*g[6:10])
        total_n, total_pct = first_pair(*g[10:14])
        out.append({
            "n":         int(g[0]),
            "statement": g[1].strip(),
            "true_n":    true_n,
            "true_pct":  true_pct,
            "false_n":   false_n,
            "false_pct": false_pct,
            "total_n":   total_n,
            "total_pct": total_pct,
        })
    return out


def _pdf_render_survey_chart(table: dict) -> bytes | None:
    """Render a clean bar chart with non-overlapping legend. Returns PNG bytes."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import Patch
        import io as _io

        plt.rcParams["font.family"] = "Times New Roman"
        fig, ax = plt.subplots(figsize=(5.5, 3.4), dpi=160)
        labels = ["TRUE", "FALSE"]
        values = [table["true_n"], table["false_n"]]
        colors = ["#2563eb", "#dc2626"]
        bars   = ax.bar(labels, values, color=colors, width=0.55)
        for bar, val, pct in zip(bars, values,
                                  [table["true_pct"], table["false_pct"]]):
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height() + 0.5,
                f"{val} ({pct}%)",
                ha="center", va="bottom",
                fontsize=11, fontweight="bold",
            )
        ax.set_ylim(0, max(values) * 1.45 if max(values) else 1)
        ax.set_ylabel("No. of Participants", fontsize=12)
        ax.set_xlabel("Response", fontsize=12)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(axis="y", alpha=0.25, linestyle=":")
        ax.legend(
            handles=[Patch(facecolor=colors[0], label="TRUE"),
                     Patch(facecolor=colors[1], label="FALSE")],
            loc="center left",
            bbox_to_anchor=(1.02, 0.5),
            frameon=True,
            fontsize=11,
            title="Response",
            title_fontsize=11,
        )
        fig.tight_layout()
        buf = _io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception as exc:
        logger.warning("survey chart render failed: %s", exc)
        return None


# ---------------------------------------------------------------------------
# Page-as-image rendering for Canva-designed form pages
# ---------------------------------------------------------------------------

_FORM_PAGE_KEYWORDS: tuple[str, ...] = (
    "TRAINEE EVALUATION", "PARTICIPANT EVALUATION",
    "EVALUATION FORM", "ASSESSMENT FORM",
)


def _pdf_render_page_png(pdf, page_no_0based: int,
                          dpi_matrix: int = 2) -> bytes:
    """Rasterize a PDF page to PNG bytes. Used for form pages whose
    Canva-designed layout can't be reliably reconstructed from text blocks.
    """
    import fitz
    page = pdf[page_no_0based]
    pix = page.get_pixmap(matrix=fitz.Matrix(dpi_matrix, dpi_matrix))
    return pix.tobytes("png")


def _pdf_is_form_page(blocks: list[tuple]) -> tuple[bool, str]:
    """Return (is_form, heading) if any block matches form keywords."""
    for text, _max, _dom, _bold, _align in blocks:
        u = text.strip(" ?:.").upper()
        for kw in _FORM_PAGE_KEYWORDS:
            if kw in u and len(text) < 80:
                return True, text.strip(" ?:.")
    return False, ""


# ---------------------------------------------------------------------------
# PDF -> formatted DOCX (deterministic, no LLM)
# ---------------------------------------------------------------------------

_FRONT_MATTER_PAGE_LIMIT: int = 4   # cover + ack + decl + ToC (typical max)


def format_fix_pdf(pdf_bytes: bytes, university_id: str,
                    skip_pages: list[int] | None = None) -> bytes:
    """Reformat a PDF into a university-styled DOCX.

    skip_pages: optional list of 1-based PDF page numbers to drop entirely
                (e.g. redundant 'Report Title' duplicate-cover pages).

    PDFs have no Word styles, so we build a fresh DOCX from extracted content:
      - Block-based positional extraction (true visual reading order)
      - Sentence merging for Canva line breaks
      - Heading detection via ALL_CAPS short lines and 'Chapter N' patterns
      - Decorative-image filter (page backgrounds + repeated icons dropped)
      - Embed remaining images inline at their source page
      - Apply university styles + black-text font discipline
    No LLM in path, no content invention.
    """
    skip_set: set[int] = {p - 1 for p in (skip_pages or [])}
    import fitz
    from docx.shared import Inches, RGBColor

    config = load_university_config(university_id)
    pdf    = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        decorative_xrefs = _pdf_build_decorative_xrefs(pdf)

        doc = Document()
        _apply_university_styles(doc, config)

        body_font = config["body_font"]
        body_size = config["body_size_pt"]
        spacing   = config["line_spacing"]
        sp_before = config["para_space_before_pt"]
        sp_after  = config["para_space_after_pt"]

        for section in doc.sections:
            _apply_margins_to_section(section, config)
            _set_section_page_numbering(section, fmt="decimal", start=1)

        def _add_body(text: str) -> None:
            p  = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing      = spacing
            pf.space_before      = Pt(sp_before)
            pf.space_after       = Pt(sp_after)
            r = p.add_run(text)
            r.font.name = body_font
            r.font.size = Pt(body_size)

        def _add_heading(text: str, level: int = 1) -> None:
            h = doc.add_heading(text, level=level)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _add_kv_table(pairs: list[tuple[str, str]]) -> None:
            tbl = doc.add_table(rows=len(pairs), cols=2)
            tbl.autofit = False
            for r_idx, (k, v) in enumerate(pairs):
                cells = tbl.rows[r_idx].cells
                cells[0].width = Cm(6.5)
                cells[1].width = Cm(9.5)
                cells[0].text  = k.rstrip(":.").strip()
                cells[1].text  = v.strip()
                for run in cells[0].paragraphs[0].runs:
                    run.bold = True

        def _add_image_from_xref(xref: int) -> None:
            try:
                pix = fitz.Pixmap(pdf, xref)
                if pix.alpha or pix.n > 4:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                buf = pix.tobytes("png")
                pix = None
                doc.add_picture(io.BytesIO(buf), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as exc:
                logger.warning("format_fix_pdf image embed failed: %s", exc)

        body_pt = _pdf_estimate_body_size(pdf)

        # First pass: detect survey-table pages (so we can render them as
        # native Word tables + regenerated charts instead of flat text).
        survey_pages: dict[int, list[dict]] = {}
        for page_no in range(len(pdf)):
            page_text = "\n".join(
                t for t, *_ in _pdf_blocks_with_font(pdf, page_no)
            )
            tables = _pdf_parse_survey_tables(page_text)
            if tables:
                survey_pages[page_no] = tables

        def _add_survey_block(t: dict) -> None:
            # Caption above the table
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cap.add_run(f"Table {t['n']}: {t['statement']}")
            r.bold = True
            # Word table
            tbl = doc.add_table(rows=4, cols=4)
            tbl.style = "Table Grid"
            for i, head in enumerate(
                ["SL. NO", "Variables", "No. of Participants",
                 "Percentage %"]
            ):
                c = tbl.rows[0].cells[i]
                c.text = head
                for r_ in c.paragraphs[0].runs:
                    r_.bold = True
            rows = [
                ("1", "TRUE",  str(t["true_n"]),  str(t["true_pct"])),
                ("2", "FALSE", str(t["false_n"]), str(t["false_pct"])),
                ("",  "TOTAL", str(t["total_n"]), str(t["total_pct"])),
            ]
            for r_idx, row in enumerate(rows, start=1):
                for c_idx, val in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = val
            doc.add_paragraph()
            # Regenerated bar chart
            png_bytes = _pdf_render_survey_chart(t)
            if png_bytes:
                doc.add_picture(io.BytesIO(png_bytes), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap2 = doc.add_paragraph()
                cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = cap2.add_run(f"Figure {t['n']}: response distribution")
                rr.italic = True
            doc.add_paragraph()

        in_form_section = False
        for page_no in range(len(pdf)):
            # User-requested page drop (e.g. redundant Report Title page)
            if page_no in skip_set:
                continue
            # Survey-table page → emit table + chart for each parsed table,
            # then SKIP normal text/image rendering for this page (avoids
            # the duplicate text dump of the same table data).
            if page_no in survey_pages:
                in_form_section = False
                for t in survey_pages[page_no]:
                    _add_survey_block(t)
                continue
            blocks = _pdf_blocks_with_font(pdf, page_no)

            # Inter-block sentence merging: Word PDFs put each visual line
            # as a separate block. Without this pass, body paragraphs render
            # as one-line-per-paragraph. Merge consecutive body-sized blocks
            # whose previous line doesn't end with sentence-terminal punct.
            blocks = _merge_body_blocks(blocks, body_pt)

            # Front-matter pages take priority over form detection (ToC may
            # list "TRAINEE EVALUATION" / "PARTICIPANT EVALUATION" as plain
            # entries — must not trigger form-page rendering).
            if page_no < _FRONT_MATTER_PAGE_LIMIT:
                in_form_section = False
                for text, max_sz, _dom, _bold, blk_align in blocks:
                    if not text.strip() or _PDF_PAGE_NUM_RE.match(text):
                        continue
                    is_title_size = max_sz > body_pt + 2
                    if is_title_size:
                        p = doc.add_heading(text.strip(" ?:."), level=1)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(text)
                    if blk_align == "center":
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif blk_align == "right":
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            # Form page (Trainee/Participant Evaluation) → render whole page
            # as image since Canva forms have no border rules and positional
            # column data interleaves badly as flowing text. Continuation
            # pages (no keyword but still inside the form section) also get
            # rendered as image.
            is_form, form_heading = _pdf_is_form_page(blocks)
            # End the form section if we hit a real h1 chapter divider
            if not is_form and in_form_section:
                has_real_h1 = any(
                    msz > body_pt + 8
                    and txt[:1].isupper()
                    and len(txt) >= 4
                    for txt, msz, *_ in blocks
                )
                if has_real_h1:
                    in_form_section = False

            if is_form or in_form_section:
                if is_form:
                    in_form_section = True
                    if form_heading:
                        _add_heading(form_heading, level=1)
                try:
                    png = _pdf_render_page_png(pdf, page_no)
                    doc.add_picture(io.BytesIO(png), width=Inches(6.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except Exception as exc:
                    logger.warning("form page render failed p%d: %s",
                                    page_no, exc)
                continue
            in_form_section = False

            is_toc = _pdf_is_toc_page(blocks)

            kv_buffer: list[tuple[str, str]] = []

            def _flush_kv() -> None:
                nonlocal kv_buffer
                if len(kv_buffer) >= 3:
                    _add_kv_table(kv_buffer)
                else:
                    for k, v in kv_buffer:
                        _add_body(f"{k}: {v}")
                kv_buffer = []

            for text, max_sz, _dom_sz, _bold, block_align in blocks:
                ln = text
                if _PDF_PAGE_NUM_RE.match(ln):
                    continue

                # Key-value detection (groups consecutive K:V lines)
                kv_m = _PDF_KV_RE.match(ln)
                if kv_m and not _PDF_ALL_CAPS_RE.match(ln):
                    kv_buffer.append((kv_m.group(1), kv_m.group(2)))
                    continue
                elif kv_buffer:
                    _flush_kv()

                cls = _pdf_classify_block(ln, max_sz, body_pt, page_no,
                                           is_toc_page=is_toc)
                if cls == "skip":
                    continue
                if cls == "h1":
                    _add_heading(ln.strip(" ?:"), level=1)
                    continue
                if cls == "h2":
                    h = doc.add_heading(ln.strip(" ?:"), level=2)
                    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    continue

                if _PDF_BULLET_RE.match(ln):
                    cleaned = _PDF_BULLET_RE.sub("", ln).strip()
                    p = doc.add_paragraph(style="List Bullet")
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.add_run(cleaned)
                    continue

                m = _PDF_NUM_LIST_RE.match(ln)
                if m:
                    cleaned = _PDF_NUM_LIST_RE.sub("", ln).strip()
                    p = doc.add_paragraph(style="List Number")
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.add_run(cleaned)
                    continue

                _add_body(ln)

            if kv_buffer:
                _flush_kv()

            page = pdf[page_no]
            for info in page.get_images(full=True):
                xref = info[0]
                if xref in decorative_xrefs:
                    continue
                _add_image_from_xref(xref)

        def _force_black_tnr(run, *, size: float, bold: bool | None = None) -> None:
            run.font.name = body_font
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"),    body_font)
            rFonts.set(qn("w:hAnsi"),    body_font)
            rFonts.set(qn("w:cs"),       body_font)
            rFonts.set(qn("w:eastAsia"), body_font)
            run.font.size = Pt(size)
            if bold is not None:
                run.font.bold = bold
            run.font.italic    = False
            run.font.underline = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            for color_el in rPr.findall(qn("w:color")):
                rPr.remove(color_el)
            new_color = OxmlElement("w:color")
            new_color.set(qn("w:val"), "000000")
            rPr.append(new_color)

        def _walk(para) -> None:
            sn = (para.style.name or "") if para.style else ""
            is_heading = sn.startswith("Heading")
            sz = body_size + 2 if is_heading else body_size
            bd = True if is_heading else None
            for run in para.runs:
                _force_black_tnr(run, size=sz, bold=bd)

        for para in doc.paragraphs:
            _walk(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _walk(para)

        # Publish-ready typography pass — overrides legacy font discipline
        # above with the strict spec (orphan control + even block spacing).
        _publish_ready_pass(doc, config)

        _add_footer_cta(doc)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    finally:
        pdf.close()
